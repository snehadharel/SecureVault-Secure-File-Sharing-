##### SecureVault - Secure File Sharing System #####
#### Developed by Sneha Dharel ####

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os
import json
import logging
from datetime import datetime, timedelta
from cryptography.fernet import Fernet, InvalidToken
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.asymmetric import rsa, padding
from cryptography.hazmat.primitives import serialization
import base64
import re
import hashlib
import threading
import time
import io
import sqlite3
from contextlib import closing
import shutil
import mimetypes
from PIL import Image, ImageFile
from PIL.ExifTags import TAGS, GPSTAGS
import PyPDF2
from docx import Document
import csv
import zipfile
import tarfile
import matplotlib
matplotlib.use('TkAgg')  # Set the backend to TkAgg
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import matplotlib.pyplot as plt

# Custom formatter to handle additional fields in logging page
class CustomFormatter(logging.Formatter):
    def format(self, record):
        if hasattr(record, 'username'):
            record.user = getattr(record, 'username', 'unknown')
        else:
            record.user = 'unknown'
            
        if hasattr(record, 'action_type'):
            record.action = getattr(record, 'action_type', 'unknown')
        else:
            record.action = 'unknown'
            
        if hasattr(record, 'target_file'):
            record.file = getattr(record, 'target_file', 'unknown')
        else:
            record.file = 'unknown'
            
        return super().format(record)

class MetadataExtractor:
    """Extracts metadata from various file types"""
    
    def __init__(self):
        self.supported_formats = {
            'images': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'],
            'documents': ['.pdf', '.docx', '.txt', '.csv', '.xlsx', '.pptx'],
            'archives': ['.zip', '.tar', '.gz', '.rar'],
            'media': ['.mp3', '.mp4', '.avi', '.mkv', '.wav']
        }
    
    def extract_all_metadata(self, file_path):
        """Extract metadata from any supported file type"""
        if not os.path.exists(file_path):
            return {"error": "File not found"}
        
        file_type = self._get_file_type(file_path)
        metadata = {
            "basic": self._extract_basic_metadata(file_path),
            "file_type": file_type,
            "extensions": []
        }
        
        # Extract specific metadata based on file types
        if file_type == "image":
            metadata.update(self._extract_image_metadata(file_path))
            metadata["extensions"].append("image_metadata")
        elif file_type == "pdf":
            metadata.update(self._extract_pdf_metadata(file_path))
            metadata["extensions"].append("pdf_metadata")
        elif file_type == "document":
            metadata.update(self._extract_document_metadata(file_path))
            metadata["extensions"].append("document_metadata")
        elif file_type == "archive":
            metadata.update(self._extract_archive_metadata(file_path))
            metadata["extensions"].append("archive_metadata")
        elif file_type == "csv":
            metadata.update(self._extract_csv_metadata(file_path))
            metadata["extensions"].append("csv_metadata")
        
        # Extract cryptographic metadata if it's for an encrypted file
        if file_path.endswith('.encrypted'):
            metadata.update(self._extract_crypto_metadata(file_path))
            metadata["extensions"].append("crypto_metadata")
        
        return metadata
    
    def _get_file_type(self, file_path):
        """Determine file type based on extension"""
        ext = os.path.splitext(file_path)[1].lower()
        
        for file_type, extensions in self.supported_formats.items():
            if ext in extensions:
                if file_type == 'images':
                    return 'image'
                elif file_type == 'documents':
                    if ext == '.pdf':
                        return 'pdf'
                    elif ext == '.csv':
                        return 'csv'
                    else:
                        return 'document'
                elif file_type == 'archives':
                    return 'archive'
                elif file_type == 'media':
                    return 'media'
        
        return 'unknown'
    
    def _extract_basic_metadata(self, file_path):
        """Extract basic file metadata"""
        try:
            stat = os.stat(file_path)
            return {
                "filename": os.path.basename(file_path),
                "file_path": file_path,
                "size_bytes": stat.st_size,
                "size_human": self._human_readable_size(stat.st_size),
                "created": time.ctime(stat.st_ctime),
                "modified": time.ctime(stat.st_mtime),
                "accessed": time.ctime(stat.st_atime),
                "file_extension": os.path.splitext(file_path)[1],
                "mime_type": mimetypes.guess_type(file_path)[0] or "unknown",
                "permissions": oct(stat.st_mode)[-3:],
                "inode": stat.st_ino,
                "device": stat.st_dev
            }
        except Exception as e:
            return {"error": f"Failed to extract basic metadata: {str(e)}"}
    
    def _extract_image_metadata(self, file_path):
        """Extract EXIF and image metadata"""
        metadata = {"exif": {}, "image_info": {}}
        
        try:
            with Image.open(file_path) as img:
                # Basic image information
                metadata["image_info"] = {
                    "format": img.format,
                    "mode": img.mode,
                    "size": img.size,
                    "width": img.width,
                    "height": img.height,
                    "is_animated": getattr(img, "is_animated", False),
                    "n_frames": getattr(img, "n_frames", 1)
                }
                
                # EXIF data info
                exif_data = img._getexif()
                if exif_data:
                    for tag_id, value in exif_data.items():
                        tag = TAGS.get(tag_id, tag_id)
                        if tag == "GPSInfo":
                            gps_data = self._extract_gps_info(value)
                            metadata["exif"]["GPSInfo"] = gps_data
                        else:
                            try:
                                # Try to decode bytes to string
                                if isinstance(value, bytes):
                                    try:
                                        value = value.decode('utf-8', errors='ignore')
                                    except:
                                        value = str(value)
                                metadata["exif"][tag] = value
                            except:
                                metadata["exif"][tag] = str(value)
                
                # Color profile
                if hasattr(img, 'info') and 'icc_profile' in img.info:
                    metadata["image_info"]["has_icc_profile"] = True
                
        except Exception as e:
            metadata["error"] = f"Failed to extract image metadata: {str(e)}"
        
        return metadata
    
    def _extract_gps_info(self, gps_info):
        """Extract GPS information from EXIF data"""
        gps_data = {}
        try:
            for key in gps_info.keys():
                decoded_key = GPSTAGS.get(key, key)
                value = gps_info[key]
                
                if decoded_key in ["GPSLatitude", "GPSLongitude"]:
                    # Convert to decimal degrees
                    degrees = value[0][0] / value[0][1]
                    minutes = value[1][0] / value[1][1]
                    seconds = value[2][0] / value[2][1]
                    decimal = degrees + (minutes / 60.0) + (seconds / 3600.0)
                    
                    # Check direction
                    direction_key = "GPSLatitudeRef" if decoded_key == "GPSLatitude" else "GPSLongitudeRef"
                    direction = gps_info.get(
                        next(k for k, v in GPSTAGS.items() if v == direction_key),
                        'N' if decoded_key == "GPSLatitude" else 'E'
                    )
                    
                    if direction in ['S', 'W']:
                        decimal = -decimal
                    
                    gps_data[decoded_key] = {
                        "raw": value,
                        "decimal": decimal,
                        "direction": direction
                    }
                else:
                    gps_data[decoded_key] = value
        except Exception as e:
            gps_data["error"] = f"Failed to extract GPS info: {str(e)}"
        
        return gps_data
    
    def _extract_pdf_metadata(self, file_path):
        """Extract PDF metadata"""
        metadata = {"pdf_info": {}, "document_info": {}}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Document info
                if pdf_reader.metadata:
                    for key, value in pdf_reader.metadata.items():
                        if value:
                            metadata["document_info"][key] = value
                
                # PDF structure info
                metadata["pdf_info"] = {
                    "pages": len(pdf_reader.pages),
                    "encrypted": pdf_reader.is_encrypted,
                    "has_outline": len(pdf_reader.outline) > 0 if hasattr(pdf_reader, 'outline') else False,
                }
                
                # Extract text from first page (preview)
                try:
                    if len(pdf_reader.pages) > 0:
                        first_page = pdf_reader.pages[0]
                        text = first_page.extract_text()
                        if text:
                            metadata["preview"] = text[:500] + "..." if len(text) > 500 else text
                except:
                    pass
                
        except Exception as e:
            metadata["error"] = f"Failed to extract PDF metadata: {str(e)}"
        
        return metadata
    
    def _extract_document_metadata(self, file_path):
        """Extract metadata from documents (DOCX, TXT, etc.)"""
        metadata = {"document_info": {}}
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.docx':
                doc = Document(file_path)
                metadata["document_info"] = {
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                    "sections": len(doc.sections),
                    "core_properties": {
                        "title": doc.core_properties.title,
                        "author": doc.core_properties.author,
                        "created": doc.core_properties.created,
                        "modified": doc.core_properties.modified,
                        "last_modified_by": doc.core_properties.last_modified_by,
                        "revision": doc.core_properties.revision,
                        "category": doc.core_properties.category,
                        "keywords": doc.core_properties.keywords,
                        "comments": doc.core_properties.comments
                    }
                }
                
                # Extract some content for preview
                if len(doc.paragraphs) > 0:
                    content = "\n".join([p.text for p in doc.paragraphs[:10] if p.text.strip()])
                    if content:
                        metadata["preview"] = content[:500] + "..." if len(content) > 500 else content
            
            elif ext == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                    metadata["document_info"] = {
                        "encoding_guessed": "utf-8",
                        "line_count": content.count('\n') + 1,
                        "word_count": len(content.split()),
                        "character_count": len(content)
                    }
                    if content:
                        metadata["preview"] = content[:500] + "..." if len(content) > 500 else content
        
        except Exception as e:
            metadata["error"] = f"Failed to extract document metadata: {str(e)}"
        
        return metadata
    
    def _extract_csv_metadata(self, file_path):
        """Extract metadata from CSV files"""
        metadata = {"csv_info": {}}
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                # Try to detect delimiter
                sample = f.read(1024)
                f.seek(0)
                
                delimiters = [',', ';', '\t', '|']
                delimiter_counts = {d: sample.count(d) for d in delimiters}
                detected_delimiter = max(delimiter_counts.items(), key=lambda x: x[1])[0]
                
                csv_reader = csv.reader(f, delimiter=detected_delimiter)
                rows = list(csv_reader)
                
                if rows:
                    metadata["csv_info"] = {
                        "delimiter": detected_delimiter,
                        "delimiter_name": {
                            ',': 'comma',
                            ';': 'semicolon',
                            '\t': 'tab',
                            '|': 'pipe'
                        }.get(detected_delimiter, 'unknown'),
                        "row_count": len(rows),
                        "column_count": len(rows[0]) if rows else 0,
                        "headers": rows[0] if rows else [],
                        "sample_data": rows[1:6] if len(rows) > 1 else []
                    }
        
        except Exception as e:
            metadata["error"] = f"Failed to extract CSV metadata: {str(e)}"
        
        return metadata
    
    def _extract_archive_metadata(self, file_path):
        """Extract metadata from archive files"""
        metadata = {"archive_info": {}}
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.zip':
                with zipfile.ZipFile(file_path, 'r') as zipf:
                    metadata["archive_info"] = {
                        "format": "ZIP",
                        "file_count": len(zipf.namelist()),
                        "files": zipf.namelist()[:20],  # First 20 files
                        "total_size": sum(zipf.getinfo(f).file_size for f in zipf.namelist()),
                        "compressed_size": os.path.getsize(file_path),
                        "comment": zipf.comment.decode('utf-8', errors='ignore') if zipf.comment else None
                    }
            
            elif ext in ['.tar', '.gz']:
                mode = 'r:gz' if file_path.endswith('.gz') else 'r'
                with tarfile.open(file_path, mode) as tarf:
                    members = tarf.getmembers()
                    metadata["archive_info"] = {
                        "format": "TAR" + (".GZ" if mode == 'r:gz' else ""),
                        "file_count": len(members),
                        "files": [m.name for m in members[:20]],  # First 20 files
                        "total_size": sum(m.size for m in members),
                        "compressed_size": os.path.getsize(file_path)
                    }
        
        except Exception as e:
            metadata["error"] = f"Failed to extract archive metadata: {str(e)}"
        
        return metadata
    
    def _extract_crypto_metadata(self, file_path):
        """Extract metadata from encrypted files"""
        metadata = {"crypto_info": {}}
        
        try:
            file_size = os.path.getsize(file_path)
            with open(file_path, 'rb') as f:
                # Read first 1KB to analyze
                header = f.read(1024)
                
                metadata["crypto_info"] = {
                    "file_type": "Encrypted file",
                    "header_size": len(header),
                    "is_encrypted": True,
                    "likely_format": self._guess_crypto_format(header)
                }
                
                # Try to detect Fernet format
                if b'gAAAAA' in header[:50]:
                    metadata["crypto_info"]["likely_algorithm"] = "Fernet (AES-128)"
        
        except Exception as e:
            metadata["error"] = f"Failed to extract crypto metadata: {str(e)}"
        
        return metadata
    
    def _guess_crypto_format(self, header):
        """Guess the cryptographic format from header"""
        # Common encryption headers
        if header.startswith(b'Salted__'):
            return "OpenSSL encrypted"
        elif header.startswith(b'U2FsdGVkX1'):
            return "OpenSSL salted"
        elif b'BEGIN PGP MESSAGE' in header[:100]:
            return "PGP/GPG encrypted"
        elif b'-----BEGIN' in header[:50]:
            return "PEM format"
        else:
            return "Unknown/Generic encryption"
    
    def _human_readable_size(self, size_bytes):
        """Convert bytes to human readable format"""
        if size_bytes == 0:
            return "0 B"
        
        units = ["B", "KB", "MB", "GB", "TB"]
        i = 0
        while size_bytes >= 1024 and i < len(units) - 1:
            size_bytes /= 1024.0
            i += 1
        
        return f"{size_bytes:.2f} {units[i]}"
    
    def export_metadata_to_json(self, metadata, output_path):
        """Export metadata to JSON file"""
        try:
            with open(output_path, 'w') as f:
                json.dump(metadata, f, indent=4, default=str)
            return True, output_path
        except Exception as e:
            return False, str(e)
    
    def get_supported_formats(self):
        """Get list of supported file formats"""
        all_formats = []
        for formats in self.supported_formats.values():
            all_formats.extend(formats)
        return sorted(set(all_formats))

class UserJsonManager:
    """Manages user data in JSON file format"""
    def __init__(self, json_path="users.json"):
        self.json_path = json_path
        self.init_json_file()
    
    def init_json_file(self):
        """Initialize JSON file if it doesn't exist"""
        if not os.path.exists(self.json_path):
            with open(self.json_path, 'w') as f:
                json.dump({"users": [], "last_updated": datetime.now().strftime('%Y-%m-%d %H:%M:%S')}, f, indent=4)
    
    def load_users(self):
        """Load users from JSON file"""
        try:
            with open(self.json_path, 'r') as f:
                data = json.load(f)
            return data.get("users", [])
        except (FileNotFoundError, json.JSONDecodeError):
            return []
    
    def save_users(self, users):
        """Save users to JSON file"""
        try:
            data = {
                "users": users,
                "last_updated": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "total_users": len(users)
            }
            with open(self.json_path, 'w') as f:
                json.dump(data, f, indent=4)
            return True
        except Exception as e:
            print(f"Error saving to JSON: {e}")
            return False
    
    def add_user(self, user_data):
        """Add a new user to JSON"""
        users = self.load_users()
        
        # Check if user already exists
        for user in users:
            if user.get("email") == user_data["email"]:
                return False
        
        users.append(user_data)
        return self.save_users(users)
    
    def get_user(self, email):
        """Get user by email from JSON"""
        users = self.load_users()
        for user in users:
            if user.get("email") == email:
                return user
        return None
    
    def update_user(self, email, updated_data):
        """Update user information in JSON"""
        users = self.load_users()
        updated = False
        
        for i, user in enumerate(users):
            if user.get("email") == email:
                users[i].update(updated_data)
                updated = True
                break
        
        if updated:
            return self.save_users(users)
        return False
    
    def delete_user(self, email):
        """Remove user from JSON (soft delete)"""
        users = self.load_users()
        updated_users = []
        deleted = False
        
        for user in users:
            if user.get("email") == email:
                # Soft delete by marking as inactive
                user["is_active"] = 0
                deleted = True
            updated_users.append(user)
        
        if deleted:
            return self.save_users(updated_users)
        return False
    
    def get_all_users(self):
        """Get all active users from JSON"""
        users = self.load_users()
        return [user for user in users if user.get("is_active", 1) == 1]
    
    def sync_from_database(self, db_manager):
        """Sync users from database to JSON file"""
        try:
            db_users = db_manager.get_all_users()
            json_users = []
            
            for db_user in db_users:
                user_data = {
                    "email": db_user["email"],
                    "password_hash": db_user["password_hash"],
                    "role": db_user["role"],
                    "security_question": db_user["security_question"],
                    "security_answer_hash": db_user["security_answer_hash"],
                    "created_at": db_user["created_at"],
                    "last_login": db_user["last_login"],
                    "is_active": db_user["is_active"]
                }
                json_users.append(user_data)
            
            return self.save_users(json_users)
        except Exception as e:
            print(f"Error syncing from database: {e}")
            return False
    
    def export_to_file(self, export_path):
        """Export users to a specified JSON file"""
        users = self.load_users()
        return self.save_users_to_path(users, export_path)
    
    def save_users_to_path(self, users, path):
        """Save users to a specific path"""
        try:
            data = {
                "users": users,
                "export_date": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "exported_by": "system",
                "total_users": len(users)
            }
            with open(path, 'w') as f:
                json.dump(data, f, indent=4)
            return True
        except Exception as e:
            print(f"Error exporting to file: {e}")
            return False
    
    def import_from_file(self, import_path):
        """Import users from a JSON file"""
        try:
            with open(import_path, 'r') as f:
                data = json.load(f)
            
            imported_users = data.get("users", [])
            current_users = self.load_users()
            
            # Merge users (avoid duplicates)
            current_emails = {user["email"] for user in current_users}
            
            for user in imported_users:
                if user["email"] not in current_emails:
                    current_users.append(user)
            
            return self.save_users(current_users)
        except Exception as e:
            print(f"Error importing from file: {e}")
            return False

class DatabaseManager:
    """Manages SQLite database operations"""
    def __init__(self, db_path="secure_file_sharing.db"):
        self.db_path = db_path
        self.init_database()
    
    def get_connection(self):
        """Get a database connection"""
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        return conn
    
    def init_database(self):
        """Initialize database tables"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            
            # Users table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS users (
                    email TEXT PRIMARY KEY,
                    password_hash TEXT NOT NULL,
                    role TEXT NOT NULL CHECK(role IN ('admin', 'user')),
                    security_question TEXT NOT NULL,
                    security_answer_hash TEXT NOT NULL,
                    created_at TIMESTAMP NOT NULL,
                    last_login TIMESTAMP,
                    is_active INTEGER DEFAULT 1 CHECK(is_active IN (0, 1))
                )
            ''')
            
            # Public keys table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS public_keys (
                    email TEXT PRIMARY KEY,
                    public_key TEXT NOT NULL,
                    upload_date TIMESTAMP NOT NULL,
                    last_updated TIMESTAMP NOT NULL,
                    FOREIGN KEY (email) REFERENCES users(email) ON DELETE CASCADE
                )
            ''')
            
            # Shared files table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS shared_files (
                    file_id TEXT PRIMARY KEY,
                    filename TEXT NOT NULL,
                    shared_by TEXT NOT NULL,
                    shared_with TEXT NOT NULL,
                    shared_date TIMESTAMP NOT NULL,
                    encrypted_file TEXT NOT NULL,
                    encrypted_aes_key TEXT NOT NULL,
                    aes_key_plain TEXT,
                    downloaded INTEGER DEFAULT 0 CHECK(downloaded IN (0, 1)),
                    FOREIGN KEY (shared_by) REFERENCES users(email) ON DELETE CASCADE,
                    FOREIGN KEY (shared_with) REFERENCES users(email) ON DELETE CASCADE
                )
            ''')
            
            # Login statistics table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS login_stats (
                    stat_id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_email TEXT NOT NULL,
                    login_time TIMESTAMP NOT NULL,
                    user_role TEXT NOT NULL,
                    FOREIGN KEY (user_email) REFERENCES users(email) ON DELETE CASCADE
                )
            ''')
            
            # Active sessions table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS active_sessions (
                    session_id TEXT PRIMARY KEY,
                    user_email TEXT NOT NULL,
                    user_role TEXT NOT NULL,
                    login_time TIMESTAMP NOT NULL,
                    last_activity TIMESTAMP NOT NULL,
                    FOREIGN KEY (user_email) REFERENCES users(email) ON DELETE CASCADE
                )
            ''')
            
            # User login counts table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS user_login_counts (
                    email TEXT PRIMARY KEY,
                    login_count INTEGER DEFAULT 0,
                    FOREIGN KEY (email) REFERENCES users(email) ON DELETE CASCADE
                )
            ''')
            
            # Failed login attempts table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS failed_login_attempts (
                    email TEXT PRIMARY KEY,
                    attempts INTEGER DEFAULT 0,
                    last_attempt TIMESTAMP NOT NULL,
                    FOREIGN KEY (email) REFERENCES users(email) ON DELETE CASCADE
                )
            ''')
            
            # Create indexes for better performance
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_shared_files_by ON shared_files(shared_by)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_shared_files_with ON shared_files(shared_with)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_login_stats_email ON login_stats(user_email)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_active_sessions_email ON active_sessions(user_email)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_users_role ON users(role)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_users_active ON users(is_active)')
            
            conn.commit()
    
    def add_user(self, email, password_hash, role, security_question, security_answer_hash):
        """Add a new user to the database"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            try:
                cursor.execute('''
                    INSERT INTO users (email, password_hash, role, security_question, 
                                     security_answer_hash, created_at, last_login, is_active)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''', (email, password_hash, role, security_question, security_answer_hash,
                     datetime.now().strftime('%Y-%m-%d %H:%M:%S'), None, 1))
                conn.commit()
                return True
            except sqlite3.IntegrityError:
                return False
    
    def get_user(self, email):
        """Get user by email"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT * FROM users WHERE email = ? AND is_active = 1', (email,))
            row = cursor.fetchone()
            return dict(row) if row else None
    
    def get_all_users(self):
        """Get all active users"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT * FROM users WHERE is_active = 1 ORDER BY created_at DESC')
            return [dict(row) for row in cursor.fetchall()]
    
    def update_user_last_login(self, email):
        """Update user's last login timestamp"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE users 
                SET last_login = ?
                WHERE email = ?
            ''', (datetime.now().strftime('%Y-%m-%d %H:%M:%S'), email))
            conn.commit()
    
    def update_user_password(self, email, new_password_hash):
        """Update user's password"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE users 
                SET password_hash = ?
                WHERE email = ?
            ''', (new_password_hash, email))
            conn.commit()
            return cursor.rowcount > 0
    
    def delete_user(self, email):
        """Soft delete a user (mark as inactive)"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE users 
                SET is_active = 0 
                WHERE email = ?
            ''', (email,))
            conn.commit()
            return cursor.rowcount > 0
    
    def add_public_key(self, email, public_key):
        """Add or update public key for a user"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            
            # Check if key exists
            cursor.execute('SELECT * FROM public_keys WHERE email = ?', (email,))
            if cursor.fetchone():
                cursor.execute('''
                    UPDATE public_keys 
                    SET public_key = ?, last_updated = ?
                    WHERE email = ?
                ''', (public_key, now, email))
            else:
                cursor.execute('''
                    INSERT INTO public_keys (email, public_key, upload_date, last_updated)
                    VALUES (?, ?, ?, ?)
                ''', (email, public_key, now, now))
            
            conn.commit()
            return True
    
    def get_public_key(self, email):
        """Get public key for a user"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT * FROM public_keys WHERE email = ?', (email,))
            row = cursor.fetchone()
            return dict(row) if row else None
    
    def get_all_public_keys(self):
        """Get all public keys"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT * FROM public_keys ORDER BY last_updated DESC')
            return [dict(row) for row in cursor.fetchall()]
    
    def delete_public_key(self, email):
        """Delete public key for a user"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('DELETE FROM public_keys WHERE email = ?', (email,))
            conn.commit()
            return cursor.rowcount > 0
    
    def add_shared_file(self, file_id, filename, shared_by, shared_with, encrypted_file, encrypted_aes_key, aes_key_plain=None):
        """Add a shared file record"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO shared_files (file_id, filename, shared_by, shared_with, shared_date,
                                        encrypted_file, encrypted_aes_key, aes_key_plain, downloaded)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (file_id, filename, shared_by, shared_with,
                 datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                 encrypted_file, encrypted_aes_key, aes_key_plain, 0))
            conn.commit()
            return True
    
    def get_shared_files_for_user(self, email):
        """Get files shared with a user"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT * FROM shared_files 
                WHERE shared_with = ? 
                ORDER BY shared_date DESC
            ''', (email,))
            return [dict(row) for row in cursor.fetchall()]
    
    def get_shared_files_by_user(self, email):
        """Get files shared by a user"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT * FROM shared_files 
                WHERE shared_by = ? 
                ORDER BY shared_date DESC
            ''', (email,))
            return [dict(row) for row in cursor.fetchall()]
    
    def update_shared_file_downloaded(self, file_id):
        """Mark a shared file as downloaded"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE shared_files 
                SET downloaded = 1 
                WHERE file_id = ?
            ''', (file_id,))
            conn.commit()
            return cursor.rowcount > 0
    
    def delete_shared_file(self, file_id):
        """Delete a shared file record"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('DELETE FROM shared_files WHERE file_id = ?', (file_id,))
            conn.commit()
            return cursor.rowcount > 0
    
    def delete_user_shared_files(self, email):
        """Delete all shared files for a user (both shared by and shared with)"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('DELETE FROM shared_files WHERE shared_by = ? OR shared_with = ?', (email, email))
            deleted_count = cursor.rowcount
            conn.commit()
            return deleted_count
    
    def add_login_record(self, email, role):
        """Add a login record"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO login_stats (user_email, login_time, user_role)
                VALUES (?, ?, ?)
            ''', (email, datetime.now().strftime('%Y-%m-%d %H:%M:%S'), role))
            
            # Update user login count
            cursor.execute('''
                INSERT OR REPLACE INTO user_login_counts (email, login_count)
                VALUES (?, COALESCE((SELECT login_count FROM user_login_counts WHERE email = ?), 0) + 1)
            ''', (email, email))
            
            conn.commit()
            return cursor.lastrowid
    
    def get_recent_logins(self, limit=100):
        """Get recent login records"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT * FROM login_stats 
                ORDER BY login_time DESC 
                LIMIT ?
            ''', (limit,))
            return [dict(row) for row in cursor.fetchall()]
    
    def get_total_logins(self):
        """Get total number of logins"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT COUNT(*) as count FROM login_stats')
            row = cursor.fetchone()
            return row['count'] if row else 0
    
    def get_user_login_count(self, email):
        """Get login count for a specific user"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT login_count FROM user_login_counts WHERE email = ?', (email,))
            row = cursor.fetchone()
            return row['login_count'] if row else 0
    
    def add_active_session(self, session_id, email, role):
        """Add an active session"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            cursor.execute('''
                INSERT OR REPLACE INTO active_sessions (session_id, user_email, user_role, login_time, last_activity)
                VALUES (?, ?, ?, ?, ?)
            ''', (session_id, email, role, now, now))
            conn.commit()
            return True
    
    def update_session_activity(self, email):
        """Update session last activity time"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE active_sessions 
                SET last_activity = ?
                WHERE user_email = ?
            ''', (datetime.now().strftime('%Y-%m-%d %H:%M:%S'), email))
            conn.commit()
            return cursor.rowcount > 0
    
    def get_active_sessions(self):
        """Get all active sessions"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT * FROM active_sessions ORDER BY last_activity DESC')
            return [dict(row) for row in cursor.fetchall()]
    
    def get_active_session_by_user(self, email):
        """Get active session for a specific user"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT * FROM active_sessions WHERE user_email = ?', (email,))
            row = cursor.fetchone()
            return dict(row) if row else None
    
    def remove_active_session(self, email):
        """Remove active session for a user"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('DELETE FROM active_sessions WHERE user_email = ?', (email,))
            conn.commit()
            return cursor.rowcount > 0
    
    def remove_inactive_sessions(self, timeout_minutes=30):
        """Remove sessions inactive for more than timeout_minutes"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            timeout_time = (datetime.now() - timedelta(minutes=timeout_minutes)).strftime('%Y-%m-%d %H:%M:%S')
            cursor.execute('DELETE FROM active_sessions WHERE last_activity < ?', (timeout_time,))
            deleted_count = cursor.rowcount
            conn.commit()
            return deleted_count
    
    def record_failed_login_attempt(self, email):
        """Record a failed login attempt"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            
            cursor.execute('''
                INSERT OR REPLACE INTO failed_login_attempts (email, attempts, last_attempt)
                VALUES (?, COALESCE((SELECT attempts FROM failed_login_attempts WHERE email = ?), 0) + 1, ?)
            ''', (email, email, now))
            
            conn.commit()
            
            # Get current attempt count
            cursor.execute('SELECT attempts FROM failed_login_attempts WHERE email = ?', (email,))
            row = cursor.fetchone()
            return row['attempts'] if row else 1
    
    def get_failed_login_attempts(self, email):
        """Get failed login attempts for a user"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT attempts, last_attempt FROM failed_login_attempts WHERE email = ?', (email,))
            row = cursor.fetchone()
            if row:
                return {'attempts': row['attempts'], 'last_attempt': row['last_attempt']}
            return {'attempts': 0, 'last_attempt': None}
    
    def reset_failed_login_attempts(self, email):
        """Reset failed login attempts for a user"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute('DELETE FROM failed_login_attempts WHERE email = ?', (email,))
            conn.commit()
            return cursor.rowcount > 0
    
    def get_system_statistics(self):
        """Get system statistics"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            
            stats = {}
            
            # Total users
            cursor.execute('SELECT COUNT(*) as count FROM users WHERE is_active = 1')
            stats['total_users'] = cursor.fetchone()['count']
            
            # Total logins
            stats['total_logins'] = self.get_total_logins()
            
            # Files shared
            cursor.execute('SELECT COUNT(*) as count FROM shared_files')
            stats['files_shared'] = cursor.fetchone()['count']
            
            # Currently logged in
            cursor.execute('SELECT COUNT(*) as count FROM active_sessions')
            stats['currently_logged_in'] = cursor.fetchone()['count']
            
            # Users with public keys
            cursor.execute('SELECT COUNT(DISTINCT email) as count FROM public_keys')
            stats['users_with_pubkeys'] = cursor.fetchone()['count']
            
            # Admin count
            cursor.execute('SELECT COUNT(*) as count FROM users WHERE role = "admin" AND is_active = 1')
            stats['admin_count'] = cursor.fetchone()['count']
            
            # User count
            cursor.execute('SELECT COUNT(*) as count FROM users WHERE role = "user" AND is_active = 1')
            stats['user_count'] = cursor.fetchone()['count']
            
            return stats
    
    def backup_database(self, backup_path):
        """Create a backup of the database"""
        try:
            shutil.copy2(self.db_path, backup_path)
            return True
        except Exception as e:
            print(f"Backup failed: {e}")
            return False
    
    def vacuum(self):
        """Optimize database"""
        with closing(self.get_connection()) as conn:
            conn.execute('VACUUM')
            conn.commit()
    
    def execute_query(self, query, params=()):
        """Execute a custom query (for advanced operations)"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute(query, params)
            conn.commit()
            return cursor.rowcount
    
    def fetch_all(self, query, params=()):
        """Fetch all rows from a query"""
        with closing(self.get_connection()) as conn:
            cursor = conn.cursor()
            cursor.execute(query, params)
            return [dict(row) for row in cursor.fetchall()]

class SecureFileSharingApp:
    def __init__(self, root):
        self.root = root
        self.root.title("🔐 SecureVault-Secure File Sharing Tool")
        self.root.geometry("1550x900")
        
        # Initialize database manager
        self.db = DatabaseManager()
        
        # Initialize JSON manager
        self.json_manager = UserJsonManager("users.json")
        
        # Initialize metadata extractor
        self.metadata_extractor = MetadataExtractor()
        
        # Configure theme colors
        self.colors = {
            'primary': '#2E3B4E',       # Dark blue
            'secondary': '#4A90E2',     # Bright blue
            'accent': '#FF6B6B',        # Coral red
            'success': '#50C878',       # Emerald green
            'warning': '#FFA500',       # Orange
            'background': '#F5F7FA',    # Light gray
            'card_bg': '#FFFFFF',       # White
            'text_dark': '#2C3E50',     # Dark text
            'text_light': '#FFFFFF',    # White text
            'sidebar': '#34495E',       # Dark sidebar
            'highlight': '#FFD700',     # Gold
            'gradient1': '#667eea',     # Purple blue
            'gradient2': '#764ba2',     # Purple
        }
        
        # Apply custom styling
        self.setup_styles()
        
        # Center the main window
        self.center_window(self.root, 1550, 900)
        
        # Initialize components
        self.current_user = None
        self.user_role = None
        self.aes_key = None
        self.user_private_key = None
        self.user_public_key = None
        self.max_login_attempts = 5
        self.lockout_time = 300
        
        # Setup directories
        self.shared_files_dir = "shared_files"
        self.user_keys_dir = "user_keys"
        self.user_rsa_keys_dir = "user_rsa_keys"
        os.makedirs(self.shared_files_dir, exist_ok=True)
        os.makedirs(self.user_keys_dir, exist_ok=True)
        os.makedirs(self.user_rsa_keys_dir, exist_ok=True)
        
        # Setup logging
        self.setup_logging()
        
        # Initialize default admin user if not exists
        self.initialize_default_admin()
        
        # Sync database to JSON on startup
        self.sync_database_to_json()
        
        # Create unified authentication frame
        self.create_unified_auth_frame()
        
        # Start session cleanup thread
        self.start_session_cleanup_thread()
    
    def setup_styles(self):
        """Setup custom styles and themes"""
        style = ttk.Style()
        
        # Configure colors for different elements
        style.theme_use('clam')
        
        # Configure main window background
        self.root.configure(bg=self.colors['background'])
        
        # Configure ttk styles
        style.configure('Title.TLabel',
                       font=('Segoe UI', 24, 'bold'),
                       foreground=self.colors['primary'],
                       background=self.colors['background'])
        
        style.configure('Subtitle.TLabel',
                       font=('Segoe UI', 12),
                       foreground=self.colors['secondary'],
                       background=self.colors['background'])
        
        style.configure('Card.TFrame',
                       background=self.colors['card_bg'],
                       relief='raised',
                       borderwidth=2)
        
        style.configure('Primary.TButton',
                       font=('Segoe UI', 10, 'bold'),
                       background=self.colors['secondary'],
                       foreground=self.colors['text_light'],
                       borderwidth=1,
                       focuscolor='none')
        
        style.map('Primary.TButton',
                 background=[('active', self.colors['primary']),
                           ('pressed', self.colors['accent'])])
        
        style.configure('Success.TButton',
                       font=('Segoe UI', 10, 'bold'),
                       background=self.colors['success'],
                       foreground=self.colors['text_light'],
                       borderwidth=1)
        
        style.configure('Danger.TButton',
                       font=('Segoe UI', 10, 'bold'),
                       background=self.colors['accent'],
                       foreground=self.colors['text_light'],
                       borderwidth=1)
        
        style.configure('Warning.TButton',
                       font=('Segoe UI', 10, 'bold'),
                       background=self.colors['warning'],
                       foreground=self.colors['text_dark'],
                       borderwidth=1)
        
        style.configure('Header.TFrame',
                       background=self.colors['primary'])
        
        style.configure('Header.TLabel',
                       font=('Segoe UI', 11),
                       foreground=self.colors['text_light'],
                       background=self.colors['primary'])
        
        style.configure('Accent.TLabel',
                       font=('Segoe UI', 10, 'bold'),
                       foreground=self.colors['accent'],
                       background=self.colors['card_bg'])
        
        # Notebook style
        style.configure('Custom.TNotebook',
                       background=self.colors['background'],
                       borderwidth=0)
        
        style.configure('Custom.TNotebook.Tab',
                       font=('Segoe UI', 10, 'bold'),
                       padding=[20, 10],
                       background=self.colors['card_bg'],
                       foreground=self.colors['text_dark'])
        
        style.map('Custom.TNotebook.Tab',
                 background=[('selected', self.colors['secondary'])],
                 foreground=[('selected', self.colors['text_light'])])
        
        # Entry style
        style.configure('Custom.TEntry',
                       fieldbackground=self.colors['card_bg'],
                       foreground=self.colors['text_dark'],
                       borderwidth=1,
                       relief='solid')
        
        # LabelFrame style
        style.configure('Custom.TLabelframe',
                       background=self.colors['card_bg'],
                       foreground=self.colors['primary'],
                       borderwidth=2,
                       relief='ridge')
        
        style.configure('Custom.TLabelframe.Label',
                       font=('Segoe UI', 11, 'bold'),
                       background=self.colors['card_bg'],
                       foreground=self.colors['primary'])
        
        # Treeview style
        style.configure('Custom.Treeview',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text_dark'],
                       fieldbackground=self.colors['card_bg'],
                       rowheight=25)
        
        style.configure('Custom.Treeview.Heading',
                       font=('Segoe UI', 10, 'bold'),
                       background=self.colors['primary'],
                       foreground=self.colors['text_light'],
                       relief='flat')
        
        style.map('Custom.Treeview',
                 background=[('selected', self.colors['secondary'])],
                 foreground=[('selected', self.colors['text_light'])])
        
        # Custom JSON button styles
        style.configure('JSONPrimary.TButton',
                       font=('Segoe UI', 10, 'bold'),
                       background='#4A6FA5',
                       foreground='white',
                       borderwidth=1)
        
        style.configure('JSONSuccess.TButton',
                       font=('Segoe UI', 10, 'bold'),
                       background='#50C878',
                       foreground='white',
                       borderwidth=1)
        
        style.configure('JSONWarning.TButton',
                       font=('Segoe UI', 10, 'bold'),
                       background='#FFA500',
                       foreground='black',
                       borderwidth=1)
        
        style.configure('JSONDanger.TButton',
                       font=('Segoe UI', 10, 'bold'),
                       background='#DC143C',
                       foreground='white',
                       borderwidth=1)
        
    def center_window(self, window, width, height):
        """Center a window on the screen"""
        screen_width = window.winfo_screenwidth()
        screen_height = window.winfo_screenheight()
        x = (screen_width - width) // 2
        y = (screen_height - height) // 2
        window.geometry(f'{width}x{height}+{x}+{y}')
        
    def setup_logging(self):
        """Setup logging configuration"""
        for handler in logging.root.handlers[:]:
            logging.root.removeHandler(handler)
            
        self.logger = logging.getLogger('SecureFileSharing')
        self.logger.setLevel(logging.INFO)
        
        formatter = CustomFormatter('%(asctime)s - %(user)s - %(action)s - %(file)s')
        
        file_handler = logging.FileHandler('file_sharing.log')
        file_handler.setFormatter(formatter)
        self.logger.addHandler(file_handler)
        
        console_handler = logging.StreamHandler()
        console_handler.setFormatter(formatter)
        self.logger.addHandler(console_handler)
        
    def log_action(self, action, filename=""):
        """Log user actions safely"""
        extra = {
            'username': self.current_user if self.current_user else 'unknown',
            'action_type': action,
            'target_file': filename
        }
        self.logger.info('File operation', extra=extra)
    
    def is_account_locked(self, email):
        """Check if account is locked due to too many failed attempts"""
        failed_attempts = self.db.get_failed_login_attempts(email)
        
        if failed_attempts['attempts'] >= self.max_login_attempts:
            if failed_attempts['last_attempt']:
                try:
                    last_attempt = datetime.strptime(failed_attempts['last_attempt'], '%Y-%m-%d %H:%M:%S')
                    time_diff = (datetime.now() - last_attempt).total_seconds()
                    if time_diff < self.lockout_time:
                        return True
                    else:
                        # Reset after lockout period
                        self.db.reset_failed_login_attempts(email)
                except:
                    pass
        return False
    
    def record_failed_attempt(self, email):
        """Record a failed login attempt"""
        attempts = self.db.record_failed_login_attempt(email)
        return attempts
    
    def validate_password_strength(self, password):
        """Validate password strength"""
        if len(password) < 8:
            return False, "Password must be at least 8 characters long"
        
        if not re.search(r"[A-Z]", password):
            return False, "Password must contain at least one uppercase letter"
        
        if not re.search(r"[a-z]", password):
            return False, "Password must contain at least one lowercase letter"
        
        if not re.search(r"\d", password):
            return False, "Password must contain at least one digit"
        
        if not re.search(r"[!@#$%^&*(),.?\":{}|<>]", password):
            return False, "Password must contain at least one special character"
        
        return True, "Password is strong"
    
    def hash_password(self, password):
        """Hash password using SHA-256"""
        return hashlib.sha256(password.encode()).hexdigest()
    
    def validate_email(self, email):
        """Validate email format"""
        pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        return re.match(pattern, email) is not None
    
    def initialize_default_admin(self):
        """Initialize default admin user in database and JSON"""
        admin_email = "admin@gmail.com"
        admin_data = self.db.get_user(admin_email)
        
        if not admin_data:
            # Create default admin user in database
            password_hash = self.hash_password("Admin@123")
            security_answer_hash = self.hash_password("blue")
            
            success = self.db.add_user(
                email=admin_email,
                password_hash=password_hash,
                role="admin",
                security_question="What is your favorite color?",
                security_answer_hash=security_answer_hash
            )
            
            if success:
                # Also create in JSON
                user_data = {
                    "email": admin_email,
                    "password_hash": password_hash,
                    "role": "admin",
                    "security_question": "What is your favorite color?",
                    "security_answer_hash": security_answer_hash,
                    "created_at": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                    "last_login": None,
                    "is_active": 1
                }
                self.json_manager.add_user(user_data)
                print("Default admin user created in both database and JSON")
            else:
                print("Failed to create default admin user")
    
    def get_public_key_from_repository(self, email):
        """Get public key from database"""
        try:
            key_data = self.db.get_public_key(email)
            if key_data:
                public_key_str = key_data['public_key']
                public_key = serialization.load_pem_public_key(public_key_str.encode())
                return public_key, "Success"
            return None, "Public key not found in repository"
        except Exception as e:
            return None, f"Error loading public key: {str(e)}"
    
    def get_all_public_keys(self):
        """Get all public keys from database"""
        return self.db.get_all_public_keys()
    
    def update_public_key(self, email, public_key_pem):
        """Update existing public key in database"""
        public_key_str = public_key_pem.decode() if isinstance(public_key_pem, bytes) else public_key_pem
        return self.db.add_public_key(email, public_key_str)
    
    def upload_public_key(self, email, public_key_pem):
        """Upload public key to database"""
        public_key_str = public_key_pem.decode() if isinstance(public_key_pem, bytes) else public_key_pem
        return self.db.add_public_key(email, public_key_str)
    
    def update_login_stats(self):
        """Update login statistics when user logs in"""
        if self.current_user and self.user_role:
            self.db.add_login_record(self.current_user, self.user_role)
            self.db.update_user_last_login(self.current_user)
    
    def add_active_session(self):
        """Add current user to active sessions"""
        if self.current_user and self.user_role:
            session_id = f"{self.current_user}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
            self.db.add_active_session(session_id, self.current_user, self.user_role)
    
    def remove_active_session(self):
        """Remove current user from active sessions"""
        if self.current_user:
            self.db.remove_active_session(self.current_user)
    
    def update_session_activity(self):
        """Update last activity time for current session"""
        if self.current_user:
            self.db.update_session_activity(self.current_user)
    
    def cleanup_inactive_sessions(self):
        """Remove sessions inactive for more than 30 minutes"""
        return self.db.remove_inactive_sessions(timeout_minutes=30)
    
    def generate_user_rsa_keys(self, email, password):
        """Generate RSA key pair for a new user and encrypt private key with password"""
        try:
            # Generate RSA key pair
            private_key = rsa.generate_private_key(
                public_exponent=65537,
                key_size=2048
            )
            public_key = private_key.public_key()
            
            # Save public key locally
            public_key_path = os.path.join(self.user_rsa_keys_dir, f"{email}_public.pem")
            with open(public_key_path, 'wb') as f:
                f.write(public_key.public_bytes(
                    encoding=serialization.Encoding.PEM,
                    format=serialization.PublicFormat.SubjectPublicKeyInfo
                ))
            
            # Encrypt private key with user's password
            private_key_bytes = private_key.private_bytes(
                encoding=serialization.Encoding.PEM,
                format=serialization.PrivateFormat.PKCS8,
                encryption_algorithm=serialization.NoEncryption()
            )
            
            # Use password as AES key (32 bytes for Fernet)
            password_hash = hashlib.sha256(password.encode()).digest()[:32]
            fernet = Fernet(base64.urlsafe_b64encode(password_hash))
            encrypted_private_key = fernet.encrypt(private_key_bytes)
            
            # Save encrypted private key locally
            private_key_path = os.path.join(self.user_rsa_keys_dir, f"{email}_private.enc")
            with open(private_key_path, 'wb') as f:
                f.write(encrypted_private_key)
            
            # Upload public key to database
            public_key_pem = public_key.public_bytes(
                encoding=serialization.Encoding.PEM,
                format=serialization.PublicFormat.SubjectPublicKeyInfo
            )
            success = self.upload_public_key(email, public_key_pem)
            
            if not success:
                print(f"Warning: Could not upload public key to database: {email}")
            
            return True
        except Exception as e:
            print(f"Error generating RSA keys for {email}: {e}")
            return False
    
    def load_user_rsa_keys(self, email, password):
        """Load user's RSA keys using password to decrypt private key"""
        try:
            # Load public key from local storage first
            public_key_path = os.path.join(self.user_rsa_keys_dir, f"{email}_public.pem")
            if os.path.exists(public_key_path):
                with open(public_key_path, 'rb') as f:
                    self.user_public_key = serialization.load_pem_public_key(f.read())
            else:
                # Try to get from database if not local
                public_key, message = self.get_public_key_from_repository(email)
                if public_key:
                    self.user_public_key = public_key
                else:
                    return False, "Public key not found"
            
            # Load and decrypt private key
            private_key_path = os.path.join(self.user_rsa_keys_dir, f"{email}_private.enc")
            if not os.path.exists(private_key_path):
                return False, "Private key not found"
            
            with open(private_key_path, 'rb') as f:
                encrypted_private_key = f.read()
            
            # Decrypt private key using password
            password_hash = hashlib.sha256(password.encode()).digest()[:32]
            fernet = Fernet(base64.urlsafe_b64encode(password_hash))
            
            try:
                decrypted_private_key = fernet.decrypt(encrypted_private_key)
                self.user_private_key = serialization.load_pem_private_key(
                    decrypted_private_key,
                    password=None
                )
                return True, "Success"
            except InvalidToken:
                return False, "Incorrect password for private key decryption"
            
        except Exception as e:
            return False, f"Error loading RSA keys: {str(e)}"
    
    def get_recipient_public_key(self, recipient_email):
        """Get recipient's public key from database"""
        return self.get_public_key_from_repository(recipient_email)
    
    def encrypt_aes_key_with_rsa(self, aes_key, recipient_public_key):
        """Encrypt AES key with recipient's RSA public key"""
        try:
            encrypted_key = recipient_public_key.encrypt(
                aes_key,
                padding.OAEP(
                    mgf=padding.MGF1(algorithm=hashes.SHA256()),
                    algorithm=hashes.SHA256(),
                    label=None
                )
            )
            return base64.b64encode(encrypted_key).decode(), "Success"
        except Exception as e:
            return None, f"Error encrypting AES key: {str(e)}"
    
    def decrypt_aes_key_with_rsa(self, encrypted_aes_key_b64):
        """Decrypt AES key with user's RSA private key"""
        try:
            if not self.user_private_key:
                return None, "User private key not loaded"
            
            encrypted_aes_key = base64.b64decode(encrypted_aes_key_b64)
            decrypted_key = self.user_private_key.decrypt(
                encrypted_aes_key,
                padding.OAEP(
                    mgf=padding.MGF1(algorithm=hashes.SHA256()),
                    algorithm=hashes.SHA256(),
                    label=None
                )
            )
            return decrypted_key, "Success"
        except Exception as e:
            return None, f"Error decrypting AES key: {str(e)}"

    def create_unified_auth_frame(self):
        """Create unified authentication frame with modern design"""
        # Create gradient background frame
        self.auth_container = tk.Frame(self.root, bg=self.colors['background'])
        self.auth_container.pack(fill=tk.BOTH, expand=True)
        
        # Create header with gradient effect
        header_frame = tk.Frame(self.auth_container, bg=self.colors['primary'], height=80)
        header_frame.pack(fill=tk.X, pady=(0, 30))
        header_frame.pack_propagate(False)
        
        # Title in header
        title_label = tk.Label(header_frame,
                              text="🔐 SecureVault",
                              font=('Segoe UI', 28, 'bold'),
                              fg=self.colors['text_light'],
                              bg=self.colors['primary'])
        title_label.pack(expand=True)
        
        # Subtitle
        subtitle_label = tk.Label(header_frame,
                                 text="Secure Encryption & File Sharing Platform",
                                 font=('Segoe UI', 12),
                                 fg=self.colors['highlight'],
                                 bg=self.colors['primary'])
        subtitle_label.pack()
        
        # Main content container
        content_frame = tk.Frame(self.auth_container, bg=self.colors['background'])
        content_frame.pack(expand=True, fill=tk.BOTH, padx=50, pady=20)
        
        # Create notebook with custom style
        self.auth_notebook = ttk.Notebook(content_frame, style='Custom.TNotebook')
        self.auth_notebook.pack(expand=True, fill=tk.BOTH, padx=20, pady=10)
        
        # Create tabs
        self.user_login_tab = self.create_tab_frame("👤 User Login")
        self.user_register_tab = self.create_tab_frame("📝 User Registration")
        self.admin_login_tab = self.create_tab_frame("🛡️ Admin Login")
        
        self.auth_notebook.add(self.user_login_tab, text="👤 User Login")
        self.auth_notebook.add(self.user_register_tab, text="📝 Register")
        self.auth_notebook.add(self.admin_login_tab, text="🛡️ Admin")
        
        self.create_user_login_tab()
        self.create_user_register_tab()
        self.create_admin_login_tab()
        
        # Footer
        footer_frame = tk.Frame(self.auth_container, bg=self.colors['primary'], height=40)
        footer_frame.pack(side=tk.BOTTOM, fill=tk.X)
        footer_frame.pack_propagate(False)
        
        footer_label = tk.Label(footer_frame,
                               text="© 2024 Secure File Sharing Tool | Version 3.0 (SQLite + JSON)",
                               font=('Segoe UI', 9),
                               fg=self.colors['text_light'],
                               bg=self.colors['primary'])
        footer_label.pack(expand=True)
    
    def create_tab_frame(self, title):
        """Create a styled tab frame"""
        frame = tk.Frame(self.auth_notebook, bg=self.colors['card_bg'])
        
        # Tab title
        title_label = tk.Label(frame,
                              text=title,
                              font=('Segoe UI', 16, 'bold'),
                              fg=self.colors['primary'],
                              bg=self.colors['card_bg'])
        title_label.pack(pady=(20, 10))
        
        # Separator
        separator = tk.Frame(frame, height=2, bg=self.colors['secondary'])
        separator.pack(fill=tk.X, padx=50, pady=(0, 20))
        
        return frame
    
    def create_user_login_tab(self):
        """Create user login tab with modern design"""
        # Main container with grid layout
        container = tk.Frame(self.user_login_tab, bg=self.colors['card_bg'])
        container.pack(expand=True, fill=tk.BOTH, padx=40, pady=20)
        
        # Welcome message
        welcome_label = tk.Label(container,
                                text="Welcome Back!",
                                font=('Segoe UI', 18, 'bold'),
                                fg=self.colors['secondary'],
                                bg=self.colors['card_bg'])
        welcome_label.grid(row=0, column=0, columnspan=2, pady=(0, 30))
        
        # Form fields
        form_frame = tk.Frame(container, bg=self.colors['card_bg'])
        form_frame.grid(row=1, column=0, columnspan=2, pady=10)
        
        # Email field
        email_frame = self.create_form_field(form_frame, "📧 Email:", 0)
        self.login_email_entry = ttk.Entry(email_frame, width=35, font=('Segoe UI', 10), style='Custom.TEntry')
        self.login_email_entry.pack(fill=tk.X, padx=10, pady=5)
        
        # Password field
        password_frame = self.create_form_field(form_frame, "🔒 Password:", 1)
        self.login_password_entry = ttk.Entry(password_frame, width=35, font=('Segoe UI', 10), show="*", style='Custom.TEntry')
        self.login_password_entry.pack(fill=tk.X, padx=10, pady=5)
        
        # Forgot password link
        forgot_frame = tk.Frame(form_frame, bg=self.colors['card_bg'])
        forgot_frame.grid(row=2, column=0, columnspan=2, pady=(10, 0), sticky=tk.E)
        
        forgot_btn = ttk.Button(forgot_frame,
                               text="🔓 Forgot Password?",
                               command=lambda: self.forgot_password_dialog(admin=False),
                               style='Warning.TButton',
                               width=20)
        forgot_btn.pack()
        
        # Login button
        button_frame = tk.Frame(container, bg=self.colors['card_bg'])
        button_frame.grid(row=3, column=0, columnspan=2, pady=30)
        
        login_btn = ttk.Button(button_frame,
                              text="🚀 Login as User",
                              command=self.login,
                              style='Success.TButton',
                              width=25)
        login_btn.pack()
        
        # Register link
        register_frame = tk.Frame(container, bg=self.colors['card_bg'])
        register_frame.grid(row=4, column=0, columnspan=2, pady=20)
        
        tk.Label(register_frame,
                text="Don't have an account?",
                font=('Segoe UI', 9),
                fg=self.colors['text_dark'],
                bg=self.colors['card_bg']).pack(side=tk.LEFT, padx=(0, 10))
        
        register_btn = ttk.Button(register_frame,
                                 text="📝 Create Account",
                                 command=lambda: self.auth_notebook.select(1),
                                 style='Primary.TButton',
                                 width=18)
        register_btn.pack(side=tk.LEFT)
    
    def create_form_field(self, parent, label_text, row):
        """Create a styled form field"""
        field_frame = tk.Frame(parent, bg=self.colors['card_bg'])
        field_frame.grid(row=row, column=0, columnspan=2, pady=10, sticky=tk.EW)
        field_frame.grid_columnconfigure(0, weight=1)
        field_frame.grid_columnconfigure(1, weight=2)
        
        label = tk.Label(field_frame,
                        text=label_text,
                        font=('Segoe UI', 10, 'bold'),
                        fg=self.colors['primary'],
                        bg=self.colors['card_bg'])
        label.grid(row=0, column=0, sticky=tk.W, padx=(0, 15))
        
        entry_frame = tk.Frame(field_frame, bg=self.colors['card_bg'])
        entry_frame.grid(row=0, column=1, sticky=tk.EW)
        
        return entry_frame
    
    def create_user_register_tab(self):
        """Create user registration tab with modern design"""
        # Create scrolled frame
        canvas = tk.Canvas(self.user_register_tab, bg=self.colors['card_bg'], highlightthickness=0)
        scrollbar = ttk.Scrollbar(self.user_register_tab, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg=self.colors['card_bg'])
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Registration form
        container = tk.Frame(scrollable_frame, bg=self.colors['card_bg'])
        container.pack(padx=40, pady=20)
        
        # Title
        title_label = tk.Label(container,
                              text="Create New Account",
                              font=('Segoe UI', 18, 'bold'),
                              fg=self.colors['secondary'],
                              bg=self.colors['card_bg'])
        title_label.grid(row=0, column=0, columnspan=2, pady=(0, 20))
        
        # Form fields
        fields = [
            ("📧 Email:", "self.reg_email_entry", False),
            ("🔒 Password:", "self.reg_password_entry", True),
            ("✅ Confirm Password:", "self.reg_confirm_password_entry", True),
            ("❓ Security Question:", "self.reg_security_question", False),
            ("💡 Security Answer:", "self.reg_security_answer", False)
        ]
        
        row = 1
        for label_text, attr_name, is_password in fields:
            field_frame = self.create_form_field(container, label_text, row)
            
            if attr_name == "self.reg_security_question":
                combo = ttk.Combobox(field_frame, 
                                   font=('Segoe UI', 10),
                                   width=33,
                                   state='readonly')
                combo['values'] = [
                    "What is your favorite color?",
                    "What is your pet's name?",
                    "What city were you born in?",
                    "What is your mother's maiden name?",
                    "What was your first school?"
                ]
                combo.current(0)
                setattr(self, attr_name[5:], combo)
                combo.pack(fill=tk.X, padx=10, pady=5)
            else:
                entry = ttk.Entry(field_frame,
                                width=35,
                                font=('Segoe UI', 10),
                                show="*" if is_password else "",
                                style='Custom.TEntry')
                entry.pack(fill=tk.X, padx=10, pady=5)
                setattr(self, attr_name[5:], entry)
            
            row += 1
        
        # Register button
        button_frame = tk.Frame(container, bg=self.colors['card_bg'])
        button_frame.grid(row=row, column=0, columnspan=2, pady=30)
        
        register_btn = ttk.Button(button_frame,
                                 text="📋 Register Account",
                                 command=self.register,
                                 style='Success.TButton',
                                 width=25)
        register_btn.pack()
        
        # Back to login
        back_frame = tk.Frame(container, bg=self.colors['card_bg'])
        back_frame.grid(row=row+1, column=0, columnspan=2, pady=20)
        
        tk.Label(back_frame,
                text="Already have an account?",
                font=('Segoe UI', 9),
                fg=self.colors['text_dark'],
                bg=self.colors['card_bg']).pack(side=tk.LEFT, padx=(0, 10))
        
        back_btn = ttk.Button(back_frame,
                             text="🔙 Back to Login",
                             command=lambda: self.auth_notebook.select(0),
                             style='Primary.TButton',
                             width=18)
        back_btn.pack(side=tk.LEFT)
    
    def create_admin_login_tab(self):
        """Create admin login tab with modern design"""
        container = tk.Frame(self.admin_login_tab, bg=self.colors['card_bg'])
        container.pack(expand=True, fill=tk.BOTH, padx=40, pady=40)
        
        # Admin icon and title
        icon_frame = tk.Frame(container, bg=self.colors['card_bg'])
        icon_frame.pack(pady=(0, 20))
        
        tk.Label(icon_frame,
                text="🛡️",
                font=('Segoe UI', 48),
                fg=self.colors['primary'],
                bg=self.colors['card_bg']).pack()
        
        title_label = tk.Label(container,
                              text="Administrator Access",
                              font=('Segoe UI', 20, 'bold'),
                              fg=self.colors['primary'],
                              bg=self.colors['card_bg'])
        title_label.pack(pady=(0, 30))
        
        # Form
        form_frame = tk.Frame(container, bg=self.colors['card_bg'])
        form_frame.pack(pady=10)
        
        # Email field
        email_frame = self.create_form_field(form_frame, "📧 Admin Email:", 0)
        self.admin_email_entry = ttk.Entry(email_frame, width=35, font=('Segoe UI', 10), style='Custom.TEntry')
        self.admin_email_entry.pack(fill=tk.X, padx=10, pady=5)
        
        # Password field
        password_frame = self.create_form_field(form_frame, "🔒 Password:", 1)
        self.admin_password_entry = ttk.Entry(password_frame, width=35, font=('Segoe UI', 10), show="*", style='Custom.TEntry')
        self.admin_password_entry.pack(fill=tk.X, padx=10, pady=5)
        
        # Forgot password
        forgot_frame = tk.Frame(form_frame, bg=self.colors['card_bg'])
        forgot_frame.grid(row=2, column=0, columnspan=2, pady=(20, 0), sticky=tk.E)
        
        forgot_btn = ttk.Button(forgot_frame,
                               text="🔓 Reset Admin Password",
                               command=lambda: self.forgot_password_dialog(admin=True),
                               style='Warning.TButton',
                               width=25)
        forgot_btn.pack()
        
        # Login button
        button_frame = tk.Frame(container, bg=self.colors['card_bg'])
        button_frame.pack(pady=40)
        
        login_btn = ttk.Button(button_frame,
                              text="⚡ Login as Admin",
                              command=self.admin_login,
                              style='Danger.TButton',
                              width=25)
        login_btn.pack()
        
        # Warning label
        warning_label = tk.Label(container,
                                text="⚠️ Restricted access - Admin privileges required",
                                font=('Segoe UI', 9, 'italic'),
                                fg=self.colors['accent'],
                                bg=self.colors['card_bg'])
        warning_label.pack(pady=20)
    
    def forgot_password_dialog(self, admin=False):
        """Create modern password reset dialog"""
        dialog = tk.Toplevel(self.root)
        dialog.title("🔓 Password Reset")
        dialog.geometry("500x600")
        dialog.resizable(False, False)
        dialog.configure(bg=self.colors['background'])
        dialog.transient(self.root)
        dialog.grab_set()
        
        self.center_window(dialog, 500, 600)
        
        # Header
        header_frame = tk.Frame(dialog, bg=self.colors['primary'], height=60)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        tk.Label(header_frame,
                text="🔓 Password Reset",
                font=('Segoe UI', 18, 'bold'),
                fg=self.colors['text_light'],
                bg=self.colors['primary']).pack(expand=True)
        
        if admin:
            tk.Label(header_frame,
                    text="Admin Password Recovery",
                    font=('Segoe UI', 10),
                    fg=self.colors['highlight'],
                    bg=self.colors['primary']).pack()
        
        # Main content
        main_frame = tk.Frame(dialog, bg=self.colors['card_bg'], padx=20, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Form
        fields = [
            ("📧 Email Address:", "email"),
            ("❓ Security Question:", "question"),
            ("💡 Security Answer:", "answer"),
            ("🔑 New Password:", "new_pass"),
            ("✅ Confirm Password:", "confirm_pass")
        ]
        
        entries = {}
        
        for i, (label, key) in enumerate(fields):
            tk.Label(main_frame,
                    text=label,
                    font=('Segoe UI', 9, 'bold'),
                    fg=self.colors['primary'],
                    bg=self.colors['card_bg'],
                    anchor='w').grid(row=i*2, column=0, sticky=tk.W, pady=(10 if i == 0 else 15))
            
            if key == "question":
                combo = ttk.Combobox(main_frame,
                                   font=('Segoe UI', 10),
                                   width=40,
                                   state='readonly')
                combo['values'] = [
                    "What is your favorite color?",
                    "What is your pet's name?",
                    "What city were you born in?",
                    "What is your mother's maiden name?",
                    "What was your first school?"
                ]
                combo.current(0)
                combo.grid(row=i*2+1, column=0, pady=5, sticky=tk.W)
                entries[key] = combo
            else:
                show_char = "*" if "pass" in key else ""
                entry = ttk.Entry(main_frame,
                                width=43,
                                font=('Segoe UI', 10),
                                show=show_char,
                                style='Custom.TEntry')
                entry.grid(row=i*2+1, column=0, pady=5, sticky=tk.W)
                entries[key] = entry
        
        # Regenerate RSA keys option
        rsa_var = tk.BooleanVar(value=True)
        rsa_check = tk.Checkbutton(main_frame,
                                  text="🔄 Regenerate RSA keys with new password",
                                  variable=rsa_var,
                                  font=('Segoe UI', 9),
                                  fg=self.colors['primary'],
                                  bg=self.colors['card_bg'],
                                  anchor='w')
        rsa_check.grid(row=10, column=0, sticky=tk.W, pady=10)
        
        # Buttons
        button_frame = tk.Frame(dialog, bg=self.colors['background'], pady=20)
        button_frame.pack(fill=tk.X, padx=20)
        
        def reset_password():
            email = entries['email'].get().strip().lower()
            new_password = entries['new_pass'].get()
            confirm_password = entries['confirm_pass'].get()
            answer = entries['answer'].get().strip()
            regenerate_rsa = rsa_var.get()
            
            if not email or not new_password or not confirm_password or not answer:
                messagebox.showerror("Error", "All fields are required", parent=dialog)
                return
            
            if not self.validate_email(email):
                messagebox.showerror("Error", "Please enter a valid email address", parent=dialog)
                return
            
            if new_password != confirm_password:
                messagebox.showerror("Error", "Passwords do not match", parent=dialog)
                return
            
            is_valid, msg = self.validate_password_strength(new_password)
            if not is_valid:
                messagebox.showerror("Error", msg, parent=dialog)
                return
            
            try:
                user_data = self.db.get_user(email)
                
                if not user_data:
                    messagebox.showerror("Error", "Email not found in system", parent=dialog)
                    return
                
                if admin and user_data["role"] != "admin":
                    messagebox.showerror("Error", "This email is not registered as admin", parent=dialog)
                    return
                
                stored_answer_hash = user_data.get("security_answer_hash", "")
                if not stored_answer_hash:
                    default_answer_hash = self.hash_password("blue")
                    if self.hash_password(answer) != default_answer_hash:
                        messagebox.showerror("Error", "Security answer is incorrect", parent=dialog)
                        return
                elif self.hash_password(answer) != stored_answer_hash:
                    messagebox.showerror("Error", "Security answer is incorrect", parent=dialog)
                    return
                
                # Update password in database
                self.db.update_user_password(email, self.hash_password(new_password))
                
                # Update password in JSON
                json_user = self.json_manager.get_user(email)
                if json_user:
                    json_user["password_hash"] = self.hash_password(new_password)
                    self.json_manager.update_user(email, json_user)
                
                # Regenerate RSA keys if requested
                if regenerate_rsa:
                    if self.generate_user_rsa_keys(email, new_password):
                        messagebox.showinfo("Success", 
                                          "✅ Password reset successfully!\n\n"
                                          "Your RSA keys have been regenerated with the new password.",
                                          parent=dialog)
                    else:
                        messagebox.showwarning("Warning",
                                             "Password reset but RSA key regeneration failed. "
                                             "Please login and regenerate keys manually.",
                                             parent=dialog)
                else:
                    messagebox.showinfo("Success", 
                                      "✅ Password reset successfully!\n\n"
                                      "Note: Your old RSA keys are still encrypted with the old password. "
                                      "You will need to regenerate them to use file sharing.",
                                      parent=dialog)
                
                self.log_action("PASSWORD_RESET", email)
                dialog.destroy()
                
                if admin:
                    self.admin_email_entry.delete(0, tk.END)
                    self.admin_password_entry.delete(0, tk.END)
                else:
                    self.login_email_entry.delete(0, tk.END)
                    self.login_password_entry.delete(0, tk.END)
                
            except Exception as e:
                messagebox.showerror("Error", f"Password reset failed: {str(e)}", parent=dialog)
        
        ttk.Button(button_frame,
                  text="🔓 Reset Password",
                  command=reset_password,
                  style='Success.TButton',
                  width=20).pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Button(button_frame,
                  text="❌ Cancel",
                  command=dialog.destroy,
                  style='Danger.TButton',
                  width=20).pack(side=tk.RIGHT)
        
        dialog.bind('<Return>', lambda e: reset_password())
        entries['email'].focus_set()
    
    def admin_login(self):
        """Handle admin login"""
        email = self.admin_email_entry.get().strip().lower()
        password = self.admin_password_entry.get()
        
        if not email or not password:
            messagebox.showerror("Error", "Please enter both email and password")
            return
        
        if self.is_account_locked(email):
            failed_attempts = self.db.get_failed_login_attempts(email)
            last_attempt = failed_attempts['last_attempt']
            if last_attempt:
                try:
                    last_attempt_time = datetime.strptime(last_attempt, '%Y-%m-%d %H:%M:%S')
                    remaining_time = int(self.lockout_time - (datetime.now() - last_attempt_time).total_seconds())
                    messagebox.showerror("Error", f"Account locked. Try again in {remaining_time} seconds.")
                except:
                    messagebox.showerror("Error", "Account locked. Try again later.")
            return
        
        try:
            user_data = self.db.get_user(email)
            
            if user_data and user_data['password_hash'] == self.hash_password(password):
                if user_data['role'] == "admin":
                    # Reset failed attempts
                    self.db.reset_failed_login_attempts(email)
                    
                    self.current_user = email
                    self.user_role = "admin"
                    
                    # Load RSA keys for admin
                    success, message = self.load_user_rsa_keys(email, password)
                    if not success:
                        # Generate RSA keys if they don't exist
                        if "not found" in message.lower():
                            if self.generate_user_rsa_keys(email, password):
                                success, message = self.load_user_rsa_keys(email, password)
                    
                    if not success:
                        messagebox.showwarning("Warning", f"RSA keys not loaded: {message}")
                    
                    # Update last login in JSON
                    json_user = self.json_manager.get_user(email)
                    if json_user:
                        json_user["last_login"] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                        self.json_manager.update_user(email, json_user)
                    
                    self.update_login_stats()
                    self.add_active_session()
                    
                    self.auth_container.destroy()
                    self.create_main_interface()
                    
                    self.log_action("ADMIN_LOGIN", "")
                else:
                    messagebox.showerror("Error", "This user is not an admin")
                    self.record_failed_attempt(email)
            else:
                attempts = self.record_failed_attempt(email)
                if attempts >= self.max_login_attempts:
                    messagebox.showerror("Error", "Account locked for 5 minutes due to too many failed attempts")
                else:
                    messagebox.showerror("Error", f"Invalid credentials. {self.max_login_attempts - attempts} attempts remaining.")
                
        except Exception as e:
            messagebox.showerror("Error", f"Admin login failed: {str(e)}")
    
    def register(self):
        """Handle user registration"""
        email = self.reg_email_entry.get().strip().lower()
        password = self.reg_password_entry.get()
        confirm_password = self.reg_confirm_password_entry.get()
        security_question = self.reg_security_question.get()
        security_answer = self.reg_security_answer.get().strip()
        
        if not email or not password or not confirm_password or not security_answer:
            messagebox.showerror("Error", "All fields are required")
            return
        
        if not self.validate_email(email):
            messagebox.showerror("Error", "Please enter a valid email address")
            return
        
        if password != confirm_password:
            messagebox.showerror("Error", "Passwords do not match")
            return
        
        is_valid, msg = self.validate_password_strength(password)
        if not is_valid:
            messagebox.showerror("Error", msg)
            return
        
        try:
            # Check if user already exists
            existing_user = self.db.get_user(email)
            if existing_user:
                messagebox.showerror("Error", "User already exists with this email")
                return
            
            # Add user to database
            success = self.db.add_user(
                email=email,
                password_hash=self.hash_password(password),
                role="user",
                security_question=security_question,
                security_answer_hash=self.hash_password(security_answer)
            )
            
            if success:
                # Also add to JSON file
                user_data = {
                    "email": email,
                    "password_hash": self.hash_password(password),
                    "role": "user",
                    "security_question": security_question,
                    "security_answer_hash": self.hash_password(security_answer),
                    "created_at": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                    "last_login": None,
                    "is_active": 1
                }
                self.json_manager.add_user(user_data)
                
                # Generate RSA keys for the new user
                if self.generate_user_rsa_keys(email, password):
                    messagebox.showinfo("Success", "✅ Registration successful! RSA keys generated and uploaded to repository.")
                else:
                    messagebox.showwarning("Warning", "Registration successful but RSA key generation failed. Please login and regenerate keys.")
                
                self.reg_email_entry.delete(0, tk.END)
                self.reg_password_entry.delete(0, tk.END)
                self.reg_confirm_password_entry.delete(0, tk.END)
                self.reg_security_answer.delete(0, tk.END)
                
                self.auth_notebook.select(0)
            else:
                messagebox.showerror("Error", "Registration failed: Could not create user")
            
        except Exception as e:
            messagebox.showerror("Error", f"Registration failed: {str(e)}")
    
    def login(self):
        """Handle user login"""
        email = self.login_email_entry.get().strip().lower()
        password = self.login_password_entry.get()
        
        if not email or not password:
            messagebox.showerror("Error", "Please enter both email and password")
            return
        
        if self.is_account_locked(email):
            failed_attempts = self.db.get_failed_login_attempts(email)
            last_attempt = failed_attempts['last_attempt']
            if last_attempt:
                try:
                    last_attempt_time = datetime.strptime(last_attempt, '%Y-%m-%d %H:%M:%S')
                    remaining_time = int(self.lockout_time - (datetime.now() - last_attempt_time).total_seconds())
                    messagebox.showerror("Error", f"Account locked. Try again in {remaining_time} seconds.")
                except:
                    messagebox.showerror("Error", "Account locked. Try again later.")
            return
        
        try:
            user_data = self.db.get_user(email)
            
            if user_data and user_data['password_hash'] == self.hash_password(password):
                # Reset failed attempts
                self.db.reset_failed_login_attempts(email)
                
                self.current_user = email
                self.user_role = user_data['role']
                
                # Load RSA keys
                success, message = self.load_user_rsa_keys(email, password)
                if not success:
                    # Generate RSA keys if they don't exist
                    if "not found" in message.lower():
                        if self.generate_user_rsa_keys(email, password):
                            success, message = self.load_user_rsa_keys(email, password)
                
                if not success:
                    messagebox.showwarning("Warning", f"RSA keys not loaded: {message}")
                
                # Update last login in JSON
                json_user = self.json_manager.get_user(email)
                if json_user:
                    json_user["last_login"] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    self.json_manager.update_user(email, json_user)
                
                # Update public key in database (in case it changed)
                if self.user_public_key:
                    public_key_pem = self.user_public_key.public_bytes(
                        encoding=serialization.Encoding.PEM,
                        format=serialization.PublicFormat.SubjectPublicKeyInfo
                    )
                    self.update_public_key(email, public_key_pem)
                
                self.update_login_stats()
                self.add_active_session()
                
                self.auth_container.destroy()
                self.create_main_interface()
            else:
                attempts = self.record_failed_attempt(email)
                if attempts >= self.max_login_attempts:
                    messagebox.showerror("Error", "Account locked for 5 minutes due to too many failed attempts")
                else:
                    messagebox.showerror("Error", f"Invalid credentials. {self.max_login_attempts - attempts} attempts remaining.")
                
        except Exception as e:
            messagebox.showerror("Error", f"Login failed: {str(e)}")

    def create_main_interface(self):
        """Create main application interface with modern design"""
        # Clear any existing widgets
        for widget in self.root.winfo_children():
            widget.destroy()
        
        # Create sidebar
        self.create_sidebar()
        
        # Create main content area
        self.main_content = tk.Frame(self.root, bg=self.colors['background'])
        self.main_content.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Create header
        self.create_header()
        
        # Create notebook for tabs
        self.notebook = ttk.Notebook(self.main_content, style='Custom.TNotebook')
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=20, pady=(10, 20))
        
        # Create tabs
        self.create_file_operations_tab()
        self.create_file_sharing_tab()
        self.create_key_management_tab()
        self.create_metadata_extractor_tab()  # NEW TAB
        
        if self.user_role == "admin":
            self.create_admin_dashboard_tab()
    
    def create_sidebar(self):
        """Create a colorful sidebar"""
        sidebar = tk.Frame(self.root, bg=self.colors['primary'], width=220)
        sidebar.pack(side=tk.LEFT, fill=tk.Y)
        sidebar.pack_propagate(False)
        
        # User info section
        user_frame = tk.Frame(sidebar, bg=self.colors['primary'])
        user_frame.pack(fill=tk.X, padx=20, pady=30)
        
        # User avatar
        avatar_label = tk.Label(user_frame,
                               text="👤",
                               font=('Segoe UI', 48),
                               fg=self.colors['highlight'],
                               bg=self.colors['primary'])
        avatar_label.pack(pady=(0, 10))
        
        # User name
        username_label = tk.Label(user_frame,
                                 text=self.current_user.split('@')[0],
                                 font=('Segoe UI', 14, 'bold'),
                                 fg=self.colors['text_light'],
                                 bg=self.colors['primary'])
        username_label.pack()
        
        # User role
        role_label = tk.Label(user_frame,
                             text="Administrator" if self.user_role == "admin" else "User",
                             font=('Segoe UI', 11),
                             fg=self.colors['highlight'],
                             bg=self.colors['primary'])
        role_label.pack(pady=(5, 0))
        
        # RSA Key Status
        key_status = "🔑 RSA: Loaded" if self.user_private_key else "🔐 RSA: Not Loaded"
        key_label = tk.Label(user_frame,
                           text=key_status,
                           font=('Segoe UI', 9),
                           fg=self.colors['highlight'] if self.user_private_key else self.colors['warning'],
                           bg=self.colors['primary'])
        key_label.pack(pady=(5, 0))
        
        # Public Key Status
        public_key_info = self.get_public_key_from_repository(self.current_user)
        pub_status = "🌐 Public Key: Available" if public_key_info[0] else "🌐 Public Key: Not Found"
        pub_label = tk.Label(user_frame,
                           text=pub_status,
                           font=('Segoe UI', 9),
                           fg=self.colors['highlight'] if public_key_info[0] else self.colors['warning'],
                           bg=self.colors['primary'])
        pub_label.pack(pady=(2, 0))
        
        # JSON Status
        json_status = "📄 JSON: Synced" if os.path.exists("users.json") else "📄 JSON: Not Found"
        json_label = tk.Label(user_frame,
                           text=json_status,
                           font=('Segoe UI', 9),
                           fg=self.colors['highlight'] if os.path.exists("users.json") else self.colors['warning'],
                           bg=self.colors['primary'])
        json_label.pack(pady=(2, 0))
        
        # Separator
        separator = tk.Frame(sidebar, height=2, bg=self.colors['secondary'])
        separator.pack(fill=tk.X, padx=20, pady=20)
        
        # Navigation buttons
        nav_buttons = [
            ("📁 File Operations", self.show_file_operations),
            ("🤝 File Sharing", self.show_file_sharing),
            ("🔑 Key Management", self.show_key_management),
            ("🔍 Metadata Extractor", self.show_metadata_extractor)  # NEW
        ]
        
        if self.user_role == "admin":
            nav_buttons.append(("📊 Admin Dashboard", self.show_admin_dashboard))
        
        for i, (text, command) in enumerate(nav_buttons):
            btn = tk.Button(sidebar,
                          text=text,
                          font=('Segoe UI', 11),
                          fg=self.colors['text_light'],
                          bg=self.colors['primary'],
                          activeforeground=self.colors['text_light'],
                          activebackground=self.colors['secondary'],
                          bd=0,
                          padx=20,
                          pady=12,
                          anchor='w',
                          cursor='hand2',
                          command=command)
            btn.pack(fill=tk.X, padx=10)
            
            # Add hover effect
            btn.bind("<Enter>", lambda e, b=btn: b.config(bg=self.colors['secondary']))
            btn.bind("<Leave>", lambda e, b=btn: b.config(bg=self.colors['primary']))
        
        # Separator
        separator2 = tk.Frame(sidebar, height=2, bg=self.colors['secondary'])
        separator2.pack(fill=tk.X, padx=20, pady=20)
        
        # Logout button
        logout_btn = tk.Button(sidebar,
                             text="🚪 Logout",
                             font=('Segoe UI', 11, 'bold'),
                             fg=self.colors['text_light'],
                             bg=self.colors['accent'],
                             activeforeground=self.colors['text_light'],
                             activebackground='#FF4D4D',
                             bd=0,
                             padx=20,
                             pady=12,
                             cursor='hand2',
                             command=self.logout)
        logout_btn.pack(side=tk.BOTTOM, fill=tk.X, padx=10, pady=(0, 20))
        
        logout_btn.bind("<Enter>", lambda e: logout_btn.config(bg='#FF4D4D'))
        logout_btn.bind("<Leave>", lambda e: logout_btn.config(bg=self.colors['accent']))
    
    def create_header(self):
        """Create header with welcome message"""
        header = tk.Frame(self.main_content, bg=self.colors['primary'], height=80)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        
        # Welcome message
        welcome_text = f"Welcome back, {self.current_user.split('@')[0]}! 👋"
        welcome_label = tk.Label(header,
                                text=welcome_text,
                                font=('Segoe UI', 20, 'bold'),
                                fg=self.colors['text_light'],
                                bg=self.colors['primary'])
        welcome_label.pack(side=tk.LEFT, padx=30, pady=20)
        
        # Current time
        self.time_label = tk.Label(header,
                                 text=datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                                 font=('Segoe UI', 11),
                                 fg=self.colors['highlight'],
                                 bg=self.colors['primary'])
        self.time_label.pack(side=tk.RIGHT, padx=30, pady=20)
        
        # Update time every second
        def update_time():
            try:
                # Check if the widget still exists before updating
                if self.time_label and self.time_label.winfo_exists():
                    self.time_label.config(text=datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
                    self.root.after(1000, update_time)
            except (tk.TclError, AttributeError):
                # Widget has been destroyed, stop the timer
                pass
        
        # Store the timer ID so we can cancel it if needed
        self.time_updater_id = self.root.after(1000, update_time)
    
    def show_file_operations(self):
        """Show file operations tab"""
        self.notebook.select(0)
    
    def show_file_sharing(self):
        """Show file sharing tab"""
        self.notebook.select(1)
    
    def show_key_management(self):
        """Show key management tab"""
        self.notebook.select(2)
    
    def show_metadata_extractor(self):
        """Show metadata extractor tab"""
        self.notebook.select(3)
    
    def show_admin_dashboard(self):
        """Show admin dashboard tab"""
        if self.user_role == "admin":
            self.notebook.select(4)
    
    def create_metadata_extractor_tab(self):
        """Create metadata extractor tab with modern design"""
        self.metadata_frame = tk.Frame(self.notebook, bg=self.colors['background'])
        self.notebook.add(self.metadata_frame, text="🔍 Metadata Extractor")
        
        # Create notebook for metadata sections
        metadata_notebook = ttk.Notebook(self.metadata_frame, style='Custom.TNotebook')
        metadata_notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Metadata Extraction Tab
        self.create_metadata_extraction_tab(metadata_notebook)
        
        # Supported Formats Tab
        self.create_supported_formats_tab(metadata_notebook)
    
    def create_metadata_extraction_tab(self, notebook):
        """Create metadata extraction tab with fixed scrolling"""
        extract_tab = tk.Frame(notebook, bg=self.colors['background'])
        notebook.add(extract_tab, text="📊 Extract Metadata")
        
        # Create main container with vertical layout
        main_container = tk.Frame(extract_tab, bg=self.colors['background'])
        main_container.pack(fill=tk.BOTH, expand=True)
        
        # Create canvas and scrollbar
        canvas = tk.Canvas(main_container, bg=self.colors['background'], highlightthickness=0)
        scrollbar = ttk.Scrollbar(main_container, orient="vertical", command=canvas.yview)
        
        # Create scrollable frame
        scrollable_frame = tk.Frame(canvas, bg=self.colors['background'])
        
        # Configure the scrollregion
        def configure_scrollregion(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
        
        # Create window in canvas
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw", tags="scrollable_frame")
        
        # Configure canvas
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Pack canvas and scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Bind configuration events
        scrollable_frame.bind("<Configure>", configure_scrollregion)
        
        # Adjust canvas window width
        def configure_canvas(event):
            canvas.itemconfig("scrollable_frame", width=event.width)
        
        canvas.bind("<Configure>", configure_canvas)
        
        # Now add all content to scrollable_frame
        container = tk.Frame(scrollable_frame, bg=self.colors['background'])
        container.pack(padx=20, pady=20, fill=tk.BOTH, expand=True)
        
        # Title
        title_label = tk.Label(container,
                              text="📊 File Metadata Extractor",
                              font=('Segoe UI', 20, 'bold'),
                              fg=self.colors['primary'],
                              bg=self.colors['background'])
        title_label.pack(pady=(0, 20))
        
        # Description
        desc_label = tk.Label(container,
                             text="Extract detailed metadata from files including images, documents, PDFs, archives, and more.",
                             font=('Segoe UI', 11),
                             fg=self.colors['text_dark'],
                             bg=self.colors['background'],
                             wraplength=800,
                             justify=tk.LEFT)
        desc_label.pack(pady=(0, 30))
        
        # File selection section
        file_section = tk.LabelFrame(container,
                                    text=" 📁 Select File",
                                    font=('Segoe UI', 12, 'bold'),
                                    fg=self.colors['primary'],
                                    bg=self.colors['card_bg'],
                                    relief='groove',
                                    borderwidth=2)
        file_section.pack(fill=tk.X, padx=10, pady=10)
        
        # File path entry
        file_path_frame = tk.Frame(file_section, bg=self.colors['card_bg'])
        file_path_frame.pack(fill=tk.X, padx=20, pady=20)
        
        tk.Label(file_path_frame,
                text="File path:",
                font=('Segoe UI', 10, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['card_bg']).pack(side=tk.LEFT, padx=(0, 10))
        
        self.metadata_file_path = tk.StringVar()
        path_entry = ttk.Entry(file_path_frame,
                          textvariable=self.metadata_file_path,
                          font=('Segoe UI', 10),
                          width=50,
                          style='Custom.TEntry')
        path_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        
        ttk.Button(file_path_frame,
                  text="📂 Browse",
                  command=self.browse_metadata_file,
                  style='Primary.TButton',
                  width=10).pack(side=tk.LEFT)
        
        # Action buttons
        action_frame = tk.Frame(file_section, bg=self.colors['card_bg'])
        action_frame.pack(fill=tk.X, padx=20, pady=(0, 20))
        
        ttk.Button(action_frame,
                  text="🔍 Extract Metadata",
                  command=self.extract_metadata,
                  style='Success.TButton',
                  width=20).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(action_frame,
                  text="💾 Export to JSON",
                  command=self.export_metadata,
                  style='Primary.TButton',
                  width=20).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(action_frame,
                  text="🧹 Clear",
                  command=self.clear_metadata_display,
                  style='Danger.TButton',
                  width=20).pack(side=tk.LEFT, padx=5)
        
        # Results display section
        results_section = tk.LabelFrame(container,
                                       text=" 📋 Extracted Metadata",
                                       font=('Segoe UI', 12, 'bold'),
                                       fg=self.colors['primary'],
                                       bg=self.colors['card_bg'],
                                       relief='groove',
                                       borderwidth=2)
        results_section.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Create notebook for different metadata views
        results_notebook = ttk.Notebook(results_section, style='Custom.TNotebook')
        results_notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Tree view tab
        tree_tab = tk.Frame(results_notebook, bg=self.colors['card_bg'])
        results_notebook.add(tree_tab, text="🌳 Tree View")
        
        # Create tree with scrollbars
        tree_frame = tk.Frame(tree_tab, bg=self.colors['card_bg'])
        tree_frame.pack(fill=tk.BOTH, expand=True)
        
        # Horizontal scrollbar
        h_scrollbar = ttk.Scrollbar(tree_frame, orient="horizontal")
        h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
        
        # Vertical scrollbar
        v_scrollbar = ttk.Scrollbar(tree_frame, orient="vertical")
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Treeview
        self.metadata_tree = ttk.Treeview(tree_frame,
                                         columns=("Property", "Value"),
                                         show="tree headings",
                                         height=20,
                                         style='Custom.Treeview',
                                         xscrollcommand=h_scrollbar.set,
                                         yscrollcommand=v_scrollbar.set)
        
        self.metadata_tree.heading("#0", text="Category")
        self.metadata_tree.heading("Property", text="Property")
        self.metadata_tree.heading("Value", text="Value")
        
        self.metadata_tree.column("#0", width=200)
        self.metadata_tree.column("Property", width=200)
        self.metadata_tree.column("Value", width=400)
        
        self.metadata_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        h_scrollbar.config(command=self.metadata_tree.xview)
        v_scrollbar.config(command=self.metadata_tree.yview)
        
        # JSON view tab
        json_tab = tk.Frame(results_notebook, bg=self.colors['card_bg'])
        results_notebook.add(json_tab, text="📝 JSON View")
        
        self.metadata_json_text = scrolledtext.ScrolledText(json_tab,
                                                           height=25,
                                                           font=('Consolas', 9),
                                                           bg='#1E1E1E',
                                                           fg='#00FF00',
                                                           relief='flat',
                                                           borderwidth=0)
        self.metadata_json_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        self.metadata_json_text.config(state=tk.DISABLED)
        
        # Basic view tab
        basic_tab = tk.Frame(results_notebook, bg=self.colors['card_bg'])
        results_notebook.add(basic_tab, text="📄 Basic Info")
        
        self.metadata_basic_text = scrolledtext.ScrolledText(basic_tab,
                                                            height=25,
                                                            font=('Consolas', 10),
                                                            bg=self.colors['card_bg'],
                                                            fg=self.colors['text_dark'],
                                                            relief='flat',
                                                            borderwidth=0)
        self.metadata_basic_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        self.metadata_basic_text.config(state=tk.DISABLED)
        
        # Status bar
        status_frame = tk.Frame(container, bg=self.colors['primary'], height=30)
        status_frame.pack(fill=tk.X, pady=(10, 0))
        status_frame.pack_propagate(False)
        
        self.metadata_status_label = tk.Label(status_frame,
                                             text="Ready to extract metadata",
                                             font=('Segoe UI', 9),
                                             fg=self.colors['text_light'],
                                             bg=self.colors['primary'])
        self.metadata_status_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        # Bind mouse wheel scrolling
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        # Bind to canvas
        canvas.bind_all("<MouseWheel>", on_mousewheel)
        
        # For Linux
        canvas.bind_all("<Button-4>", lambda e: canvas.yview_scroll(-1, "units"))
        canvas.bind_all("<Button-5>", lambda e: canvas.yview_scroll(1, "units"))
    
    def create_supported_formats_tab(self, notebook):
        """Create supported formats tab"""
        formats_tab = tk.Frame(notebook, bg=self.colors['background'])
        notebook.add(formats_tab, text="📋 Supported Formats")
        
        container = tk.Frame(formats_tab, bg=self.colors['background'])
        container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        # Title
        tk.Label(container,
                text="📋 Supported File Formats",
                font=('Segoe UI', 20, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['background']).pack(pady=(0, 20))
        
        # Card
        card = tk.Frame(container, bg=self.colors['card_bg'], relief='ridge', borderwidth=1)
        card.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Supported formats by category
        categories = [
            ("🖼️ Images", self.metadata_extractor.supported_formats['images']),
            ("📄 Documents", self.metadata_extractor.supported_formats['documents']),
            ("📦 Archives", self.metadata_extractor.supported_formats['archives']),
            ("🎵 Media Files", self.metadata_extractor.supported_formats['media'])
        ]
        
        row = 0
        for category_name, formats in categories:
            # Category label
            cat_frame = tk.Frame(card, bg=self.colors['card_bg'])
            cat_frame.grid(row=row, column=0, sticky=tk.W, padx=30, pady=(20 if row == 0 else 10))
            
            tk.Label(cat_frame,
                    text=category_name,
                    font=('Segoe UI', 12, 'bold'),
                    fg=self.colors['secondary'],
                    bg=self.colors['card_bg']).pack(anchor=tk.W)
            
            # Formats
            formats_frame = tk.Frame(card, bg=self.colors['card_bg'])
            formats_frame.grid(row=row + 1, column=0, sticky=tk.W, padx=50, pady=(0, 10))
            
            # Display formats in a grid
            for i, fmt in enumerate(formats):
                col = i % 4
                format_label = tk.Label(formats_frame,
                                       text=fmt,
                                       font=('Consolas', 10),
                                       fg=self.colors['text_dark'],
                                       bg=self.colors['card_bg'],
                                       padx=10,
                                       pady=5,
                                       relief='ridge',
                                       borderwidth=1)
                format_label.grid(row=i // 4, column=col, padx=5, pady=2, sticky=tk.W)
            
            row += 2
        
        # Special formats note
        note_frame = tk.Frame(card, bg=self.colors['card_bg'])
        note_frame.grid(row=row, column=0, sticky=tk.W, padx=30, pady=20)
        
        tk.Label(note_frame,
                text="💡 Note: Also supports encrypted files (.encrypted) for basic crypto metadata extraction",
                font=('Segoe UI', 6, 'italic'),
                fg=self.colors['warning'],
                bg=self.colors['card_bg']).pack(anchor=tk.W)
        
        # Capabilities section
        capabilities_frame = tk.Frame(card, bg=self.colors['card_bg'])
        capabilities_frame.grid(row=row + 1, column=0, sticky=tk.W, padx=30, pady=(0, 20))
        
        tk.Label(capabilities_frame,
                text="🔧 Extraction Capabilities:",
                font=('Segoe UI', 11, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['card_bg']).pack(anchor=tk.W, pady=(0, 10))
        
        capabilities = [
            "✓ Basic file metadata (size, dates, permissions)",
            "✓ Image EXIF data (camera settings, GPS coordinates)",
            "✓ PDF document info and structure",
            "✓ Document properties (DOCX, TXT)",
            "✓ CSV file structure and sample data",
            "✓ Archive contents (ZIP, TAR)",
            "✓ Cryptographic file analysis",
            "✓ JSON export capabilities"
        ]
        
        for capability in capabilities:
            tk.Label(capabilities_frame,
                    text=capability,
                    font=('Segoe UI', 10),
                    fg=self.colors['text_dark'],
                    bg=self.colors['card_bg']).pack(anchor=tk.W)
    
    def browse_metadata_file(self):
        """Browse for file to extract metadata"""
        filename = filedialog.askopenfilename(
            title="Select file for metadata extraction",
            filetypes=[
                ("All files", "*.*"),
                ("Image files", "*.jpg *.jpeg *.png *.gif *.bmp *.tiff"),
                ("PDF files", "*.pdf"),
                ("Documents", "*.docx *.txt *.csv"),
                ("Archives", "*.zip *.tar *.gz"),
                ("Encrypted files", "*.encrypted")
            ]
        )
        
        if filename:
            self.metadata_file_path.set(filename)
            self.update_metadata_status(f"Selected: {os.path.basename(filename)}")
    
    def extract_metadata(self):
        """Extract metadata from selected file"""
        file_path = self.metadata_file_path.get()
        
        if not file_path or not os.path.exists(file_path):
            messagebox.showerror("Error", "Please select a valid file")
            return
        
        try:
            self.update_metadata_status("🔍 Extracting metadata...")
            self.root.update()
            
            # Extract metadata
            metadata = self.metadata_extractor.extract_all_metadata(file_path)
            
            # Display in tree view
            self.display_metadata_tree(metadata)
            
            # Display in JSON view
            self.display_metadata_json(metadata)
            
            # Display basic info
            self.display_metadata_basic(metadata)
            
            self.update_metadata_status(f"✅ Metadata extracted successfully from {os.path.basename(file_path)}")
            
        except Exception as e:
            self.update_metadata_status(f"❌ Error extracting metadata: {str(e)}")
            messagebox.showerror("Error", f"Failed to extract metadata: {str(e)}")
    
    def display_metadata_tree(self, metadata):
        """Display metadata in tree view"""
        # Clear existing items
        for item in self.metadata_tree.get_children():
            self.metadata_tree.delete(item)
        
        def add_to_tree(parent, data, prefix=""):
            if isinstance(data, dict):
                for key, value in data.items():
                    if isinstance(value, (dict, list)):
                        node = self.metadata_tree.insert(parent, "end", text=key, values=("", ""))
                        add_to_tree(node, value, f"{prefix}.{key}" if prefix else key)
                    else:
                        self.metadata_tree.insert(parent, "end", text="", values=(key, str(value)))
            elif isinstance(data, list):
                for i, item in enumerate(data):
                    if isinstance(item, (dict, list)):
                        node = self.metadata_tree.insert(parent, "end", text=f"[{i}]", values=("", ""))
                        add_to_tree(node, item, f"{prefix}[{i}]")
                    else:
                        self.metadata_tree.insert(parent, "end", text="", values=(f"[{i}]", str(item)))
        
        # Add main categories
        for category, data in metadata.items():
            if data and not (isinstance(data, dict) and "error" in data):
                cat_node = self.metadata_tree.insert("", "end", text=category.upper(), values=("", ""))
                add_to_tree(cat_node, data, category)
    
    def display_metadata_json(self, metadata):
        """Display metadata in JSON format"""
        self.metadata_json_text.config(state=tk.NORMAL)
        self.metadata_json_text.delete(1.0, tk.END)
        
        try:
            formatted_json = json.dumps(metadata, indent=2, default=str)
            self.metadata_json_text.insert(tk.END, formatted_json)
        except Exception as e:
            self.metadata_json_text.insert(tk.END, f"Error formatting JSON: {str(e)}")
        
        self.metadata_json_text.config(state=tk.DISABLED)
    
    def display_metadata_basic(self, metadata):
        """Display basic metadata information"""
        self.metadata_basic_text.config(state=tk.NORMAL)
        self.metadata_basic_text.delete(1.0, tk.END)
        
        try:
            if "basic" in metadata:
                basic = metadata["basic"]
                text = "📄 BASIC FILE INFORMATION\n"
                text += "=" * 40 + "\n\n"
                
                for key, value in basic.items():
                    text += f"{key.replace('_', ' ').title()}: {value}\n"
                
                text += "\n" + "=" * 40 + "\n\n"
                
                if "file_type" in metadata:
                    text += f"📁 File Type: {metadata['file_type'].upper()}\n"
                
                if "extensions" in metadata:
                    text += f"🔧 Extracted Metadata Types: {', '.join(metadata['extensions'])}\n"
                
                # Show any errors
                for key, value in metadata.items():
                    if isinstance(value, dict) and "error" in value:
                        text += f"\n⚠️ Error in {key}: {value['error']}\n"
                
                self.metadata_basic_text.insert(tk.END, text)
            else:
                self.metadata_basic_text.insert(tk.END, "No basic metadata available")
        
        except Exception as e:
            self.metadata_basic_text.insert(tk.END, f"Error displaying basic info: {str(e)}")
        
        self.metadata_basic_text.config(state=tk.DISABLED)
    
    def export_metadata(self):
        """Export metadata to JSON file"""
        file_path = self.metadata_file_path.get()
        
        if not file_path or not os.path.exists(file_path):
            messagebox.showerror("Error", "No file selected or file doesn't exist")
            return
        
        # Ask for export location
        default_name = f"{os.path.splitext(os.path.basename(file_path))[0]}_metadata.json"
        export_path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
            initialfile=default_name
        )
        
        if export_path:
            try:
                # Extract metadata if not already extracted
                self.update_metadata_status("📤 Exporting metadata...")
                metadata = self.metadata_extractor.extract_all_metadata(file_path)
                
                # Export to JSON
                success, result = self.metadata_extractor.export_metadata_to_json(metadata, export_path)
                
                if success:
                    self.update_metadata_status(f"✅ Metadata exported to {os.path.basename(export_path)}")
                    messagebox.showinfo("Success", f"Metadata exported successfully to:\n{export_path}")
                else:
                    self.update_metadata_status(f"❌ Export failed: {result}")
                    messagebox.showerror("Error", f"Export failed: {result}")
            
            except Exception as e:
                self.update_metadata_status(f"❌ Export error: {str(e)}")
                messagebox.showerror("Error", f"Export failed: {str(e)}")
    
    def clear_metadata_display(self):
        """Clear all metadata displays"""
        # Clear tree
        for item in self.metadata_tree.get_children():
            self.metadata_tree.delete(item)
        
        # Clear JSON text
        self.metadata_json_text.config(state=tk.NORMAL)
        self.metadata_json_text.delete(1.0, tk.END)
        self.metadata_json_text.config(state=tk.DISABLED)
        
        # Clear basic text
        self.metadata_basic_text.config(state=tk.NORMAL)
        self.metadata_basic_text.delete(1.0, tk.END)
        self.metadata_basic_text.config(state=tk.DISABLED)
        
        # Clear file path
        self.metadata_file_path.set("")
        
        # Update status
        self.update_metadata_status("Ready to extract metadata")
    
    def update_metadata_status(self, message):
        """Update metadata status bar"""
        self.metadata_status_label.config(text=message)
    
    def create_file_operations_tab(self):
        """Create file encryption/decryption interface with modern design"""
        self.file_frame = tk.Frame(self.notebook, bg=self.colors['background'])
        self.notebook.add(self.file_frame, text="🔒 File Operations")
        
        # Main container
        container = tk.Frame(self.file_frame, bg=self.colors['background'])
        container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        # Title
        title_frame = tk.Frame(container, bg=self.colors['background'])
        title_frame.pack(fill=tk.X, pady=(0, 20))
        
        tk.Label(title_frame,
                text="🔐 File Encryption & Decryption",
                font=('Segoe UI', 20, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['background']).pack(side=tk.LEFT)
        
        # Card container
        card = tk.Frame(container, bg=self.colors['card_bg'], relief='ridge', borderwidth=1)
        card.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # File selection section
        file_section = tk.LabelFrame(card,
                                    text=" 📁 File Selection",
                                    font=('Segoe UI', 12, 'bold'),
                                    fg=self.colors['primary'],
                                    bg=self.colors['card_bg'],
                                    relief='groove',
                                    borderwidth=2)
        file_section.pack(fill=tk.X, padx=20, pady=20)
        
        tk.Label(file_section,
                text="Select a file to encrypt or decrypt:",
                font=('Segoe UI', 10),
                fg=self.colors['text_dark'],
                bg=self.colors['card_bg']).grid(row=0, column=0, sticky=tk.W, pady=10, padx=20)
        
        self.file_path_var = tk.StringVar()
        ttk.Entry(file_section,
                 textvariable=self.file_path_var,
                 font=('Segoe UI', 10),
                 width=50,
                 style='Custom.TEntry').grid(row=1, column=0, padx=20, pady=(0, 10))
        
        ttk.Button(file_section,
                  text="📂 Browse",
                  command=self.browse_file,
                  style='Primary.TButton',
                  width=15).grid(row=1, column=1, padx=(0, 20), pady=(0, 10))
        
        # Key section
        key_section = tk.LabelFrame(card,
                                   text=" 🔑 Encryption Key",
                                   font=('Segoe UI', 12, 'bold'),
                                   fg=self.colors['primary'],
                                   bg=self.colors['card_bg'],
                                   relief='groove',
                                    borderwidth=2)
        key_section.pack(fill=tk.X, padx=20, pady=20)
        
        tk.Label(key_section,
                text="AES Key (leave empty to generate new):",
                font=('Segoe UI', 10),
                fg=self.colors['text_dark'],
                bg=self.colors['card_bg']).grid(row=0, column=0, sticky=tk.W, pady=10, padx=20)
        
        self.aes_key_entry = ttk.Entry(key_section,
                                      font=('Segoe UI', 10),
                                      width=50,
                                      style='Custom.TEntry')
        self.aes_key_entry.grid(row=1, column=0, padx=20, pady=(0, 10))
        
        ttk.Button(key_section,
                  text="🎲 Generate Key",
                  command=self.generate_aes_key,
                  style='Primary.TButton',
                  width=15).grid(row=1, column=1, padx=(0, 20), pady=(0, 10))
        
        # Operation buttons
        button_frame = tk.Frame(card, bg=self.colors['card_bg'])
        button_frame.pack(pady=30)
        
        ttk.Button(button_frame,
                  text="🔒 Encrypt File",
                  command=self.encrypt_file,
                  style='Success.TButton',
                  width=20).pack(side=tk.LEFT, padx=10)
        
        ttk.Button(button_frame,
                  text="🔓 Decrypt File",
                  command=self.decrypt_file,
                  style='Danger.TButton',
                  width=20).pack(side=tk.LEFT, padx=10)
        
        # Status display
        status_section = tk.LabelFrame(card,
                                      text=" 📋 Activity Log",
                                      font=('Segoe UI', 12, 'bold'),
                                      fg=self.colors['primary'],
                                      bg=self.colors['card_bg'],
                                      relief='groove',
                                      borderwidth=2)
        status_section.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 20))
        
        self.status_text = scrolledtext.ScrolledText(status_section,
                                                   height=10,
                                                   font=('Consolas', 9),
                                                   bg=self.colors['card_bg'],
                                                   fg=self.colors['text_dark'],
                                                   relief='flat')
        self.status_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        self.status_text.config(state=tk.DISABLED)
    
    def create_file_sharing_tab(self):
        """Create file sharing interface with modern design"""
        self.share_frame = tk.Frame(self.notebook, bg=self.colors['background'])
        self.notebook.add(self.share_frame, text="🤝 File Sharing")
        
        # Create notebook for sharing sections
        share_notebook = ttk.Notebook(self.share_frame, style='Custom.TNotebook')
        share_notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Share File Tab
        self.create_share_file_tab(share_notebook)
        
        # Shared Files Tab
        self.create_shared_files_tab(share_notebook)
        
        # Public Key Directory Tab
        self.create_public_key_directory_tab(share_notebook)
    
    def create_share_file_tab(self, notebook):
        """Create share file tab"""
        share_tab = tk.Frame(notebook, bg=self.colors['background'])
        notebook.add(share_tab, text="📤 Share File")
        
        container = tk.Frame(share_tab, bg=self.colors['background'])
        container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        # Title
        tk.Label(container,
                text="Share File with Another User",
                font=('Segoe UI', 20, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['background']).pack(pady=(0, 20))
        
        # Card
        card = tk.Frame(container, bg=self.colors['card_bg'], relief='ridge', borderwidth=1)
        card.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Form fields
        fields = [
            ("📁 File to Share:", "share_file_path_var", True),
            ("📧 Recipient Email:", "recipient_email_var", False),
            ("🔑 AES Key for Encryption:", "share_aes_key_entry", False)
        ]
        
        for i, (label, attr_name, has_browse) in enumerate(fields):
            tk.Label(card,
                    text=label,
                    font=('Segoe UI', 10, 'bold'),
                    fg=self.colors['primary'],
                    bg=self.colors['card_bg']).grid(row=i*2, column=0, sticky=tk.W, pady=(20 if i == 0 else 10), padx=30)
            
            if attr_name == "share_file_path_var":
                self.share_file_path_var = tk.StringVar()
                entry = ttk.Entry(card,
                                 textvariable=self.share_file_path_var,
                                 font=('Segoe UI', 10),
                                 width=50,
                                 style='Custom.TEntry')
                entry.grid(row=i*2+1, column=0, padx=30, pady=(0, 10))
                
                ttk.Button(card,
                          text="📂 Browse",
                          command=lambda: self.browse_file(self.share_file_path_var),
                          style='Primary.TButton',
                          width=15).grid(row=i*2+1, column=1, padx=(0, 30), pady=(0, 10))
            elif attr_name == "share_aes_key_entry":
                self.share_aes_key_entry = ttk.Entry(card,
                                                    font=('Segoe UI', 10),
                                                    width=50,
                                                    style='Custom.TEntry')
                self.share_aes_key_entry.grid(row=i*2+1, column=0, padx=30, pady=(0, 10))
                
                ttk.Button(card,
                          text="🎲 Generate Key",
                          command=lambda: self.generate_aes_key(self.share_aes_key_entry),
                          style='Primary.TButton',
                          width=15).grid(row=i*2+1, column=1, padx=(0, 30), pady=(0, 10))
            else:
                self.recipient_email_var = tk.StringVar()
                recipient_frame = tk.Frame(card, bg=self.colors['card_bg'])
                recipient_frame.grid(row=i*2+1, column=0, columnspan=2, padx=30, pady=(0, 10), sticky=tk.W)
                
                entry = ttk.Entry(recipient_frame,
                                 textvariable=self.recipient_email_var,
                                 font=('Segoe UI', 10),
                                 width=40,
                                 style='Custom.TEntry')
                entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
                
                # Verify recipient button
                ttk.Button(recipient_frame,
                          text="✅ Verify",
                          command=self.verify_recipient,
                          style='Primary.TButton',
                          width=10).pack(side=tk.RIGHT, padx=(10, 0))
        
        # Recipient status display
        self.recipient_status_label = tk.Label(card,
                                              text="",
                                              font=('Segoe UI', 9),
                                              fg=self.colors['text_dark'],
                                              bg=self.colors['card_bg'])
        self.recipient_status_label.grid(row=6, column=0, columnspan=2, padx=30, pady=(5, 20), sticky=tk.W)
        
        # Share button
        button_frame = tk.Frame(card, bg=self.colors['card_bg'])
        button_frame.grid(row=7, column=0, columnspan=2, pady=20)
        
        ttk.Button(button_frame,
                  text="🚀 Share File",
                  command=self.share_file,
                  style='Success.TButton',
                  width=25).pack()
    
    def verify_recipient(self):
        """Verify if recipient exists and has a public key"""
        recipient_email = self.recipient_email_var.get().strip().lower()
        
        if not recipient_email:
            self.recipient_status_label.config(text="❌ Please enter recipient email", fg=self.colors['accent'])
            return
        
        if not self.validate_email(recipient_email):
            self.recipient_status_label.config(text="❌ Invalid email format", fg=self.colors['accent'])
            return
        
        # Check if recipient exists
        recipient_data = self.db.get_user(recipient_email)
        if not recipient_data:
            self.recipient_status_label.config(text="❌ Recipient not found in system", fg=self.colors['accent'])
            return
        
        # Check if recipient has public key
        public_key, message = self.get_recipient_public_key(recipient_email)
        if public_key:
            self.recipient_status_label.config(text=f"✅ Recipient verified! Public key available.", 
                                              fg=self.colors['success'])
        else:
            self.recipient_status_label.config(text=f"⚠️ Recipient found but no public key: {message}", 
                                              fg=self.colors['warning'])
    
    def create_public_key_directory_tab(self, notebook):
        """Create public key directory tab"""
        pubkey_tab = tk.Frame(notebook, bg=self.colors['background'])
        notebook.add(pubkey_tab, text="📇 Public Key Directory")
        
        container = tk.Frame(pubkey_tab, bg=self.colors['background'])
        container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        # Title
        tk.Label(container,
                text="Public Key Directory",
                font=('Segoe UI', 20, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['background']).pack(pady=(0, 20))
        
        # Card
        card = tk.Frame(container, bg=self.colors['card_bg'], relief='ridge', borderwidth=1)
        card.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Control buttons
        control_frame = tk.Frame(card, bg=self.colors['card_bg'])
        control_frame.pack(fill=tk.X, padx=20, pady=20)
        
        ttk.Button(control_frame,
                  text="🔄 Refresh Directory",
                  command=self.refresh_public_key_directory,
                  style='Primary.TButton',
                  width=20).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(control_frame,
                  text="📤 Upload My Public Key",
                  command=self.upload_my_public_key,
                  style='Success.TButton',
                  width=25).pack(side=tk.LEFT, padx=5)
        
        # Treeview for public keys
        tree_frame = tk.Frame(card, bg=self.colors['card_bg'])
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 20))
        
        columns = ("Email", "Last Updated", "Status")
        self.public_keys_tree = ttk.Treeview(tree_frame,
                                            columns=columns,
                                            show="headings",
                                            height=12,
                                            style='Custom.Treeview')
        
        column_widths = {"Email": 250, "Last Updated": 150, "Status": 100}
        for col in columns:
            self.public_keys_tree.heading(col, text=col)
            width = column_widths.get(col, 120)
            self.public_keys_tree.column(col, width=width, anchor='center')
        
        scrollbar = ttk.Scrollbar(tree_frame,
                                 orient="vertical",
                                 command=self.public_keys_tree.yview)
        self.public_keys_tree.configure(yscrollcommand=scrollbar.set)
        
        self.public_keys_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Refresh directory
        self.refresh_public_key_directory()
    
    def refresh_public_key_directory(self):
        """Refresh the public key directory from database"""
        try:
            # Clear existing items
            for item in self.public_keys_tree.get_children():
                self.public_keys_tree.delete(item)
            
            # Get all users
            all_users = self.db.get_all_users()
            # Get all public keys from database
            public_keys = self.db.get_all_public_keys()
            
            # Create a set of emails that have public keys
            pubkey_emails = {key['email'] for key in public_keys}
            
            # Add each user to the treeview
            for user in all_users:
                email = user['email']
                if email in pubkey_emails:
                    # Find the key info for this user
                    key_info = next((k for k in public_keys if k['email'] == email), None)
                    status = "✅ Available"
                    last_updated = key_info.get('last_updated', 'Unknown') if key_info else "Unknown"
                else:
                    status = "❌ Not Available"
                    last_updated = "Never"
                
                self.public_keys_tree.insert("", "end", values=(
                    email,
                    last_updated,
                    status
                ))
            
        except Exception as e:
            print(f"Error refreshing public key directory: {e}")
            messagebox.showerror("Error", f"Failed to refresh directory: {str(e)}")
    
    def upload_my_public_key(self):
        """Upload current user's public key to database"""
        if not self.user_public_key:
            messagebox.showerror("Error", "Public key not loaded. Please login again.")
            return
        
        try:
            public_key_pem = self.user_public_key.public_bytes(
                encoding=serialization.Encoding.PEM,
                format=serialization.PublicFormat.SubjectPublicKeyInfo
            )
            
            success = self.update_public_key(self.current_user, public_key_pem)
            if success:
                messagebox.showinfo("Success", "✅ Your public key has been uploaded to the repository!")
                self.refresh_public_key_directory()
            else:
                messagebox.showerror("Error", "Failed to upload public key to database")
                
        except Exception as e:
            messagebox.showerror("Error", f"Failed to upload public key: {str(e)}")
    
    def create_shared_files_tab(self, notebook):
        """Create shared files tab"""
        files_tab = tk.Frame(notebook, bg=self.colors['background'])
        notebook.add(files_tab, text="📥 Shared Files")
        
        container = tk.Frame(files_tab, bg=self.colors['background'])
        container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        # Title
        tk.Label(container,
                text="Files Shared with You",
                font=('Segoe UI', 20, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['background']).pack(pady=(0, 20))
        
        # Card
        card = tk.Frame(container, bg=self.colors['card_bg'], relief='ridge', borderwidth=1)
        card.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Control buttons
        control_frame = tk.Frame(card, bg=self.colors['card_bg'])
        control_frame.pack(fill=tk.X, padx=20, pady=20)
        
        ttk.Button(control_frame,
                  text="🔄 Refresh List",
                  command=self.refresh_shared_files,
                  style='Primary.TButton',
                  width=15).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(control_frame,
                  text="⬇️ Download & Decrypt",
                  command=self.download_shared_file,
                  style='Success.TButton',
                  width=25).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(control_frame,
                  text="🔑 Get Decryption Key",
                  command=self.decrypt_shared_file_key,
                  style='Warning.TButton',
                  width=25).pack(side=tk.LEFT, padx=5)
        
        # Treeview for shared files
        tree_frame = tk.Frame(card, bg=self.colors['card_bg'])
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 20))
        
        columns = ("Filename", "Shared By", "Date", "Status")
        self.shared_files_tree = ttk.Treeview(tree_frame,
                                             columns=columns,
                                             show="headings",
                                             height=12,
                                             style='Custom.Treeview')
        
        for col in columns:
            self.shared_files_tree.heading(col, text=col)
            width = 180 if col == "Filename" else 120
            self.shared_files_tree.column(col, width=width, anchor='center')
        
        scrollbar = ttk.Scrollbar(tree_frame,
                                 orient="vertical",
                                 command=self.shared_files_tree.yview)
        self.shared_files_tree.configure(yscrollcommand=scrollbar.set)
        
        self.shared_files_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Key input for decryption
        key_frame = tk.Frame(card, bg=self.colors['card_bg'])
        key_frame.pack(fill=tk.X, padx=20, pady=(0, 20))
        
        tk.Label(key_frame,
                text="🔑 AES Key for Decryption:",
                font=('Segoe UI', 10, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['card_bg']).pack(side=tk.LEFT, padx=(0, 10))
        
        self.download_aes_key_entry = ttk.Entry(key_frame,
                                               font=('Segoe UI', 10),
                                               width=50,
                                               style='Custom.TEntry')
        self.download_aes_key_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        
        self.refresh_shared_files()
    
    def create_key_management_tab(self):
        """Create key management interface with modern design"""
        self.key_frame = tk.Frame(self.notebook, bg=self.colors['background'])
        self.notebook.add(self.key_frame, text="🔑 Key Management")
        
        # Create notebook for key management sections
        key_notebook = ttk.Notebook(self.key_frame, style='Custom.TNotebook')
        key_notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # RSA Keys Tab
        self.create_rsa_keys_tab(key_notebook)
        
        # AES Operations Tab
        self.create_aes_operations_tab(key_notebook)
        
        # User RSA Key Management Tab (Integrated with test)
        self.create_user_rsa_key_tab(key_notebook)
    
    def create_rsa_keys_tab(self, notebook):
        """Create RSA keys tab"""
        rsa_tab = tk.Frame(notebook, bg=self.colors['background'])
        notebook.add(rsa_tab, text="🔐 RSA Keys")
        
        container = tk.Frame(rsa_tab, bg=self.colors['background'])
        container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        # Title
        tk.Label(container,
                text="RSA Key Information",
                font=('Segoe UI', 20, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['background']).pack(pady=(0, 20))
        
        # Card
        card = tk.Frame(container, bg=self.colors['card_bg'], relief='ridge', borderwidth=1)
        card.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Public key display
        tk.Label(card,
                text="Your Public Key:",
                font=('Segoe UI', 12, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['card_bg']).pack(anchor='w', padx=30, pady=(20, 10))
        
        public_key_frame = tk.Frame(card, bg=self.colors['card_bg'])
        public_key_frame.pack(fill=tk.BOTH, expand=True, padx=30, pady=(0, 20))
        
        self.public_key_text = scrolledtext.ScrolledText(public_key_frame,
                                                        height=8,
                                                        font=('Consolas', 9),
                                                        bg='#F8F9FA',
                                                        fg=self.colors['primary'],
                                                        relief='solid',
                                                        borderwidth=1)
        self.public_key_text.pack(fill=tk.BOTH, expand=True)
        
        # Load current user's public key
        if self.user_public_key:
            public_key_pem = self.user_public_key.public_bytes(
                encoding=serialization.Encoding.PEM,
                format=serialization.PublicFormat.SubjectPublicKeyInfo
            ).decode()
            self.public_key_text.insert(tk.END, public_key_pem)
            self.public_key_text.config(state=tk.DISABLED)
        else:
            self.public_key_text.insert(tk.END, "Public key not loaded. Please login again.")
            self.public_key_text.config(state=tk.DISABLED)
        
        # Info
        info_frame = tk.Frame(card, bg=self.colors['card_bg'])
        info_frame.pack(fill=tk.X, padx=30, pady=(0, 20))
        
        tk.Label(info_frame,
                text="ℹ️ Your public key is used by others to encrypt AES keys for you.",
                font=('Segoe UI', 9, 'italic'),
                fg=self.colors['secondary'],
                bg=self.colors['card_bg']).pack(anchor='w')
        
        tk.Label(info_frame,
                text="ℹ️ Your private key is securely encrypted with your password.",
                font=('Segoe UI', 9, 'italic'),
                fg=self.colors['secondary'],
                bg=self.colors['card_bg']).pack(anchor='w')
        
        tk.Label(info_frame,
                text="ℹ️ Make sure your public key is uploaded to the repository for sharing.",
                font=('Segoe UI', 9, 'italic'),
                fg=self.colors['secondary'],
                bg=self.colors['card_bg']).pack(anchor='w')
    
    def create_user_rsa_key_tab(self, notebook):
        """Create user RSA key management tab with integrated test section"""
        user_rsa_tab = tk.Frame(notebook, bg=self.colors['background'])
        notebook.add(user_rsa_tab, text="👤 User RSA Keys")
        
        # Main container with vertical layout
        main_container = tk.Frame(user_rsa_tab, bg=self.colors['background'])
        main_container.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Title
        title_label = tk.Label(main_container,
                              text="👤 User RSA Key Management",
                              font=('Segoe UI', 20, 'bold'),
                              fg=self.colors['primary'],
                              bg=self.colors['background'])
        title_label.pack(pady=(0, 20))
        
        # Create a notebook for the two main sections
        rsa_notebook = ttk.Notebook(main_container, style='Custom.TNotebook')
        rsa_notebook.pack(fill=tk.BOTH, expand=True)
        
        # Tab 1: RSA Key Management
        key_management_tab = tk.Frame(rsa_notebook, bg=self.colors['card_bg'])
        rsa_notebook.add(key_management_tab, text="🔑 RSA Key Management")
        
        # Tab 2: RSA Encryption/Decryption Test
        test_tab = tk.Frame(rsa_notebook, bg=self.colors['card_bg'])
        rsa_notebook.add(test_tab, text="🔐 RSA Encryption/Decryption Test")
        
        # Build the tabs
        self.create_rsa_key_management_section(key_management_tab)
        self.create_rsa_test_section(test_tab)
    
    def create_rsa_key_management_section(self, parent):
        """Create RSA key management section"""
        container = tk.Frame(parent, bg=self.colors['card_bg'])
        container.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Status header
        status_frame = tk.Frame(container, bg=self.colors['card_bg'])
        status_frame.pack(fill=tk.X, pady=(0, 20))
        
        key_status = "✅ RSA Keys: Loaded" if self.user_private_key else "❌ RSA Keys: Not Loaded"
        color = self.colors['success'] if self.user_private_key else self.colors['accent']
        
        tk.Label(status_frame,
                text=key_status,
                font=('Segoe UI', 14, 'bold'),
                fg=color,
                bg=self.colors['card_bg']).pack()
        
        # Options header
        tk.Label(container,
                text="📋 RSA Key Options:",
                font=('Segoe UI', 12, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['card_bg']).pack(anchor='w', pady=(0, 10))
        
        # Buttons frame
        buttons_frame = tk.Frame(container, bg=self.colors['card_bg'])
        buttons_frame.pack(fill=tk.X, pady=(0, 30))
        
        # Create buttons in a grid
        button_definitions = [
            ("🔄 Generate/Regenerate RSA Keys", self.regenerate_rsa_keys, "Primary.TButton"),
            ("📤 Export Public Key", self.export_public_key, "Success.TButton"),
            ("🌐 Upload to Repository", self.upload_my_public_key, "Warning.TButton"),
            ("📧 E-mail Public Key", lambda: messagebox.showinfo("Info", "E-mail functionality would be added here"), "JSONPrimary.TButton")
        ]
        
        for i, (text, command, style_name) in enumerate(button_definitions):
            btn = ttk.Button(buttons_frame,
                           text=text,
                           command=command,
                           style=style_name,
                           width=30)
            btn.grid(row=i//2, column=i%2, padx=10, pady=10, sticky='ew')
        
        # Configure grid
        buttons_frame.grid_columnconfigure(0, weight=1)
        buttons_frame.grid_columnconfigure(1, weight=1)
        
        # Current RSA Key Info
        info_frame = tk.LabelFrame(container,
                                  text="📄 Current RSA Key Information",
                                  font=('Segoe UI', 11, 'bold'),
                                  fg=self.colors['primary'],
                                  bg=self.colors['card_bg'],
                                  relief='groove',
                                  borderwidth=2)
        info_frame.pack(fill=tk.BOTH, expand=True, pady=(10, 0))
        
        # Info text
        info_text = tk.Text(info_frame,
                           height=8,
                           font=('Consolas', 10),
                           bg='#F8F9FA',
                           fg=self.colors['text_dark'],
                           relief='flat',
                           borderwidth=1,
                           wrap=tk.WORD)
        info_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Add key information
        if self.user_public_key:
            public_key_pem = self.user_public_key.public_bytes(
                encoding=serialization.Encoding.PEM,
                format=serialization.PublicFormat.SubjectPublicKeyInfo
            ).decode()
            
            key_info = f"""🔐 RSA Key Status: Loaded
🔑 Key Size: 2048 bits
📍 Public Key Location: {os.path.join(self.user_rsa_keys_dir, f"{self.current_user}_public.pem")}
📍 Private Key Location: {os.path.join(self.user_rsa_keys_dir, f"{self.current_user}_private.enc")}

📋 Public Key (first 100 chars):
{public_key_pem[:100]}..."""
            
            info_text.insert(tk.END, key_info)
        else:
            info_text.insert(tk.END, "❌ RSA keys not loaded. Please login or generate new keys.")
        
        info_text.config(state=tk.DISABLED)
    
    def create_rsa_test_section(self, parent):
        """Create RSA encryption/decryption test section"""
        container = tk.Frame(parent, bg=self.colors['card_bg'])
        container.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Title
        tk.Label(container,
                text="🔐 RSA Encryption/Decryption Test",
                font=('Segoe UI', 16, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['card_bg']).pack(pady=(0, 20))
        
        # Test input section
        input_frame = tk.LabelFrame(container,
                                   text="📝 Test Message",
                                   font=('Segoe UI', 11, 'bold'),
                                   fg=self.colors['primary'],
                                   bg=self.colors['card_bg'],
                                   relief='groove',
                                   borderwidth=2)
        input_frame.pack(fill=tk.X, pady=(0, 15))
        
        tk.Label(input_frame,
                text="Enter a test message to encrypt and decrypt:",
                font=('Segoe UI', 10),
                fg=self.colors['text_dark'],
                bg=self.colors['card_bg']).pack(anchor='w', padx=10, pady=(10, 5))
        
        # Message input
        message_frame = tk.Frame(input_frame, bg=self.colors['card_bg'])
        message_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        
        self.test_message_var = tk.StringVar(value="Test RSA encryption/decryption")
        test_entry = ttk.Entry(message_frame,
                              textvariable=self.test_message_var,
                              font=('Segoe UI', 11),
                              style='Custom.TEntry')
        test_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        
        test_btn = ttk.Button(message_frame,
                            text="🔒 Test Encrypt/Decrypt",
                            command=self.test_rsa_encryption,
                            style='Warning.TButton',
                            width=25)
        test_btn.pack(side=tk.RIGHT)
        
        # Results section
        results_frame = tk.LabelFrame(container,
                                     text="📊 Test Results",
                                     font=('Segoe UI', 11, 'bold'),
                                     fg=self.colors['primary'],
                                     bg=self.colors['card_bg'],
                                     relief='groove',
                                     borderwidth=2)
        results_frame.pack(fill=tk.BOTH, expand=True, pady=(10, 0))
        
        # Create a frame with a border for the text widget
        text_container = tk.Frame(results_frame, bg=self.colors['primary'], relief='sunken', borderwidth=2)
        text_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Results text area
        self.test_result_text = scrolledtext.ScrolledText(text_container,
                                                         height=20,
                                                         font=('Courier New', 10),
                                                         bg='#0C0C0C',
                                                         fg='#00FF00',
                                                         relief='flat',
                                                         borderwidth=0,
                                                         wrap=tk.WORD,
                                                         spacing1=2,
                                                         spacing3=2)
        self.test_result_text.pack(fill=tk.BOTH, expand=True, padx=1, pady=1)
        
        # Add initial placeholder text
        placeholder = """RSA ENCRYPTION/DECRYPTION TEST RESULTS

This area will display the complete RSA encryption/decryption test results.
Enter a test message above and click "Test Encrypt/Decrypt" to see:

1. Original message
2. Encrypted message (Base64 format)
3. Decrypted message
4. Verification result
5. Encryption/Decryption times

─────────────────────────────────────────────────────
Example output format:
Original: Test message
Encrypted: MIICXQIBAAKBgQCqGKukO1De7zhZj6+H0qtjTkVxwTCpvKe4eCZ0FPqri0cb2JZfXJ/DgYSF6vUp...
Decrypted: Test message
✅ Success: Messages match!
─────────────────────────────────────────────────────"""
        
        self.test_result_text.insert(tk.END, placeholder)
        self.test_result_text.config(state=tk.DISABLED)
        
        # Add context menu for copy/paste
        def make_selectable():
            if self.test_result_text['state'] == tk.DISABLED:
                self.test_result_text.config(state=tk.NORMAL)
                self.root.after(100, lambda: self.test_result_text.config(state=tk.DISABLED))
        
        context_menu = tk.Menu(self.test_result_text, tearoff=0)
        context_menu.add_command(label="📋 Copy", 
                               command=lambda: self.root.clipboard_clear() or 
                                              self.root.clipboard_append(self.test_result_text.selection_get()) 
                                              if self.test_result_text.tag_ranges('sel') else None)
        context_menu.add_command(label="📄 Select All", 
                               command=lambda: self.test_result_text.tag_add('sel', '1.0', 'end'))
        context_menu.add_separator()
        context_menu.add_command(label="🧹 Clear Results", 
                               command=lambda: [self.test_result_text.config(state=tk.NORMAL),
                                               self.test_result_text.delete(1.0, tk.END),
                                               self.test_result_text.insert(tk.END, placeholder),
                                               self.test_result_text.config(state=tk.DISABLED)])
        
        def show_context_menu(event):
            try:
                context_menu.tk_popup(event.x_root, event.y_root)
            finally:
                context_menu.grab_release()
        
        self.test_result_text.bind("<Button-3>", show_context_menu)
    
    def regenerate_rsa_keys(self):
        """Regenerate RSA keys for current user"""
        if not messagebox.askyesno("Confirm", 
                                 "This will regenerate your RSA keys.\n"
                                 "Old files shared with you will become inaccessible unless re-shared.\n"
                                 "Continue?"):
            return
        
        # Get current password from user
        password_dialog = tk.Toplevel(self.root)
        password_dialog.title("Enter Password")
        password_dialog.geometry("400x250")
        password_dialog.configure(bg=self.colors['background'])
        password_dialog.transient(self.root)
        password_dialog.grab_set()
        
        self.center_window(password_dialog, 400, 250)
        
        tk.Label(password_dialog,
                text="Enter your current password to regenerate RSA keys:",
                font=('Segoe UI', 10),
                fg=self.colors['primary'],
                bg=self.colors['background']).pack(pady=20)
        
        password_var = tk.StringVar()
        ttk.Entry(password_dialog,
                 textvariable=password_var,
                 font=('Segoe UI', 10),
                 show="*",
                 style='Custom.TEntry',
                 width=30).pack(pady=10)
        
        # Upload to repository option
        upload_var = tk.BooleanVar(value=True)
        upload_check = tk.Checkbutton(password_dialog,
                                     text="Upload new public key to repository",
                                     variable=upload_var,
                                     font=('Segoe UI', 9),
                                     fg=self.colors['primary'],
                                     bg=self.colors['background'])
        upload_check.pack(pady=10)
        
        def confirm_regenerate():
            password = password_var.get()
            if not password:
                messagebox.showerror("Error", "Password required", parent=password_dialog)
                return
            
            # Verify password
            user_data = self.db.get_user(self.current_user)
            if not user_data:
                messagebox.showerror("Error", "User not found", parent=password_dialog)
                return
            
            if user_data["password_hash"] != self.hash_password(password):
                messagebox.showerror("Error", "Incorrect password", parent=password_dialog)
                return
            
            # Regenerate keys
            if self.generate_user_rsa_keys(self.current_user, password):
                # Reload keys
                success, message = self.load_user_rsa_keys(self.current_user, password)
                if success:
                    # Upload to repository if requested
                    if upload_var.get() and self.user_public_key:
                        public_key_pem = self.user_public_key.public_bytes(
                            encoding=serialization.Encoding.PEM,
                            format=serialization.PublicFormat.SubjectPublicKeyInfo
                        )
                        self.update_public_key(self.current_user, public_key_pem)
                    
                    messagebox.showinfo("Success", 
                                      "✅ RSA keys regenerated successfully!\n"
                                      "Public key has been updated in the repository.",
                                      parent=password_dialog)
                    
                    # Refresh public key directory if it exists
                    if hasattr(self, 'public_keys_tree'):
                        self.refresh_public_key_directory()
                else:
                    messagebox.showwarning("Warning", 
                                         f"Keys generated but not loaded: {message}", 
                                         parent=password_dialog)
            else:
                messagebox.showerror("Error", "Failed to regenerate RSA keys", parent=password_dialog)
            
            password_dialog.destroy()
        
        button_frame = tk.Frame(password_dialog, bg=self.colors['background'])
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame,
                  text="🔄 Regenerate",
                  command=confirm_regenerate,
                  style='Success.TButton',
                  width=15).pack(side=tk.LEFT, padx=10)
        
        ttk.Button(button_frame,
                  text="❌ Cancel",
                  command=password_dialog.destroy,
                  style='Danger.TButton',
                  width=15).pack(side=tk.RIGHT, padx=10)
    
    def export_public_key(self):
        """Export public key to a file"""
        if not self.user_public_key:
            messagebox.showerror("Error", "Public key not loaded")
            return
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".pem",
            filetypes=[("PEM files", "*.pem"), ("All files", "*.*")],
            initialfile=f"{self.current_user}_public_key.pem"
        )
        
        if filename:
            try:
                public_key_pem = self.user_public_key.public_bytes(
                    encoding=serialization.Encoding.PEM,
                    format=serialization.PublicFormat.SubjectPublicKeyInfo
                )
                
                with open(filename, 'wb') as f:
                    f.write(public_key_pem)
                
                messagebox.showinfo("Success", f"✅ Public key exported to:\n{filename}")
                self.log_action("EXPORT_PUBLIC_KEY", filename)
                
            except Exception as e:
                messagebox.showerror("Error", f"Failed to export public key: {str(e)}")
    
    def test_rsa_encryption(self):
        """Test RSA encryption and decryption"""
        if not self.user_private_key or not self.user_public_key:
            messagebox.showerror("Error", "RSA keys not loaded")
            return
        
        message = self.test_message_var.get()
        if not message:
            messagebox.showerror("Error", "Please enter a test message")
            return
        
        try:
            start_time = time.time()
            
            # Encrypt with public key
            encrypted = self.user_public_key.encrypt(
                message.encode(),
                padding.OAEP(
                    mgf=padding.MGF1(algorithm=hashes.SHA256()),
                    algorithm=hashes.SHA256(),
                    label=None
                )
            )
            
            encryption_time = time.time() - start_time
            
            # Decrypt with private key
            decrypted = self.user_private_key.decrypt(
                encrypted,
                padding.OAEP(
                    mgf=padding.MGF1(algorithm=hashes.SHA256()),
                    algorithm=hashes.SHA256(),
                    label=None
                )
            )
            
            decryption_time = time.time() - start_time - encryption_time
            
            # Display results
            self.test_result_text.config(state=tk.NORMAL)
            self.test_result_text.delete(1.0, tk.END)
            
            result = f"🔐 RSA ENCRYPTION/DECRYPTION TEST RESULTS\n"
            result += "=" * 60 + "\n\n"
            
            result += f"📝 Original Message:\n{message}\n\n"
            
            result += f"🔒 Encrypted (Base64):\n{base64.b64encode(encrypted).decode()}\n\n"
            
            result += f"🔓 Decrypted Message:\n{decrypted.decode()}\n\n"
            
            result += f"✅ Verification: {'SUCCESS - Messages match!' if message.encode() == decrypted else 'FAILED - Messages do not match!'}\n\n"
            
            result += f"⏱️ Performance Metrics:\n"
            result += f"   • Encryption Time: {encryption_time:.6f} seconds\n"
            result += f"   • Decryption Time: {decryption_time:.6f} seconds\n"
            result += f"   • Total Time: {(encryption_time + decryption_time):.6f} seconds\n\n"
            
            result += f"🔑 Key Information:\n"
            result += f"   • Key Size: 2048 bits\n"
            result += f"   • Algorithm: RSA with OAEP padding\n"
            result += f"   • Hash Algorithm: SHA-256\n"
            
            result += "=" * 60 + "\n"
            result += "✅ Test completed successfully!"
            
            self.test_result_text.insert(tk.END, result)
            self.test_result_text.config(state=tk.DISABLED)
            
            # Auto-scroll to top
            self.test_result_text.see(1.0)
            
        except Exception as e:
            messagebox.showerror("Error", f"RSA test failed: {str(e)}")
    
    def create_aes_operations_tab(self, notebook):
        """Create AES operations tab"""
        aes_tab = tk.Frame(notebook, bg=self.colors['background'])
        notebook.add(aes_tab, text="🔑 AES Operations")
        
        container = tk.Frame(aes_tab, bg=self.colors['background'])
        container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        # Title
        tk.Label(container,
                text="AES Key Encryption/Decryption",
                font=('Segoe UI', 20, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['background']).pack(pady=(0, 20))
        
        # Card
        card = tk.Frame(container, bg=self.colors['card_bg'], relief='ridge', borderwidth=1)
        card.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Input section
        tk.Label(card,
                text="AES Key to Encrypt/Decrypt:",
                font=('Segoe UI', 11, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['card_bg']).pack(anchor='w', padx=30, pady=(20, 10))
        
        input_frame = tk.Frame(card, bg=self.colors['card_bg'])
        input_frame.pack(fill=tk.X, padx=30, pady=(0, 20))
        
        self.aes_key_to_encrypt = ttk.Entry(input_frame,
                                           font=('Segoe UI', 10),
                                           width=60,
                                           style='Custom.TEntry')
        self.aes_key_to_encrypt.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        
        # Buttons
        button_frame = tk.Frame(card, bg=self.colors['card_bg'])
        button_frame.pack(pady=10)
        
        ttk.Button(button_frame,
                  text="🔒 Encrypt Key",
                  command=self.encrypt_aes_key,
                  style='Success.TButton',
                  width=20).pack(side=tk.LEFT, padx=10)
        
        ttk.Button(button_frame,
                  text="🔓 Decrypt Key",
                  command=self.decrypt_aes_key,
                  style='Danger.TButton',
                  width=20).pack(side=tk.LEFT, padx=10)
        
        # Result section
        tk.Label(card,
                text="Result:",
                font=('Segoe UI', 11, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['card_bg']).pack(anchor='w', padx=30, pady=(20, 10))
        
        result_frame = tk.Frame(card, bg=self.colors['card_bg'])
        result_frame.pack(fill=tk.BOTH, expand=True, padx=30, pady=(0, 20))
        
        self.key_result_text = scrolledtext.ScrolledText(result_frame,
                                                        height=6,
                                                        font=('Consolas', 9),
                                                        bg='#F8F9FA',
                                                        fg=self.colors['primary'],
                                                        relief='solid',
                                                        borderwidth=1)
        self.key_result_text.pack(fill=tk.BOTH, expand=True)
        self.key_result_text.config(state=tk.DISABLED)
    
    def create_admin_dashboard_tab(self):
        """Create admin dashboard with modern design"""
        self.log_frame = tk.Frame(self.notebook, bg=self.colors['background'])
        self.notebook.add(self.log_frame, text="📊 Admin Dashboard")
        
        # Create notebook for admin sections
        admin_notebook = ttk.Notebook(self.log_frame, style='Custom.TNotebook')
        admin_notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Statistics Tab
        self.create_stats_tab(admin_notebook)
        
        # Graphs Tab (NEW)
        self.create_graphs_tab(admin_notebook)
        
        # Sessions Tab
        self.create_sessions_tab(admin_notebook)
        
        # User Management Tab
        self.create_user_management_tab(admin_notebook)
        
        # Logs Tab
        self.create_logs_tab(admin_notebook)
        
        # Database Management Tab
        self.create_database_management_tab(admin_notebook)
    
    def create_stats_tab(self, notebook):
        """Create statistics tab without graphs"""
        stats_tab = tk.Frame(notebook, bg=self.colors['background'])
        notebook.add(stats_tab, text="📈 Statistics")
        
        container = tk.Frame(stats_tab, bg=self.colors['background'])
        container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        # Title
        tk.Label(container,
                text="System Statistics",
                font=('Segoe UI', 20, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['background']).pack(pady=(0, 20))
        
        # Get statistics from database
        stats = self.db.get_system_statistics()
        
        # Create stats cards in a grid
        stats_frame = tk.Frame(container, bg=self.colors['background'])
        stats_frame.pack(fill=tk.BOTH, expand=True)
        
        # Stat cards
        stat_cards = [
            ("👥 Total Users", stats['total_users'], self.colors['secondary'], "Users registered in the system"),
            ("🔑 Total Logins", stats['total_logins'], self.colors['success'], "Total login attempts"),
            ("📁 Files Shared", stats['files_shared'], self.colors['warning'], "Total files shared"),
            ("🟢 Active Sessions", stats['currently_logged_in'], self.colors['accent'], "Currently logged in users"),
            ("👑 Administrators", stats['admin_count'], self.colors['primary'], "Admin users"),
            ("👤 Regular Users", stats['user_count'], self.colors['secondary'], "Regular users"),
            ("🔐 Users with Keys", stats['users_with_pubkeys'], self.colors['success'], "Users with RSA keys"),
            ("📊 Login Average", f"{stats['total_logins']/max(stats['total_users'],1):.1f}", 
             self.colors['warning'], "Logins per user")
        ]
        
        # Create cards in a grid
        for i, (title, value, color, desc) in enumerate(stat_cards):
            row = i // 4
            col = i % 4
            
            card = tk.Frame(stats_frame,
                       bg=self.colors['card_bg'],
                       relief='ridge',
                       borderwidth=2,
                       padx=10,
                       pady=10)
            card.grid(row=row, column=col, padx=10, pady=10, sticky='nsew')
            
            # Title
            tk.Label(card,
                    text=title,
                    font=('Segoe UI', 11, 'bold'),
                    fg=color,
                    bg=self.colors['card_bg']).pack(pady=(5, 0))
            
            # Value
            tk.Label(card,
                    text=str(value),
                    font=('Segoe UI', 24, 'bold'),
                    fg=self.colors['primary'],
                    bg=self.colors['card_bg']).pack(pady=(5, 5))
            
            # Description
            tk.Label(card,
                    text=desc,
                    font=('Segoe UI', 8),
                    fg=self.colors['text_dark'],
                    bg=self.colors['card_bg'],
                    wraplength=150).pack(pady=(0, 5))
        
        # Configure grid
        for i in range(4):
            stats_frame.grid_columnconfigure(i, weight=1)
        for i in range(2):
            stats_frame.grid_rowconfigure(i, weight=1)
        
        # View Graphs button
        button_frame = tk.Frame(container, bg=self.colors['background'])
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame,
                  text="📊 View Detailed Graphs",
                  command=lambda: notebook.select(notebook.index(notebook.select()) + 1),
                  style='Primary.TButton',
                  width=25).pack()
        
        # JSON Statistics
        try:
            json_stats = self.get_json_stats()
            json_frame = tk.Frame(container, bg=self.colors['card_bg'], relief='ridge', borderwidth=1)
            json_frame.pack(fill=tk.X, pady=20, padx=10)
            
            tk.Label(json_frame,
                    text="📄 JSON Statistics",
                    font=('Segoe UI', 12, 'bold'),
                    fg=self.colors['primary'],
                    bg=self.colors['card_bg']).pack(pady=10)
            
            json_info = tk.Frame(json_frame, bg=self.colors['card_bg'])
            json_info.pack(pady=(0, 10))
            
            tk.Label(json_info,
                    text=f"Users in JSON: {json_stats['total_users']}",
                    font=('Segoe UI', 10),
                    fg=self.colors['text_dark'],
                    bg=self.colors['card_bg']).pack(side=tk.LEFT, padx=20)
            
            tk.Label(json_info,
                    text=f"Last Updated: {json_stats['last_updated']}",
                    font=('Segoe UI', 10),
                    fg=self.colors['text_dark'],
                    bg=self.colors['card_bg']).pack(side=tk.LEFT, padx=20)
            
        except Exception as e:
            print(f"JSON stats error: {e}")
    
    def create_graphs_tab(self, notebook):
        """Create graphs tab with visualizations"""
        graphs_tab = tk.Frame(notebook, bg=self.colors['background'])
        notebook.add(graphs_tab, text="📊 Graphs")
        
        container = tk.Frame(graphs_tab, bg=self.colors['background'])
        container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        # Title
        tk.Label(container,
                text="Visual Statistics Dashboard",
                font=('Segoe UI', 20, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['background']).pack(pady=(0, 20))
        
        # Control buttons
        control_frame = tk.Frame(container, bg=self.colors['background'])
        control_frame.pack(fill=tk.X, pady=(0, 20))
        
        ttk.Button(control_frame,
                  text="🔄 Refresh Graphs",
                  command=lambda: self.refresh_graphs(canvas_container),
                  style='Primary.TButton',
                  width=20).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(control_frame,
                  text="💾 Save Graphs",
                  command=self.save_graphs,
                  style='Success.TButton',
                  width=20).pack(side=tk.LEFT, padx=5)
        
        # Canvas for scrolling
        canvas = tk.Canvas(container, bg=self.colors['background'], highlightthickness=0)
        scrollbar = ttk.Scrollbar(container, orient="vertical", command=canvas.yview)
        canvas_container = tk.Frame(canvas, bg=self.colors['background'])
        
        canvas_container.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=canvas_container, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True, padx=(0, 10))
        scrollbar.pack(side="right", fill="y")
        
        # Generate initial graphs
        self.refresh_graphs(canvas_container)
        
        # Bind mouse wheel
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        canvas.bind_all("<MouseWheel>", on_mousewheel)
    
    def refresh_graphs(self, container):
        """Refresh and display all graphs"""
        # Clear previous graphs
        for widget in container.winfo_children():
            widget.destroy()
        
        # Get statistics
        stats = self.db.get_system_statistics()
        
        # Graph 1: Main statistics bar chart
        graph1_frame = tk.Frame(container, bg=self.colors['card_bg'], relief='ridge', borderwidth=2)
        graph1_frame.pack(fill=tk.X, padx=10, pady=10)
        
        tk.Label(graph1_frame,
                text="📊 System Overview",
                font=('Segoe UI', 14, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['card_bg']).pack(pady=10)
        
        fig1 = Figure(figsize=(10, 5), dpi=80, facecolor=self.colors['card_bg'])
        ax1 = fig1.add_subplot(111)
        
        categories = ['Total Users', 'Total Logins', 'Files Shared', 'Active Sessions']
        values = [stats['total_users'], stats['total_logins'], 
                 stats['files_shared'], stats['currently_logged_in']]
        colors = [self.colors['secondary'], self.colors['success'], 
                 self.colors['warning'], self.colors['accent']]
        
        bars = ax1.bar(categories, values, color=colors, edgecolor='black')
        ax1.set_title('System Statistics Overview', fontsize=14, fontweight='bold')
        ax1.set_ylabel('Count', fontsize=12)
        ax1.grid(True, alpha=0.3)
        
        # Add value labels
        for bar, value in zip(bars, values):
            height = bar.get_height()
            ax1.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                    f'{value}', ha='center', va='bottom', fontweight='bold')
        
        canvas1 = FigureCanvasTkAgg(fig1, master=graph1_frame)
        canvas1.draw()
        canvas1.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Graph 2: User distribution pie chart
        graph2_frame = tk.Frame(container, bg=self.colors['card_bg'], relief='ridge', borderwidth=2)
        graph2_frame.pack(fill=tk.X, padx=10, pady=10)
        
        tk.Label(graph2_frame,
                text="👥 User Distribution",
                font=('Segoe UI', 14, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['card_bg']).pack(pady=10)
        
        fig2 = Figure(figsize=(10, 5), dpi=80, facecolor=self.colors['card_bg'])
        
        # Pie chart
        ax2 = fig2.add_subplot(121)
        user_types = ['Admins', 'Regular Users']
        user_counts = [stats['admin_count'], stats['user_count']]
        user_colors = [self.colors['accent'], self.colors['secondary']]
        
        wedges, texts, autotexts = ax2.pie(user_counts, labels=user_types, colors=user_colors,
                                          autopct='%1.1f%%', startangle=90)
        ax2.set_title('User Type Distribution', fontsize=12, fontweight='bold')
        
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_fontweight('bold')
        
        # Bar chart for users with keys
        ax3 = fig2.add_subplot(122)
        key_categories = ['With Keys', 'Without Keys']
        key_values = [stats['users_with_pubkeys'], 
                     max(0, stats['total_users'] - stats['users_with_pubkeys'])]
        key_colors = [self.colors['success'], self.colors['warning']]
        
        bars2 = ax3.bar(key_categories, key_values, color=key_colors, edgecolor='black')
        ax3.set_title('Users with RSA Keys', fontsize=12, fontweight='bold')
        ax3.set_ylabel('Count', fontsize=10)
        ax3.grid(True, alpha=0.3)
        
        for bar, value in zip(bars2, key_values):
            height = bar.get_height()
            ax3.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                    f'{value}', ha='center', va='bottom', fontweight='bold')
        
        canvas2 = FigureCanvasTkAgg(fig2, master=graph2_frame)
        canvas2.draw()
        canvas2.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Graph 3: Recent login activity (simulated)
        graph3_frame = tk.Frame(container, bg=self.colors['card_bg'], relief='ridge', borderwidth=2)
        graph3_frame.pack(fill=tk.X, padx=10, pady=10)
        
        tk.Label(graph3_frame,
                text="📈 Activity Trends",
                font=('Segoe UI', 14, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['card_bg']).pack(pady=10)
        
        fig3 = Figure(figsize=(10, 5), dpi=80, facecolor=self.colors['card_bg'])
        ax4 = fig3.add_subplot(111)
        
        # Simulated daily logins for last 7 days
        days = ['Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun']
        daily_logins = [max(1, stats['total_logins'] // 7 + i*2) for i in range(7)]
        
        ax4.plot(days, daily_logins, marker='o', color=self.colors['secondary'], 
                linewidth=2, markersize=8)
        ax4.fill_between(days, daily_logins, alpha=0.3, color=self.colors['secondary'])
        ax4.set_title('Daily Login Activity (Last 7 Days)', fontsize=12, fontweight='bold')
        ax4.set_xlabel('Day', fontsize=10)
        ax4.set_ylabel('Logins', fontsize=10)
        ax4.grid(True, alpha=0.3)
        
        canvas3 = FigureCanvasTkAgg(fig3, master=graph3_frame)
        canvas3.draw()
        canvas3.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
    
    def save_graphs(self):
        """Save graphs to files"""
        try:
            import os
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            folder = f"graphs_{timestamp}"
            os.makedirs(folder, exist_ok=True)
            
            # Get statistics
            stats = self.db.get_system_statistics()
            
            # Create and save bar chart
            fig1 = Figure(figsize=(10, 6), dpi=100)
            ax1 = fig1.add_subplot(111)
            
            categories = ['Total Users', 'Total Logins', 'Files Shared', 'Active Sessions']
            values = [stats['total_users'], stats['total_logins'], 
                     stats['files_shared'], stats['currently_logged_in']]
            
            bars = ax1.bar(categories, values, color=['blue', 'green', 'orange', 'red'])
            ax1.set_title('System Statistics')
            ax1.set_ylabel('Count')
            
            for bar, value in zip(bars, values):
                height = bar.get_height()
                ax1.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                        f'{value}', ha='center', va='bottom')
            
            fig1.savefig(os.path.join(folder, 'system_stats.png'))
            
            # Create and save pie chart
            fig2 = Figure(figsize=(8, 8), dpi=100)
            ax2 = fig2.add_subplot(111)
            
            user_types = ['Admins', 'Regular Users']
            user_counts = [stats['admin_count'], stats['user_count']]
            
            ax2.pie(user_counts, labels=user_types, autopct='%1.1f%%', startangle=90)
            ax2.set_title('User Distribution')
            
            fig2.savefig(os.path.join(folder, 'user_distribution.png'))
            
            messagebox.showinfo("Success", f"Graphs saved to folder:\n{folder}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save graphs: {str(e)}")
    
    def create_user_management_tab(self, notebook):
        """Create user management tab for admin"""
        user_tab = tk.Frame(notebook, bg=self.colors['background'])
        notebook.add(user_tab, text="👥 User Management")
        
        container = tk.Frame(user_tab, bg=self.colors['background'])
        container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        # Title
        tk.Label(container,
                text="User Management",
                font=('Segoe UI', 20, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['background']).pack(pady=(0, 20))
        
        # Card
        card = tk.Frame(container, bg=self.colors['card_bg'], relief='ridge', borderwidth=1)
        card.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Control buttons frame
        control_frame = tk.Frame(card, bg=self.colors['card_bg'])
        control_frame.pack(fill=tk.X, padx=20, pady=20)
        
        # Search frame
        search_frame = tk.Frame(control_frame, bg=self.colors['card_bg'])
        search_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        tk.Label(search_frame,
                text="🔍 Search:",
                font=('Segoe UI', 10, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['card_bg']).pack(side=tk.LEFT, padx=(0, 10))
        
        self.user_search_var = tk.StringVar()
        self.user_search_var.trace('w', self.filter_users)
        search_entry = ttk.Entry(search_frame,
                                textvariable=self.user_search_var,
                                font=('Segoe UI', 10),
                                width=30,
                                style='Custom.TEntry')
        search_entry.pack(side=tk.LEFT, padx=(0, 10))
        
        # Role filter
        tk.Label(search_frame,
                text="Role:",
                font=('Segoe UI', 10, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['card_bg']).pack(side=tk.LEFT, padx=(20, 10))
        
        self.role_filter_var = tk.StringVar(value="all")
        role_combo = ttk.Combobox(search_frame,
                                 textvariable=self.role_filter_var,
                                 values=["all", "admin", "user"],
                                 state="readonly",
                                 width=10)
        role_combo.pack(side=tk.LEFT, padx=(0, 10))
        role_combo.bind('<<ComboboxSelected>>', self.filter_users)
        
        # Action buttons
        action_frame = tk.Frame(control_frame, bg=self.colors['card_bg'])
        action_frame.pack(side=tk.RIGHT)
        
        ttk.Button(action_frame,
                  text="🔄 Refresh",
                  command=self.load_users,
                  style='Primary.TButton',
                  width=12).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(action_frame,
                  text="👁️ View Details",
                  command=self.view_user_details,
                  style='Primary.TButton',
                  width=20).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(action_frame,
                  text="❌ Delete User",
                  command=self.delete_user,
                  style='Danger.TButton',
                  width=20).pack(side=tk.LEFT, padx=5)
        
        # Treeview for users
        tree_frame = tk.Frame(card, bg=self.colors['card_bg'])
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 20))
        
        columns = ("Email", "Role", "Created At", "Last Login", "Status")
        self.users_tree = ttk.Treeview(tree_frame,
                                      columns=columns,
                                      show="headings",
                                      height=15,
                                      style='Custom.Treeview')
        
        column_widths = {"Email": 250, "Role": 80, "Created At": 150, "Last Login": 150, "Status": 100}
        for col in columns:
            self.users_tree.heading(col, text=col)
            self.users_tree.column(col, width=column_widths.get(col, 120), anchor='center')
        
        # Add scrollbar
        scrollbar = ttk.Scrollbar(tree_frame,
                                 orient="vertical",
                                 command=self.users_tree.yview)
        self.users_tree.configure(yscrollcommand=scrollbar.set)
        
        self.users_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Statistics frame
        stats_frame = tk.Frame(card, bg=self.colors['card_bg'])
        stats_frame.pack(fill=tk.X, padx=20, pady=(0, 20))
        
        self.user_stats_label = tk.Label(stats_frame,
                                        text="Total Users: 0 | Admins: 0 | Active: 0",
                                        font=('Segoe UI', 10, 'bold'),
                                        fg=self.colors['secondary'],
                                        bg=self.colors['card_bg'])
        self.user_stats_label.pack()
        
        # Load initial user data
        self.load_users()
    
    def load_users(self):
        """Load users into the treeview from database"""
        try:
            # Check if treeview exists
            if not hasattr(self, 'users_tree') or not self.users_tree.winfo_exists():
                return
            
            # Clear existing items
            for item in self.users_tree.get_children():
                self.users_tree.delete(item)
            
            all_users = self.db.get_all_users()
            active_sessions = self.db.get_active_sessions()
            
            # Get active users
            active_users = {session['user_email'] for session in active_sessions}
            
            admin_count = 0
            user_count = 0
            active_count = 0
            
            for user in all_users:
                email = user['email']
                role = user.get("role", "user").upper()
                created_at = user.get("created_at", "Unknown")
                last_login = user.get("last_login", "Never")
                status = "🟢 Active" if email in active_users else "⚪ Inactive"
                
                if role == "ADMIN":
                    admin_count += 1
                else:
                    user_count += 1
                
                if email in active_users:
                    active_count += 1
                
                self.users_tree.insert("", "end", values=(
                    email,
                    role,
                    created_at,
                    last_login,
                    status
                ))
            
            # Update statistics
            total_users = len(all_users)
            self.user_stats_label.config(
                text=f"Total Users: {total_users} | Admins: {admin_count} | Users: {user_count} | Active: {active_count}"
            )
            
        except Exception as e:
            print(f"Error loading users: {e}")
            # Don't show error message to avoid recursion
            return
    
    def filter_users(self, *args):
        """Filter users based on search criteria"""
        search_term = self.user_search_var.get().lower()
        role_filter = self.role_filter_var.get()
        
        try:
            all_users = self.db.get_all_users()
            active_sessions = self.db.get_active_sessions()
            
            # Get active users
            active_users = {session['user_email'] for session in active_sessions}
            
            # Clear existing items
            for item in self.users_tree.get_children():
                self.users_tree.delete(item)
            
            admin_count = 0
            user_count = 0
            active_count = 0
            
            for user in all_users:
                email = user['email']
                
                # Apply filters
                if search_term and search_term not in email.lower():
                    continue
                
                if role_filter != "all" and user.get("role") != role_filter:
                    continue
                
                role = user.get("role", "user").upper()
                created_at = user.get("created_at", "Unknown")
                last_login = user.get("last_login", "Never")
                status = "🟢 Active" if email in active_users else "⚪ Inactive"
                
                if role == "ADMIN":
                    admin_count += 1
                else:
                    user_count += 1
                
                if email in active_users:
                    active_count += 1
                
                self.users_tree.insert("", "end", values=(
                    email,
                    role,
                    created_at,
                    last_login,
                    status
                ))
            
            # Update statistics
            total_users = admin_count + user_count
            self.user_stats_label.config(
                text=f"Showing: {total_users} | Admins: {admin_count} | Users: {user_count} | Active: {active_count}"
            )
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to filter users: {str(e)}")
    
    def view_user_details(self):
        """View detailed information about selected user"""
        selected_item = self.users_tree.selection()
        if not selected_item:
            messagebox.showwarning("Warning", "Please select a user to view details")
            return
        
        email = self.users_tree.item(selected_item[0])["values"][0]
        
        try:
            user_data = self.db.get_user(email)
            
            if not user_data:
                messagebox.showerror("Error", "User not found")
                return
            
            # Create detail dialog
            detail_window = tk.Toplevel(self.root)
            detail_window.title(f"User Details: {email}")
            detail_window.geometry("600x500")
            detail_window.configure(bg=self.colors['background'])
            detail_window.transient(self.root)
            detail_window.grab_set()
            
            self.center_window(detail_window, 600, 500)
        
            # Header
            header_frame = tk.Frame(detail_window, bg=self.colors['primary'], height=80)
            header_frame.pack(fill=tk.X)
            header_frame.pack_propagate(False)
            
            tk.Label(header_frame,
                    text=f"👤 User Details",
                    font=('Segoe UI', 18, 'bold'),
                    fg=self.colors['text_light'],
                    bg=self.colors['primary']).pack(expand=True, pady=(10, 0))
            
            tk.Label(header_frame,
                    text=email,
                    font=('Segoe UI', 12),
                    fg=self.colors['highlight'],
                    bg=self.colors['primary']).pack(pady=(0, 10))
            
            # Content
            content_frame = tk.Frame(detail_window, bg=self.colors['card_bg'], padx=20, pady=20)
            content_frame.pack(fill=tk.BOTH, expand=True)
            
            # User information
            info_text = f"""
User Information:

📧 Email: {email}
👤 Role: {user_data.get('role', 'user').upper()}
📅 Created: {user_data.get('created_at', 'Unknown')}
🔐 Last Login: {user_data.get('last_login', 'Never')}
❓ Security Question: {user_data.get('security_question', 'Not set')}

Account Statistics:
───────────────"""
            
            # Add login statistics
            total_logins = self.db.get_user_login_count(email)
            info_text += f"\n📊 Total Logins: {total_logins}"
            
            # Failed attempts
            failed_attempts = self.db.get_failed_login_attempts(email)
            info_text += f"\n❌ Failed Attempts: {failed_attempts['attempts']}"
            
            # Add shared files info
            files_shared = len(self.db.get_shared_files_by_user(email))
            files_received = len(self.db.get_shared_files_for_user(email))
            
            info_text += f"\n\n📁 Files Shared: {files_shared}"
            info_text += f"\n📥 Files Received: {files_received}"
            
            # Check RSA keys
            public_key_info = self.get_public_key_from_repository(email)
            rsa_status = "✅ RSA Keys: Available" if public_key_info[0] else "❌ RSA Keys: Not in Repository"
            info_text += f"\n\n🔑 RSA Key Status: {rsa_status}"
            
            # Check if user is active
            active_session = self.db.get_active_session_by_user(email)
            session_status = "🟢 Currently Active" if active_session else "⚪ Not Active"
            info_text += f"\n\n🔐 Session Status: {session_status}"
            
            # JSON Status
            json_user = self.json_manager.get_user(email)
            json_status = "📄 JSON: Available" if json_user else "📄 JSON: Not in JSON file"
            info_text += f"\n\n📄 JSON Status: {json_status}"
            
            # Display info
            info_display = scrolledtext.ScrolledText(content_frame,
                                                    wrap=tk.WORD,
                                                    font=('Consolas', 10),
                                                    bg=self.colors['card_bg'],
                                                    fg=self.colors['text_dark'],
                                                    relief='flat',
                                                    borderwidth=1,
                                                    height=20)
            info_display.pack(fill=tk.BOTH, expand=True)
            info_display.insert(tk.END, info_text)
            info_display.config(state=tk.DISABLED)
            
            # Close button
            button_frame = tk.Frame(detail_window, bg=self.colors['background'], pady=20)
            button_frame.pack(fill=tk.X, padx=20)
            
            ttk.Button(button_frame,
                      text="Close",
                      command=detail_window.destroy,
                      style='Primary.TButton',
                      width=20).pack()
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load user details: {str(e)}")
    
    def delete_user(self):
        """Delete selected user with confirmation and cleanup"""
        selected_item = self.users_tree.selection()
        if not selected_item:
            messagebox.showwarning("Warning", "Please select a user to delete")
            return
        
        email = self.users_tree.item(selected_item[0])["values"][0]
        
        # Prevent deleting current admin
        if email == self.current_user:
            messagebox.showerror("Error", "Cannot delete your own account while logged in")
            return
        
        # Get user data for confirmation
        user_data = self.db.get_user(email)
        if not user_data:
            messagebox.showerror("Error", "User not found")
            return
        
        # Prevent deleting the last admin
        if user_data['role'] == 'admin':
            all_users = self.db.get_all_users()
            admin_count = sum(1 for u in all_users if u['role'] == 'admin')
            if admin_count <= 1:
                messagebox.showerror("Error", "Cannot delete the last admin user")
                return
        
        # Confirmation dialog
        confirm = messagebox.askyesno(
            "Confirm Deletion",
            f"Are you sure you want to delete user:\n\n{email}\n\n"
            f"This will permanently delete:\n"
            f"• User account\n"
            f"• User's shared files\n"
            f"• User's login history\n"
            f"• User's RSA keys\n"
            f"• User's public key from repository\n"
            f"• User from JSON file\n\n"
            f"This action cannot be undone!",
            icon='warning'
        )
        
        if not confirm:
            return
        
        try:
            # Delete user's RSA keys
            self.delete_user_rsa_keys(email)
            
            # Delete user's public key
            self.db.delete_public_key(email)
            
            # Delete user's shared files
            deleted_files = self.db.delete_user_shared_files(email)
            
            # Remove from active sessions
            self.db.remove_active_session(email)
            
            # Delete from JSON
            self.json_manager.delete_user(email)
            
            # Soft delete user (mark as inactive)
            success = self.db.delete_user(email)
            
            if success:
                # Log the action
                self.log_action(f"DELETE_USER_{email}", f"Role: {user_data['role']}")
                
                # Show success message
                messagebox.showinfo(
                    "Success",
                    f"✅ User '{email}' has been deleted successfully.\n\n"
                    f"• Account removed\n"
                    f"• {deleted_files} shared files cleaned up\n"
                    f"• Login history cleared\n"
                    f"• RSA keys removed\n"
                    f"• Public key removed from repository\n"
                    f"• Removed from JSON file"
                )
                
                # Refresh user list
                self.load_users()
                
                # Refresh public key directory if it exists
                if hasattr(self, 'public_keys_tree'):
                    self.refresh_public_key_directory()
                
            else:
                messagebox.showerror("Error", "Failed to delete user from database")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to delete user: {str(e)}")
    
    def delete_user_rsa_keys(self, email):
        """Delete user's RSA keys from local storage"""
        try:
            public_key_path = os.path.join(self.user_rsa_keys_dir, f"{email}_public.pem")
            private_key_path = os.path.join(self.user_rsa_keys_dir, f"{email}_private.enc")
            
            if os.path.exists(public_key_path):
                os.remove(public_key_path)
            
            if os.path.exists(private_key_path):
                os.remove(private_key_path)
                
        except Exception as e:
            print(f"Error deleting RSA keys for {email}: {e}")
    
    def create_sessions_tab(self, notebook):
        """Create sessions tab"""
        sessions_tab = tk.Frame(notebook, bg=self.colors['background'])
        notebook.add(sessions_tab, text="👥 Sessions")
        
        container = tk.Frame(sessions_tab, bg=self.colors['background'])
        container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        # Two column layout
        left_frame = tk.Frame(container, bg=self.colors['background'])
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        right_frame = tk.Frame(container, bg=self.colors['background'])
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(10, 0))
        
        # Active sessions
        tk.Label(left_frame,
                text="🟢 Active Sessions",
                font=('Segoe UI', 16, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['background']).pack(anchor='w', pady=(0, 10))
        
        active_card = tk.Frame(left_frame, bg=self.colors['card_bg'], relief='ridge', borderwidth=1)
        active_card.pack(fill=tk.BOTH, expand=True)
        
        columns = ("User", "Role", "Login Time", "Last Activity")
        self.active_users_tree = ttk.Treeview(active_card,
                                             columns=columns,
                                             show="headings",
                                             height=12,
                                             style='Custom.Treeview')
        
        for col in columns:
            self.active_users_tree.heading(col, text=col)
            width = 120 if col in ["User", "Role"] else 150
            self.active_users_tree.column(col, width=width, anchor='center')
        
        scrollbar = ttk.Scrollbar(active_card,
                                 orient="vertical",
                                 command=self.active_users_tree.yview)
        self.active_users_tree.configure(yscrollcommand=scrollbar.set)
        
        self.active_users_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y, pady=10)
        
        # Login history
        tk.Label(right_frame,
                text="📋 Recent Login History",
                font=('Segoe UI', 16, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['background']).pack(anchor='w', pady=(0, 10))
        
        history_card = tk.Frame(right_frame, bg=self.colors['card_bg'], relief='ridge', borderwidth=1)
        history_card.pack(fill=tk.BOTH, expand=True)
        
        columns = ("User", "Role", "Login Time")
        self.login_history_tree = ttk.Treeview(history_card,
                                              columns=columns,
                                              show="headings",
                                              height=12,
                                              style='Custom.Treeview')
        
        for col in columns:
            self.login_history_tree.heading(col, text=col)
            self.login_history_tree.column(col, width=150, anchor='center')
        
        scrollbar2 = ttk.Scrollbar(history_card,
                                  orient="vertical",
                                  command=self.login_history_tree.yview)
        self.login_history_tree.configure(yscrollcommand=scrollbar2.set)
        
        self.login_history_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
        scrollbar2.pack(side=tk.RIGHT, fill=tk.Y, pady=10)
        
        # Action buttons for sessions
        action_frame = tk.Frame(container, bg=self.colors['background'])
        action_frame.pack(side=tk.BOTTOM, pady=20)
        
        ttk.Button(action_frame,
                  text="🔄 Refresh Sessions",
                  command=self.refresh_admin_stats,
                  style='Primary.TButton',
                  width=15).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(action_frame,
                  text="🚫 Force Logout",
                  command=self.force_logout_user,
                  style='Danger.TButton',
                  width=15).pack(side=tk.LEFT, padx=5)
        
        # Load initial data
        self.refresh_admin_stats()
    
    def force_logout_user(self):
        """Force logout selected user from sessions"""
        selected_item = self.active_users_tree.selection()
        if not selected_item:
            messagebox.showwarning("Warning", "Please select a user to force logout")
            return
        
        user_data = self.active_users_tree.item(selected_item[0])["values"]
        username = user_data[0]
        
        # Prevent force logging out yourself
        if username == self.current_user:
            messagebox.showerror("Error", "Cannot force logout yourself")
            return
        
        confirm = messagebox.askyesno(
            "Confirm Force Logout",
            f"Are you sure you want to force logout user:\n\n{username}\n\n"
            f"This will immediately terminate their session.",
            icon='warning'
        )
        
        if not confirm:
            return
        
        try:
            success = self.db.remove_active_session(username)
            
            if success:
                self.log_action(f"FORCE_LOGOUT_{username}", "")
                messagebox.showinfo("Success", f"✅ User '{username}' has been force logged out.")
                self.refresh_admin_stats()
            else:
                messagebox.showerror("Error", f"Failed to force logout user '{username}'")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to force logout: {str(e)}")
    
    def create_logs_tab(self, notebook):
        """Create logs tab"""
        logs_tab = tk.Frame(notebook, bg=self.colors['background'])
        notebook.add(logs_tab, text="📜 Activity Logs")
        
        container = tk.Frame(logs_tab, bg=self.colors['background'])
        container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        # Title
        tk.Label(container,
                text="Activity Logs",
                font=('Segoe UI', 20, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['background']).pack(pady=(0, 20))
        
        # Card
        card = tk.Frame(container, bg=self.colors['card_bg'], relief='ridge', borderwidth=1)
        card.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Control buttons
        control_frame = tk.Frame(card, bg=self.colors['card_bg'])
        control_frame.pack(fill=tk.X, padx=20, pady=20)
        
        ttk.Button(control_frame,
                  text="🔄 Refresh Logs",
                  command=self.refresh_logs,
                  style='Primary.TButton',
                  width=15).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(control_frame,
                  text="📊 Export Logs",
                  command=self.export_logs,
                  style='Success.TButton',
                  width=15).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(control_frame,
                  text="🗑️ Clear Logs",
                  command=self.clear_logs,
                  style='Danger.TButton',
                  width=15).pack(side=tk.LEFT, padx=5)
        
        # Logs display
        logs_frame = tk.Frame(card, bg=self.colors['card_bg'])
        logs_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 20))
        
        self.logs_text = scrolledtext.ScrolledText(logs_frame,
                                                  height=20,
                                                  font=('Consolas', 9),
                                                  bg='#1E1E1E',
                                                  fg='#00FF00',
                                                  relief='solid',
                                                  borderwidth=1)
        self.logs_text.pack(fill=tk.BOTH, expand=True)
        self.logs_text.config(state=tk.DISABLED)
        
        self.refresh_logs()
    
    def create_database_management_tab(self, notebook):
        """Create database management tab with JSON support"""
        db_tab = tk.Frame(notebook, bg=self.colors['background'])
        notebook.add(db_tab, text="💾 Database")
        
        container = tk.Frame(db_tab, bg=self.colors['background'])
        container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        # Title
        tk.Label(container,
                text="Database Management",
                font=('Segoe UI', 20, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['background']).pack(pady=(0, 20))
        
        # Card
        card = tk.Frame(container, bg=self.colors['card_bg'], relief='ridge', borderwidth=1)
        card.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Database info
        info_frame = tk.Frame(card, bg=self.colors['card_bg'])
        info_frame.pack(fill=tk.X, padx=30, pady=30)
        
        tk.Label(info_frame,
                text="💾 SQLite Database Information",
                font=('Segoe UI', 14, 'bold'),
                fg=self.colors['secondary'],
                bg=self.colors['card_bg']).pack(anchor='w', pady=(0, 10))
        
        try:
            import sqlite3
            conn = sqlite3.connect(self.db.db_path)
            cursor = conn.cursor()
            
            # Get database info
            cursor.execute("PRAGMA page_count;")
            page_count = cursor.fetchone()[0]
            cursor.execute("PRAGMA page_size;")
            page_size = cursor.fetchone()[0]
            db_size_mb = (page_count * page_size) / (1024 * 1024)
            
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
            tables = cursor.fetchall()
            table_count = len(tables)
            
            cursor.execute("PRAGMA foreign_keys;")
            foreign_keys = cursor.fetchone()[0]
            
            conn.close()
            
            info_text = f"""
Database Path: {os.path.abspath(self.db.db_path)}
Database Size: {db_size_mb:.2f} MB
Number of Tables: {table_count}
Foreign Keys: {'Enabled' if foreign_keys else 'Disabled'}
SQLite Version: {sqlite3.sqlite_version}"""
            
            tk.Label(info_frame,
                    text=info_text,
                    font=('Consolas', 10),
                    fg=self.colors['text_dark'],
                    bg=self.colors['card_bg'],
                    justify=tk.LEFT).pack(anchor='w')
            
        except Exception as e:
            tk.Label(info_frame,
                    text=f"Error getting database info: {str(e)}",
                    font=('Consolas', 10),
                    fg=self.colors['accent'],
                    bg=self.colors['card_bg']).pack(anchor='w')
        
        # Database actions
        actions_frame = tk.Frame(card, bg=self.colors['card_bg'])
        actions_frame.pack(fill=tk.X, padx=30, pady=30)
        
        tk.Label(actions_frame,
                text="⚙️ Database Actions",
                font=('Segoe UI', 10, 'bold'),
                fg=self.colors['secondary'],
                bg=self.colors['card_bg']).pack(anchor='w', pady=(0, 10))
        
        # Buttons grid
        buttons_grid = tk.Frame(actions_frame, bg=self.colors['card_bg'])
        buttons_grid.pack()
        
        action_buttons = [
            ("🔄 Optimize DB", self.optimize_database, "Primary.TButton"),
            ("📤 Backup DB", self.backup_database, "Success.TButton"),
            ("🔍 Check Integrity", self.check_database_integrity, "Warning.TButton"),
            ("🧹 Vacuum DB", self.vacuum_database, "Danger.TButton")
        ]
        
        for i, (text, command, style_name) in enumerate(action_buttons):
            btn = ttk.Button(buttons_grid,
                           text=text,
                           command=command,
                           style=style_name,
                           width=20)
            btn.grid(row=i//2, column=i%2, padx=5, pady=5)
        
        # Add JSON operations section
        json_frame = tk.Frame(card, bg=self.colors['card_bg'])
        json_frame.pack(fill=tk.X, padx=30, pady=30)
        
        tk.Label(json_frame,
                text="📄 JSON File Operations",
                font=('Segoe UI', 10, 'bold'),
                fg=self.colors['secondary'],
                bg=self.colors['card_bg']).pack(anchor='w', pady=(0, 10))
        
        # JSON buttons grid - IMPROVED VERSION
        json_grid = tk.Frame(json_frame, bg=self.colors['card_bg'])
        json_grid.pack()
        
        # Define JSON operations with clearer names
        json_operations = [
            ("🔄 Sync Database to JSON", self.sync_database_to_json, "JSONPrimary.TButton"),
            ("📥 Sync JSON to Database", self.sync_json_to_database, "JSONSuccess.TButton"),
            ("📤 Export Users to JSON", self.export_users_json, "JSONWarning.TButton"),
            ("📥 Import Users from JSON", self.import_users_json, "JSONDanger.TButton")
        ]
        
        # Create a frame for descriptions
        desc_frame = tk.Frame(json_grid, bg=self.colors['card_bg'])
        desc_frame.grid(row=0, column=0, columnspan=2, pady=(0, 10))
        
        tk.Label(desc_frame,
                text="📊 JSON Operations - User Data Management",
                font=('Segoe UI', 10, 'bold'),
                fg=self.colors['primary'],
                bg=self.colors['card_bg']).pack()
        
        # Create buttons in a 2x2 grid
        for i, (text, command, style_name) in enumerate(json_operations):
            row = (i // 2) + 1  # Start from row 1
            col = i % 2
            
            # Create button with improved styling
            btn = ttk.Button(json_grid,
                           text=text,
                           command=command,
                           style=style_name,
                           width=25)  # Wider for better text display
            
            btn.grid(row=row, column=col, padx=10, pady=10, sticky='ew')
            
            # Add hover effects
            btn.bind("<Enter>", lambda e, b=btn: b.config(cursor="hand2"))
            btn.bind("<Leave>", lambda e, b=btn: b.config(cursor=""))
        
        # Add descriptive labels for each button
        descriptions = [
            "Copy all users from database to JSON file",
            "Import users from JSON file to database",
            "Export current users to a new JSON file",
            "Import users from an external JSON file"
        ]
        
        for i, desc in enumerate(descriptions):
            row = (i // 2) + 2  # Start from row 2
            col = i % 2
            
            tk.Label(json_grid,
                    text=desc,
                    font=('Segoe UI', 9, 'italic'),
                    fg=self.colors['text_dark'],
                    bg=self.colors['card_bg'],
                    wraplength=200).grid(row=row, column=col, padx=10, pady=(0, 15), sticky='w')
        
        # Add JSON file info with better visibility
        try:
            json_stats = self.get_json_stats()
            info_text = f"""
📄 JSON File Information:
────────────────────────────────
File: users.json
Total Users: {json_stats['total_users']}
Last Updated: {json_stats['last_updated']}
File Size: {os.path.getsize('users.json')/1024:.1f} KB"""
            
            info_frame = tk.Frame(json_frame, bg='#F0F8FF', relief='ridge', borderwidth=2)
            info_frame.pack(fill=tk.X, padx=30, pady=(20, 10))
            
            info_label = tk.Label(info_frame,
                                text=info_text,
                                font=('Consolas', 10),
                                fg=self.colors['primary'],
                                bg='#F0F8FF',
                                justify=tk.LEFT)
            info_label.pack(padx=10, pady=10)
            
        except Exception as e:
            error_frame = tk.Frame(json_frame, bg='#FFE6E6', relief='ridge', borderwidth=2)
            error_frame.pack(fill=tk.X, padx=30, pady=(20, 10))
            
            error_label = tk.Label(error_frame,
                                  text=f"⚠️ JSON file not found or error: {str(e)}",
                                  font=('Segoe UI', 10),
                                  fg=self.colors['accent'],
                                  bg='#FFE6E6')
            error_label.pack(padx=10, pady=10)
    
    def optimize_database(self):
        """Optimize the database"""
        try:
            self.db.vacuum()
            messagebox.showinfo("Success", "✅ Database optimized successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to optimize database: {str(e)}")
    
    def backup_database(self):
        """Create a backup of the database"""
        filename = filedialog.asksaveasfilename(
            defaultextension=".db",
            filetypes=[("SQLite Database", "*.db"), ("All files", "*.*")],
            initialfile=f"secure_file_sharing_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.db"
        )
        
        if filename:
            try:
                if self.db.backup_database(filename):
                    messagebox.showinfo("Success", f"✅ Database backup created:\n{filename}")
                else:
                    messagebox.showerror("Error", "Failed to create database backup")
            except Exception as e:
                messagebox.showerror("Error", f"Backup failed: {str(e)}")
    
    def check_database_integrity(self):
        """Check database integrity"""
        try:
            with closing(self.db.get_connection()) as conn:
                cursor = conn.cursor()
                cursor.execute("PRAGMA integrity_check;")
                result = cursor.fetchone()
                
                if result and result[0] == "ok":
                    messagebox.showinfo("Integrity Check", "✅ Database integrity check passed!")
                else:
                    messagebox.showerror("Integrity Check", f"❌ Database integrity issues:\n{result}")
        except Exception as e:
            messagebox.showerror("Error", f"Integrity check failed: {str(e)}")
    
    def vacuum_database(self):
        """Vacuum the database"""
        if messagebox.askyesno("Confirm", "This will reorganize the database file to reclaim unused space.\nContinue?"):
            try:
                self.db.vacuum()
                messagebox.showinfo("Success", "✅ Database vacuum completed!")
            except Exception as e:
                messagebox.showerror("Error", f"Vacuum failed: {str(e)}")
    
    def start_session_cleanup_thread(self):
        """Start a background thread for cleaning up inactive sessions"""
        def cleanup_task():
            while True:
                try:
                    self.cleanup_inactive_sessions()
                    time.sleep(60)
                except Exception as e:
                    print(f"Error in session cleanup thread: {e}")
        
        cleanup_thread = threading.Thread(target=cleanup_task, daemon=True)
        cleanup_thread.start()
    
    def get_system_statistics(self):
        """Get system statistics for admin dashboard"""
        return self.db.get_system_statistics()
    
    def refresh_admin_stats(self):
        """Refresh admin dashboard statistics"""
        self.update_session_activity()
        
        try:
            # Refresh active sessions
            for item in self.active_users_tree.get_children():
                self.active_users_tree.delete(item)
            
            active_sessions = self.db.get_active_sessions()
            
            for session in active_sessions:
                self.active_users_tree.insert("", "end", values=(
                    session.get('user_email', 'Unknown'),
                    session.get('user_role', 'Unknown'),
                    session.get('login_time', 'Unknown'),
                    session.get('last_activity', 'Unknown')
                ))
                
        except Exception as e:
            print(f"Error refreshing active users: {e}")
        
        try:
            # Refresh login history
            for item in self.login_history_tree.get_children():
                self.login_history_tree.delete(item)
            
            recent_logins = self.db.get_recent_logins(limit=10)
            for login in reversed(recent_logins):
                self.login_history_tree.insert("", "end", values=(
                    login.get('user_email', 'Unknown'),
                    login.get('user_role', 'Unknown'),
                    login.get('login_time', 'Unknown')
                ))
                
        except Exception as e:
            print(f"Error refreshing admin stats: {e}")
        
        # Refresh user management tab if it exists
        if hasattr(self, 'users_tree'):
            self.load_users()
        
        # Refresh public key directory if it exists
        if hasattr(self, 'public_keys_tree'):
            self.refresh_public_key_directory()
    
    def export_logs(self):
        """Export logs to a file"""
        try:
            filename = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("Text files", "*.txt"), ("Log files", "*.log"), ("All files", "*.*")]
            )
            
            if filename:
                with open('file_sharing.log', 'r') as source:
                    logs = source.read()
                
                with open(filename, 'w') as target:
                    target.write(f"Secure File Sharing Tool - Activity Logs\n")
                    target.write(f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    target.write(f"Exported by: {self.current_user}\n")
                    target.write(f"Database: {self.db.db_path}\n")
                    target.write(f"JSON File: users.json\n")
                    target.write("="*80 + "\n\n")
                    target.write(logs)
                
                messagebox.showinfo("Success", f"Logs exported successfully to:\n{filename}")
                
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export logs: {str(e)}")
    
    def clear_logs(self):
        """Clear activity logs"""
        if messagebox.askyesno("Clear Logs", "Are you sure you want to clear all logs?"):
            try:
                with open('file_sharing.log', 'w') as f:
                    f.write("")
                self.refresh_logs()
                messagebox.showinfo("Success", "Logs cleared successfully!")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to clear logs: {str(e)}")
    
    def logout(self):
        """Handle user logout"""
        self.log_action("LOGOUT", "")
        self.remove_active_session()
        
        # Cancel the time updater if it's running
        if hasattr(self, 'time_updater_id'):
            try:
                self.root.after_cancel(self.time_updater_id)
            except:
                pass
        
        # Clear current interface
        for widget in self.root.winfo_children():
            widget.destroy()
        
        # Recreate authentication interface
        self.create_unified_auth_frame()
        
        # Reset user variables
        self.current_user = None
        self.user_role = None
        self.aes_key = None
        self.user_private_key = None
        self.user_public_key = None
        
        messagebox.showinfo("Logged Out", "You have been successfully logged out.")
    
    def browse_file(self, path_var=None):
        """Browse and select file"""
        if path_var is None:
            path_var = self.file_path_var
            
        filename = filedialog.askopenfilename()
        if filename:
            path_var.set(filename)
    
    def generate_aes_key(self, entry_widget=None):
        """Generate a new AES key"""
        if entry_widget is None:
            entry_widget = self.aes_key_entry
        
        aes_key = Fernet.generate_key()
        entry_widget.delete(0, tk.END)
        entry_widget.insert(0, aes_key.decode())
        self.update_status("🎲 New AES key generated successfully")
        return aes_key
    
    def encrypt_file(self):
        """Encrypt selected file using AES"""
        file_path = self.file_path_var.get()
        if not file_path or not os.path.exists(file_path):
            messagebox.showerror("Error", "❌ Please select a valid file")
            return
        
        aes_key_input = self.aes_key_entry.get().strip()
        if aes_key_input:
            try:
                self.aes_key = aes_key_input.encode()
                Fernet(self.aes_key)
            except Exception as e:
                messagebox.showerror("Error", f"❌ Invalid AES key: {str(e)}")
                return
        else:
            self.generate_aes_key()
            aes_key_input = self.aes_key_entry.get().strip()
            self.aes_key = aes_key_input.encode()
        
        try:
            with open(file_path, 'rb') as file:
                file_data = file.read()
            
            fernet = Fernet(self.aes_key)
            encrypted_data = fernet.encrypt(file_data)
            
            encrypted_file_path = file_path + '.encrypted'
            with open(encrypted_file_path, 'wb') as file:
                file.write(encrypted_data)
            
            self.update_status(f"✅ File encrypted successfully!\n📁 Encrypted file: {encrypted_file_path}")
            self.log_action("ENCRYPT", file_path)
            
        except Exception as e:
            messagebox.showerror("Error", f"❌ Encryption failed: {str(e)}")
    
    def decrypt_file(self):
        """Decrypt selected file using AES"""
        file_path = self.file_path_var.get()
        if not file_path or not os.path.exists(file_path):
            messagebox.showerror("Error", "❌ Please select a valid file")
            return
        
        aes_key_input = self.aes_key_entry.get().strip()
        if not aes_key_input:
            messagebox.showerror("Error", "❌ Please provide AES key for decryption")
            return
        
        try:
            self.aes_key = aes_key_input.encode()
            fernet = Fernet(self.aes_key)
            
            with open(file_path, 'rb') as file:
                encrypted_data = file.read()
            
            decrypted_data = fernet.decrypt(encrypted_data)
            
            if file_path.endswith('.encrypted'):
                decrypted_file_path = file_path.replace('.encrypted', '.decrypted')
            else:
                decrypted_file_path = file_path + '.decrypted'
            
            with open(decrypted_file_path, 'wb') as file:
                file.write(decrypted_data)
            
            success_msg = f"✅ File decrypted successfully!\n📁 Decrypted file: {decrypted_file_path}"
            self.update_status(success_msg)
            self.log_action("DECRYPT", file_path)
            
            messagebox.showinfo("Success", success_msg)
            
        except InvalidToken:
            messagebox.showerror("Error", "❌ Invalid AES key. The key may be incorrect or the file is corrupted.")
        except Exception as e:
            messagebox.showerror("Error", f"❌ Decryption failed: {str(e)}")
    
    def share_file(self):
        """Share a file with another user using RSA-encrypted AES key"""
        file_path = self.share_file_path_var.get()
        recipient_email = self.recipient_email_var.get().strip().lower()
        aes_key_input = self.share_aes_key_entry.get().strip()
        
        if not file_path or not os.path.exists(file_path):
            messagebox.showerror("Error", "❌ Please select a valid file to share")
            return
        
        if not recipient_email:
            messagebox.showerror("Error", "❌ Please enter recipient email")
            return
        
        if not self.validate_email(recipient_email):
            messagebox.showerror("Error", "❌ Please enter a valid recipient email address")
            return
        
        # Check if recipient exists
        recipient_data = self.db.get_user(recipient_email)
        if not recipient_data:
            messagebox.showerror("Error", "❌ Recipient email not found in system")
            return
        
        # Check if recipient has RSA public key
        recipient_public_key, message = self.get_recipient_public_key(recipient_email)
        if not recipient_public_key:
            messagebox.showerror("Error", f"❌ {message}\n\nPlease ask the recipient to upload their public key to the repository.")
            return
        
        # Get or generate AES key
        if aes_key_input:
            try:
                aes_key = aes_key_input.encode()
                Fernet(aes_key)
            except:
                messagebox.showerror("Error", "❌ Invalid AES key")
                return
        else:
            aes_key = Fernet.generate_key()
            self.share_aes_key_entry.delete(0, tk.END)
            self.share_aes_key_entry.insert(0, aes_key.decode())
        
        try:
            # Encrypt file with AES
            with open(file_path, 'rb') as file:
                file_data = file.read()
            
            fernet = Fernet(aes_key)
            encrypted_data = fernet.encrypt(file_data)
            
            # Create unique filename
            original_filename = os.path.basename(file_path)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            shared_filename = f"{self.current_user}_{recipient_email}_{timestamp}_{original_filename}.encrypted"
            shared_file_path = os.path.join(self.shared_files_dir, shared_filename)
            
            # Save encrypted file to shared directory
            with open(shared_file_path, 'wb') as file:
                file.write(encrypted_data)
            
            # Encrypt AES key with recipient's RSA public key
            encrypted_aes_key_b64, encrypt_msg = self.encrypt_aes_key_with_rsa(aes_key, recipient_public_key)
            if not encrypted_aes_key_b64:
                messagebox.showerror("Error", f"❌ Failed to encrypt AES key: {encrypt_msg}")
                return
            
            # Store in database
            file_id = f"{self.current_user}_{recipient_email}_{timestamp}"
            success = self.db.add_shared_file(
                file_id=file_id,
                filename=original_filename,
                shared_by=self.current_user,
                shared_with=recipient_email,
                encrypted_file=shared_filename,
                encrypted_aes_key=encrypted_aes_key_b64,
                aes_key_plain=aes_key.decode()  # Store plain for testing
            )
            
            if success:
                # Show success with details
                success_msg = f"✅ File shared successfully with {recipient_email}!\n\n"
                success_msg += f"📁 Original file: {original_filename}\n"
                success_msg += f"🔑 AES Key (plain): {aes_key.decode()}\n"
                success_msg += f"🔐 AES Key (RSA-encrypted): {encrypted_aes_key_b64[:50]}...\n"
                success_msg += f"📦 Encrypted file saved as: {shared_filename}\n\n"
                success_msg += f"The recipient can now download and decrypt this file using the 'Get Decryption Key' button."
                
                messagebox.showinfo("Success", success_msg)
                self.update_status(f"🤝 File '{original_filename}' shared with {recipient_email}")
                self.log_action("SHARE", file_path)
                self.refresh_shared_files()
            else:
                messagebox.showerror("Error", "❌ Failed to save shared file record to database")
            
        except Exception as e:
            messagebox.showerror("Error", f"❌ File sharing failed: {str(e)}")
    
    def refresh_shared_files(self):
        """Refresh the list of shared files from database"""
        try:
            for item in self.shared_files_tree.get_children():
                self.shared_files_tree.delete(item)
            
            if self.current_user:
                shared_files = self.db.get_shared_files_for_user(self.current_user)
                
                for file_info in shared_files:
                    status = "✅ Downloaded" if file_info.get('downloaded') else "📥 Available"
                    self.shared_files_tree.insert("", "end", values=(
                        file_info['filename'],
                        file_info['shared_by'],
                        file_info['shared_date'],
                        status
                    ), iid=file_info['file_id'])
                    
        except Exception as e:
            messagebox.showerror("Error", f"❌ Failed to refresh shared files: {str(e)}")
    
    def download_shared_file(self):
        """Download and decrypt a shared file"""
        selected_item = self.shared_files_tree.selection()
        if not selected_item:
            messagebox.showerror("Error", "❌ Please select a file to download")
            return
        
        file_id = selected_item[0]
        aes_key_input = self.download_aes_key_entry.get().strip()
        
        # Debug: Show what we're working with
        print(f"DEBUG: File ID: {file_id}")
        print(f"DEBUG: AES Key provided (first 20 chars): {aes_key_input[:20] if aes_key_input else 'None'}")
        
        # If no AES key provided, try to decrypt from database
        if not aes_key_input:
            print("DEBUG: No AES key in field, trying to get from database...")
            success = self.get_aes_key_from_database(file_id)
            if not success:
                messagebox.showerror("Error", "❌ No AES key provided. Please click 'Get Decryption Key' first.")
                return
            aes_key_input = self.download_aes_key_entry.get().strip()
            print(f"DEBUG: Got AES key from database: {aes_key_input[:20]}...")
        
        try:
            # Get file info from database
            shared_files = self.db.get_shared_files_for_user(self.current_user)
            file_info = next((f for f in shared_files if f['file_id'] == file_id), None)
            
            if not file_info:
                messagebox.showerror("Error", "❌ Selected file not found in database")
                return
            
            print(f"DEBUG: File info found: {file_info['filename']}")
            
            # Get the encrypted file path
            encrypted_file_path = os.path.join(self.shared_files_dir, file_info["encrypted_file"])
            
            if not os.path.exists(encrypted_file_path):
                messagebox.showerror("Error", f"❌ Encrypted file not found: {encrypted_file_path}")
                return
            
            print(f"DEBUG: Encrypted file path: {encrypted_file_path}")
            print(f"DEBUG: Encrypted file size: {os.path.getsize(encrypted_file_path)} bytes")
            
            # Read encrypted file
            with open(encrypted_file_path, 'rb') as file:
                encrypted_data = file.read()
            
            print(f"DEBUG: Read {len(encrypted_data)} bytes of encrypted data")
            
            # Try to decrypt
            try:
                aes_key = aes_key_input.encode()
                print(f"DEBUG: Using AES key: {aes_key[:10]}... (length: {len(aes_key)})")
                
                fernet = Fernet(aes_key)
                print("DEBUG: Fernet object created")
                
                decrypted_data = fernet.decrypt(encrypted_data)
                print(f"DEBUG: Decryption successful! Decrypted {len(decrypted_data)} bytes")
                
                # Check if decrypted data looks reasonable
                if len(decrypted_data) == 0:
                    print("DEBUG: WARNING: Decrypted data is 0 bytes!")
                
                # Ask user where to save
                original_filename = file_info["filename"]
                # Remove any .encrypted extension if present
                if original_filename.endswith('.encrypted'):
                    suggested_name = original_filename.replace('.encrypted', '')
                else:
                    suggested_name = original_filename
                
                download_path = filedialog.asksaveasfilename(
                    initialfile=suggested_name,
                    title="Save Decrypted File",
                    defaultextension=""
                )
                
                if download_path:
                    # Ensure we're not overwriting the encrypted file
                    if download_path == encrypted_file_path:
                        messagebox.showerror("Error", "❌ Cannot overwrite encrypted file. Choose a different location.")
                        return
                    
                    # Write decrypted data
                    with open(download_path, 'wb') as file:
                        file.write(decrypted_data)
                    
                    print(f"DEBUG: Decrypted file saved to: {download_path}")
                    print(f"DEBUG: Saved file size: {os.path.getsize(download_path)} bytes")
                    
                    # Update shared files database
                    self.db.update_shared_file_downloaded(file_id)
                    
                    # Show success with more details
                    success_msg = f"✅ File successfully decrypted!\n\n"
                    success_msg += f"📁 Original: {file_info['filename']}\n"
                    success_msg += f"🔑 Key used: {aes_key_input[:30]}...\n"
                    success_msg += f"💾 Saved to: {download_path}\n"
                    success_msg += f"📊 Decrypted size: {len(decrypted_data)} bytes"
                    
                    messagebox.showinfo("Success", success_msg)
                    self.update_status(f"⬇️ Downloaded & Decrypted: {file_info['filename']}")
                    self.log_action("DOWNLOAD_DECRYPT", file_info["filename"])
                    self.refresh_shared_files()
                    
                    # Try to open the file automatically
                    try:
                        os.startfile(download_path)
                    except:
                        pass
                    
            except InvalidToken as e:
                print(f"DEBUG: InvalidToken error: {e}")
                messagebox.showerror("Error", 
                    f"❌ Invalid AES key. Cannot decrypt the file.\n\n"
                    f"Possible reasons:\n"
                    f"1. Wrong AES key\n"
                    f"2. File corrupted\n"
                    f"3. Key was not used to encrypt this file\n\n"
                    f"Try getting the key again with 'Get Decryption Key' button.")
                return
            except Exception as e:
                print(f"DEBUG: Decryption error: {type(e).__name__}: {e}")
                messagebox.showerror("Error", f"❌ Decryption failed: {type(e).__name__}: {str(e)}")
                return
                
        except Exception as e:
            print(f"DEBUG: General error: {type(e).__name__}: {e}")
            messagebox.showerror("Error", f"❌ Download failed: {str(e)}")
    
    def get_aes_key_from_database(self, file_id):
        """Get AES key from database for the selected file"""
        try:
            print(f"DEBUG: Trying to get AES key for file_id: {file_id}")
            
            shared_files = self.db.get_shared_files_for_user(self.current_user)
            file_info = next((f for f in shared_files if f['file_id'] == file_id), None)
            
            if not file_info:
                print(f"DEBUG: No file info found for file_id: {file_id}")
                return False
            
            print(f"DEBUG: Found file info: {file_info['filename']}")
            
            # Try to get plain AES key first
            if "aes_key_plain" in file_info and file_info["aes_key_plain"]:
                aes_key = file_info["aes_key_plain"]
                print(f"DEBUG: Found plain AES key in database: {aes_key[:30]}...")
                self.download_aes_key_entry.delete(0, tk.END)
                self.download_aes_key_entry.insert(0, aes_key)
                self.update_status(f"🔑 AES key retrieved from database for: {file_info['filename']}")
                return True
            
            # If no plain key, try RSA decryption
            elif "encrypted_aes_key" in file_info and file_info["encrypted_aes_key"]:
                encrypted_aes_key_b64 = file_info["encrypted_aes_key"]
                print(f"DEBUG: Found encrypted AES key (first 50 chars): {encrypted_aes_key_b64[:50]}...")
                
                # Try to decrypt with user's RSA private key
                if self.user_private_key:
                    print("DEBUG: User private key available, attempting RSA decryption...")
                    decrypted_aes_key, message = self.decrypt_aes_key_with_rsa(encrypted_aes_key_b64)
                    if decrypted_aes_key:
                        print(f"DEBUG: RSA decryption successful! Decrypted key: {decrypted_aes_key[:30]}...")
                        self.download_aes_key_entry.delete(0, tk.END)
                        self.download_aes_key_entry.insert(0, decrypted_aes_key.decode())
                        self.update_status(f"🔑 AES key decrypted with RSA for: {file_info['filename']}")
                        return True
                    else:
                        print(f"DEBUG: RSA decryption failed: {message}")
                        messagebox.showerror("Error", f"❌ Failed to decrypt AES key: {message}")
                        return False
                else:
                    print("DEBUG: User private key not loaded")
                    messagebox.showerror("Error", "❌ RSA private key not loaded. Please login again.")
                    return False
            else:
                print("DEBUG: No AES key found in database (neither plain nor encrypted)")
                messagebox.showerror("Error", "❌ No AES key found in database for this file")
                return False
            
        except Exception as e:
            print(f"DEBUG: Error in get_aes_key_from_database: {type(e).__name__}: {e}")
            messagebox.showerror("Error", f"❌ Failed to get AES key: {str(e)}")
            return False
    
    def decrypt_shared_file_key(self):
        """Get the decryption key for a shared file"""
        selected_item = self.shared_files_tree.selection()
        if not selected_item:
            messagebox.showerror("Error", "❌ Please select a shared file")
            return False
        
        file_id = selected_item[0]
        
        try:
            print(f"DEBUG: Decrypting shared file key for file_id: {file_id}")
            
            shared_files = self.db.get_shared_files_for_user(self.current_user)
            file_info = next((f for f in shared_files if f['file_id'] == file_id), None)
            
            if not file_info:
                messagebox.showerror("Error", "❌ Selected file not found")
                return False
            
            # Try to get plain AES key first
            if "aes_key_plain" in file_info and file_info["aes_key_plain"]:
                aes_key = file_info["aes_key_plain"]
                print(f"DEBUG: Found plain AES key in database: {aes_key[:30]}...")
                self.download_aes_key_entry.delete(0, tk.END)
                self.download_aes_key_entry.insert(0, aes_key)
                self.update_status(f"🔑 AES key retrieved from database for: {file_info['filename']}")
                messagebox.showinfo("Success", "✅ AES key retrieved from database successfully!")
                return True
            
            # If no plain key, try RSA decryption
            elif "encrypted_aes_key" in file_info and file_info["encrypted_aes_key"]:
                encrypted_aes_key_b64 = file_info["encrypted_aes_key"]
                print(f"DEBUG: Found encrypted AES key (first 50 chars): {encrypted_aes_key_b64[:50]}...")
                
                # Try to decrypt with user's RSA private key
                if self.user_private_key:
                    print("DEBUG: User private key available, attempting RSA decryption...")
                    decrypted_aes_key, message = self.decrypt_aes_key_with_rsa(encrypted_aes_key_b64)
                    if decrypted_aes_key:
                        print(f"DEBUG: RSA decryption successful! Decrypted key: {decrypted_aes_key[:30]}...")
                        self.download_aes_key_entry.delete(0, tk.END)
                        self.download_aes_key_entry.insert(0, decrypted_aes_key.decode())
                        self.update_status(f"🔑 AES key decrypted with RSA for: {file_info['filename']}")
                        messagebox.showinfo("Success", "✅ AES key decrypted successfully using your RSA private key!")
                        return True
                    else:
                        print(f"DEBUG: RSA decryption failed: {message}")
                        messagebox.showerror("Error", f"❌ Failed to decrypt AES key: {message}")
                        return False
                else:
                    print("DEBUG: User private key not loaded")
                    messagebox.showerror("Error", "❌ RSA private key not loaded. Please login again.")
                    return False
            else:
                print("DEBUG: No AES key found in database (neither plain nor encrypted)")
                messagebox.showerror("Error", "❌ No AES key found in database for this file")
                return False
            
        except Exception as e:
            print(f"DEBUG: Error in decrypt_shared_file_key: {type(e).__name__}: {e}")
            messagebox.showerror("Error", f"❌ Failed to get decryption key: {str(e)}")
            return False
    
    def encrypt_aes_key(self):
        """Encrypt AES key using current user's RSA public key"""
        aes_key_str = self.aes_key_to_encrypt.get().strip()
        if not aes_key_str:
            messagebox.showerror("Error", "❌ Please enter an AES key to encrypt")
            return
        
        if not self.user_public_key:
            messagebox.showerror("Error", "❌ User public key not loaded")
            return
        
        try:
            aes_key = aes_key_str.encode()
            
            # Encrypt with user's own public key (for testing/demo)
            encrypted_key_b64, message = self.encrypt_aes_key_with_rsa(aes_key, self.user_public_key)
            if not encrypted_key_b64:
                messagebox.showerror("Error", f"❌ Key encryption failed: {message}")
                return
            
            self.key_result_text.config(state=tk.NORMAL)
            self.key_result_text.delete(1.0, tk.END)
            self.key_result_text.insert(tk.END, encrypted_key_b64)
            self.key_result_text.config(state=tk.DISABLED)
            self.update_status("🔒 AES key encrypted successfully with your RSA public key")
            
        except Exception as e:
            messagebox.showerror("Error", f"❌ Key encryption failed: {str(e)}")
    
    def decrypt_aes_key(self):
        """Decrypt AES key using current user's RSA private key"""
        encrypted_key_b64 = self.aes_key_to_encrypt.get().strip()
        if not encrypted_key_b64:
            messagebox.showerror("Error", "❌ Please enter an encrypted key to decrypt")
            return
        
        try:
            # Decrypt with user's private key
            decrypted_key, message = self.decrypt_aes_key_with_rsa(encrypted_key_b64)
            if not decrypted_key:
                messagebox.showerror("Error", f"❌ Key decryption failed: {message}")
                return
            
            self.key_result_text.config(state=tk.NORMAL)
            self.key_result_text.delete(1.0, tk.END)
            self.key_result_text.insert(tk.END, decrypted_key.decode())
            self.key_result_text.config(state=tk.DISABLED)
            self.update_status("🔓 AES key decrypted successfully with your RSA private key")
            
        except Exception as e:
            messagebox.showerror("Error", f"❌ Key decryption failed: {str(e)}")
    
    def refresh_logs(self):
        """Refresh and display activity logs"""
        try:
            with open('file_sharing.log', 'r') as log_file:
                logs = log_file.read()
            
            self.logs_text.config(state=tk.NORMAL)
            self.logs_text.delete(1.0, tk.END)
            self.logs_text.insert(tk.END, logs)
            self.logs_text.see(tk.END)
            self.logs_text.config(state=tk.DISABLED)
            
        except FileNotFoundError:
            self.logs_text.config(state=tk.NORMAL)
            self.logs_text.delete(1.0, tk.END)
            self.logs_text.insert(tk.END, "No logs found")
            self.logs_text.config(state=tk.DISABLED)
    
    def update_status(self, message):
        """Update status display"""
        self.status_text.config(state=tk.NORMAL)
        timestamp = datetime.now().strftime('%H:%M:%S')
        self.status_text.insert(tk.END, f"[{timestamp}] {message}\n")
        self.status_text.see(tk.END)
        self.status_text.config(state=tk.DISABLED)
    
    def sync_database_to_json(self):
        """Sync all users from database to JSON file"""
        try:
            success = self.json_manager.sync_from_database(self.db)
            if success:
                print("✅ Users synchronized from database to JSON file!")
                return True
            else:
                print("❌ Failed to synchronize users to JSON")
                return False
        except Exception as e:
            print(f"❌ Synchronization failed: {str(e)}")
            return False
    
    def sync_json_to_database(self):
        """Sync users from JSON file to database"""
        try:
            json_users = self.json_manager.get_all_users()
            imported_count = 0
            skipped_count = 0
            
            for user_data in json_users:
                email = user_data["email"]
                
                # Check if user already exists in database
                existing_user = self.db.get_user(email)
                if not existing_user:
                    # Add user to database
                    success = self.db.add_user(
                        email=email,
                        password_hash=user_data["password_hash"],
                        role=user_data["role"],
                        security_question=user_data["security_question"],
                        security_answer_hash=user_data["security_answer_hash"]
                    )
                    if success:
                        imported_count += 1
                    else:
                        skipped_count += 1
            
            messagebox.showinfo("Import Complete", 
                              f"✅ Imported {imported_count} users from JSON to database.\n"
                              f"⚠️ Skipped {skipped_count} users (already exist).")
            return True
        except Exception as e:
            messagebox.showerror("Error", f"❌ Import failed: {str(e)}")
            return False
    
    def export_users_json(self):
        """Export users to a JSON file"""
        filename = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
            initialfile=f"users_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        )
        
        if filename:
            try:
                if self.json_manager.export_to_file(filename):
                    messagebox.showinfo("Success", f"✅ Users exported to:\n{filename}")
                else:
                    messagebox.showerror("Error", "Failed to export users to JSON")
            except Exception as e:
                messagebox.showerror("Error", f"Export failed: {str(e)}")
    
    def import_users_json(self):
        """Import users from a JSON file"""
        filename = filedialog.askopenfilename(
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
        )
        
        if filename:
            if messagebox.askyesno("Confirm Import", 
                                 "Import users from JSON file?\n\n"
                                 "Note: This will add new users to the system.\n"
                                 "Existing users with same email will be skipped."):
                try:
                    success = self.json_manager.import_from_file(filename)
                    if success:
                        messagebox.showinfo("Success", "✅ Users imported from JSON file!")
                        # Optionally sync to database
                        if messagebox.askyesno("Sync to Database", 
                                              "Do you want to sync imported users to the database?"):
                            self.sync_json_to_database()
                    else:
                        messagebox.showerror("Error", "Failed to import users from JSON")
                except Exception as e:
                    messagebox.showerror("Error", f"Import failed: {str(e)}")
    
    def get_json_stats(self):
        """Get statistics about JSON file"""
        try:
            users = self.json_manager.load_users()
            return {
                "total_users": len(users),
                "last_updated": datetime.fromtimestamp(os.path.getmtime("users.json")).strftime('%Y-%m-%d %H:%M:%S') if os.path.exists("users.json") else "Never"
            }
        except Exception as e:
            return {"total_users": 0, "last_updated": "Unknown"}

def main():
    root = tk.Tk()
    
    # Set application icon (if available)
    try:
        root.iconbitmap('icon.ico')
    except:
        pass
    
    app = SecureFileSharingApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()